import asyncio
import os
import logging
import sqlite3
import json


COVERAGE_TEMPLATE_PATH = "Шаблон подтверждение охватов.docx"

# ═══════════════════════════════════════════════════════════════
# БЛОК: ГЕНЕРАЦИЯ ПОДТВЕРЖДЕНИЙ ОХВАТОВ (встроен в bot.py)
# Источник: coverage_confirmation.py
# ═══════════════════════════════════════════════════════════════

import copy
import random
from datetime import date as _date, timedelta as _timedelta

# ── Базы имён (Краснодарский край) ────────────────────────────

_SURNAMES_MALE = [
    "Иванов","Петров","Сидоров","Козлов","Новиков","Морозов","Попов","Волков",
    "Соколов","Михайлов","Кузнецов","Лебедев","Семёнов","Егоров","Фёдоров",
    "Дорофеев","Щербаков","Тимошенко","Бондаренко","Харченко","Ткаченко",
    "Лысенко","Пономарёв","Мельник","Кириленко","Горбаченко","Шевченко",
    "Савченко","Руденко","Ковалёв","Кубанов","Донской","Черноморец",
    "Лиманский","Кубаненко","Криворучко","Хлопяников","Дубовик","Стрига","Голота",
    # Армянские (много в Краснодарском крае)
    "Акопян","Саркисян","Григорян","Арутюнян","Карапетян",
    "Мкртчян","Петросян","Авакян","Абрамян","Казарян",
    # Адыгейские
    "Хуако","Бгане","Шхалахов","Наурзоков","Тхагушев",
    "Хашба","Агумаа","Чамоков","Жане","Хакунов",
    # Греческие (Сочи)
    "Папандопулос","Христофоров","Ставрополис",
]

_SURNAMES_FEMALE_MAP = {
    "Иванов":"Иванова","Петров":"Петрова","Сидоров":"Сидорова",
    "Козлов":"Козлова","Новиков":"Новикова","Морозов":"Морозова",
    "Попов":"Попова","Волков":"Волкова","Соколов":"Соколова",
    "Михайлов":"Михайлова","Кузнецов":"Кузнецова","Лебедев":"Лебедева",
    "Семёнов":"Семёнова","Егоров":"Егорова","Фёдоров":"Фёдорова",
    "Дорофеев":"Дорофеева","Щербаков":"Щербакова","Тимошенко":"Тимошенко",
    "Бондаренко":"Бондаренко","Харченко":"Харченко","Ткаченко":"Ткаченко",
    "Лысенко":"Лысенко","Пономарёв":"Пономарёва","Мельник":"Мельник",
    "Кириленко":"Кириленко","Горбаченко":"Горбаченко","Шевченко":"Шевченко",
    "Савченко":"Савченко","Руденко":"Руденко","Ковалёв":"Ковалёва",
    "Кубанов":"Кубанова","Донской":"Донская","Черноморец":"Черноморец",
    "Лиманский":"Лиманская","Кубаненко":"Кубаненко","Криворучко":"Криворучко",
    "Хлопяников":"Хлопяникова","Дубовик":"Дубовик","Стрига":"Стрига","Голота":"Голота",
    "Акопян":"Акопян","Саркисян":"Саркисян","Григорян":"Григорян",
    "Арутюнян":"Арутюнян","Карапетян":"Карапетян","Мкртчян":"Мкртчян",
    "Петросян":"Петросян","Авакян":"Авакян","Абрамян":"Абрамян","Казарян":"Казарян",
    "Хуако":"Хуако","Бгане":"Бгане","Шхалахов":"Шхалахова","Наурзоков":"Наурзокова",
    "Тхагушев":"Тхагушева","Хашба":"Хашба","Агумаа":"Агумаа","Чамоков":"Чамокова",
    "Жане":"Жане","Хакунов":"Хакунова","Папандопулос":"Папандопулос",
    "Христофоров":"Христофорова","Ставрополис":"Ставрополис",
}

_FIRST_M = [
    "Александр","Алексей","Андрей","Артём","Борис","Вадим","Василий","Виктор",
    "Владимир","Владислав","Геннадий","Георгий","Григорий","Даниил","Денис",
    "Дмитрий","Евгений","Иван","Игорь","Илья","Кирилл","Константин","Максим",
    "Михаил","Никита","Николай","Павел","Роман","Руслан","Сергей","Тимур",
    "Фёдор","Филипп","Юрий","Яков","Степан","Гавриил","Макар","Терентий","Кузьма",
]
_FIRST_F = [
    "Александра","Алина","Алиса","Анастасия","Анна","Валентина","Валерия",
    "Вероника","Виктория","Галина","Дарья","Диана","Екатерина","Елена","Елизавета",
    "Ирина","Карина","Кристина","Ксения","Людмила","Маргарита","Марина","Мария",
    "Надежда","Наталья","Нина","Оксана","Ольга","Полина","Светлана","София",
    "Тамара","Татьяна","Юлия","Яна","Пелагея","Ефросинья","Прасковья","Глафира",
]
_PATR_M = [
    "Александрович","Алексеевич","Андреевич","Борисович","Васильевич",
    "Викторович","Владимирович","Дмитриевич","Евгеньевич","Иванович",
    "Игоревич","Ильич","Константинович","Максимович","Михайлович","Николаевич",
    "Павлович","Романович","Сергеевич","Юрьевич","Степанович","Фёдорович",
    "Петрович","Анатольевич",
]
_PATR_F = [p[:-2] + "на" for p in _PATR_M]


def _cov_gen_person(age_min: int, age_max: int, rng: random.Random) -> tuple:
    """Генерирует (ФИО, дата_рождения_DD.MM.YYYY) для возрастного диапазона."""
    is_male = rng.random() < 0.45
    base    = rng.choice(_SURNAMES_MALE)
    surname = base if is_male else _SURNAMES_FEMALE_MAP.get(base, base + "а")
    first   = rng.choice(_FIRST_M if is_male else _FIRST_F)
    patr    = rng.choice(_PATR_M  if is_male else _PATR_F)

    today   = _date.today()
    latest  = _date(today.year - age_min, today.month, today.day) - _timedelta(days=1)
    earliest = _date(today.year - age_max - 1, today.month, today.day) + _timedelta(days=1)
    delta   = (latest - earliest).days
    bday    = earliest + _timedelta(days=rng.randint(0, max(delta, 0)))

    return f"{surname} {first} {patr}", bday.strftime("%d.%m.%Y")


def _cov_gen_participants(cnt14: int, cnt18: int, cnt30: int, seed: int) -> list:
    """Возвращает список (ФИО, ДР) для всех групп участников."""
    rng = random.Random(seed)
    people = []
    for _ in range(cnt14): people.append(_cov_gen_person(14, 17, rng))
    for _ in range(cnt18): people.append(_cov_gen_person(18, 29, rng))
    for _ in range(cnt30): people.append(_cov_gen_person(30, 35, rng))
    return people


def _cov_fill_template(template_path: str, event_name: str,
                        participants: list, output_path: str) -> str:
    """
    Заполняет шаблон подтверждения охватов:
    - заменяет НАЗВАНИЕ_МЕРОПРИЯТИЯ в заголовке
    - удаляет исходные строки таблицы и клонирует первую для каждого участника
    """
    doc = Document(template_path)

    # Заголовок
    for para in doc.paragraphs:
        if "НАЗВАНИЕ_МЕРОПРИЯТИЯ" in para.text:
            full = "".join(r.text for r in para.runs)
            new  = full.replace("НАЗВАНИЕ_МЕРОПРИЯТИЯ", event_name)
            if para.runs:
                para.runs[0].text = new
                for r in para.runs[1:]: r.text = ""
            else:
                para.text = new

    if not doc.tables:
        doc.save(output_path)
        return output_path

    table    = doc.tables[0]
    tmpl_row = copy.deepcopy(table.rows[0]._tr)   # образец первой строки
    tbl_xml  = table._tbl

    # Удаляем все существующие строки
    for row in list(table.rows):
        tbl_xml.remove(row._tr)

    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    for full_name, birth_date in participants:
        new_tr = copy.deepcopy(tmpl_row)
        cells  = new_tr.findall(f"{{{ns}}}tc")

        def _set(cell_xml, text: str) -> None:
            paras = cell_xml.findall(f".//{{{ns}}}p")
            if not paras:
                return
            # Очищаем все runs
            for r in paras[0].findall(f".//{{{ns}}}r"):
                for t in r.findall(f"{{{ns}}}t"):
                    t.text = ""
            runs = paras[0].findall(f".//{{{ns}}}r")
            if runs:
                t_els = runs[0].findall(f"{{{ns}}}t")
                if t_els:
                    t_els[0].text = text
                    return
            # Нет runs — создаём
            import lxml.etree as _et
            run = _et.SubElement(paras[0], f"{{{ns}}}r")
            t   = _et.SubElement(run, f"{{{ns}}}t")
            t.text = text

        if len(cells) >= 2:
            _set(cells[0], full_name)
            _set(cells[1], birth_date)

        tbl_xml.append(new_tr)

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    return output_path


def generate_coverage_confirmations(
    events: list,
    club_name: str,
    month_year: str,
    output_dir: str = "temp",
    template_path: str = COVERAGE_TEMPLATE_PATH,
) -> list:
    """
    Генерирует документы «Охваты НАЗВАНИЕ_МЕРОПРИЯТИЯ.docx» для каждого
    мероприятия с ненулевым охватом. Возвращает список путей к файлам.
    """
    if not os.path.exists(template_path):
        logger.error(f"[COVERAGE] Template not found: {template_path}")
        return []

    os.makedirs(output_dir, exist_ok=True)
    paths = []

    for idx, event in enumerate(events):
        cnt14  = int(event.get("age_14_17", 0))
        cnt18  = int(event.get("age_18_29", 0))
        cnt30  = int(event.get("age_30_35", 0))
        total  = cnt14 + cnt18 + cnt30

        if total == 0:
            logger.info(f"[COVERAGE] Skip event #{idx+1} — zero participants")
            continue

        event_name   = event.get("event_name", f"Мероприятие {idx+1}")
        seed         = hash(f"{club_name}|{month_year}|{event_name}") & 0xFFFFFFFF
        participants = _cov_gen_participants(cnt14, cnt18, cnt30, seed)

        # Имя файла: «Охваты НАЗВАНИЕ_МЕРОПРИЯТИЯ.docx»
        # Обрезаем до 60 символов чтобы не превышать ограничения ФС
        safe_name    = event_name[:60].strip()
        filename     = f"Охваты {safe_name}.docx"
        output_path  = os.path.join(output_dir, filename)

        try:
            _cov_fill_template(template_path, event_name, participants, output_path)
            paths.append(output_path)
            logger.info(f"[COVERAGE] OK: {filename} ({total} participants)")
        except Exception as exc:
            logger.error(f"[COVERAGE] Error for '{event_name}': {exc}", exc_info=True)

    logger.info(f"[COVERAGE] {len(paths)}/{len(events)} docs generated")
    return paths


# ═══════════════════════════════════════════════════════════════
# КОНЕЦ БЛОКА: ГЕНЕРАЦИЯ ПОДТВЕРЖДЕНИЙ ОХВАТОВ
# ═══════════════════════════════════════════════════════════════
from typing import Dict, List, Tuple, Optional
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
import re
from datetime import datetime, timedelta
from calendar import monthrange, monthcalendar

from aiogram import Bot, Dispatcher, types, F
from aiogram.filters import Command
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from aiogram.fsm.storage.memory import MemoryStorage
from aiogram.types import FSInputFile, InlineKeyboardMarkup, InlineKeyboardButton
from aiogram.utils.keyboard import InlineKeyboardBuilder

import html

import io
import zipfile

# Настройка логирования
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Инициализация бота и диспетчера
BOT_TOKEN = os.getenv('BOT_TOKEN', '8286412023:AAEV_VKdmz_vYI93BBOr5wdCoeVynYhkOTM')
if not BOT_TOKEN:
    raise RuntimeError("Установите переменную окружения BOT_TOKEN")
bot = Bot(token=BOT_TOKEN)
storage = MemoryStorage()
dp = Dispatcher(storage=storage)

# Создаем директории если их нет
TEMPLATES_DIR = 'templates'
TEMP_DIR = 'temp'
os.makedirs(TEMPLATES_DIR, exist_ok=True)
os.makedirs(TEMP_DIR, exist_ok=True)


def resolve_file_path(stored_path: str):
    """
    Возвращает реальный путь к файлу с учётом возможного переноса папки.
    Порядок:
    1. Оригинальный путь (если бот не переносился).
    2. Нормализация разделителей Windows ↔ macOS/Linux.
    3. Поиск по имени файла в temp/.
    """
    if not stored_path:
        return None
    # 1. Оригинальный путь
    if os.path.exists(stored_path):
        return stored_path
    # 2. Нормализация разделителей
    normalized = stored_path.replace('\\', '/')
    if os.path.exists(normalized):
        return normalized
    # 3. Имя файла в temp/
    filename = os.path.basename(stored_path.replace('\\', '/'))
    candidate = os.path.join(TEMP_DIR, filename)
    if os.path.exists(candidate):
        logger.info(f"[PATH] Resolved '{stored_path}' \u2192 '{candidate}'")
        return candidate
    logger.warning(f"[PATH] File not found: '{stored_path}'")
    return None

# ID администратора (замените на ваш ID)
ADMIN_IDS = [531452225, 348263993]  # Замените на ваш Telegram ID

# Стандартные ссылки для планов
STANDARD_LINKS = {
    "telegram_molod": "https://t.me/molod_sochi",
    "telegram_crm": "https://t.me/crm_sochi"
}

# Данные клубов
CLUBS_DATA = {
    "Новое время": {"position": "Главный специалист", "responsible": "Калинина Т.Г."},
    "МТБ": {"position": "Ведущий специалист", "responsible": "Яицкий М.В."},
    "Брейкинг-лаборатория": {"position": "Ведущий специалист", "responsible": "Мацукатова Ю.И."},
    "В игре": {"position": "Ведущий специалист", "responsible": "Евтушенко Н.А."},
    "Вершина": {"position": "Специалист 2 категории", "responsible": "Камалетдинова Е.В."},
    "Нити": {"position": "Ведущий специалист", "responsible": "Чесных Н.В."},
    "КВН": {"position": "Ведущий специалист", "responsible": "Рыбальченко О.В."},
    "Лабиринт": {"position": "Ведущий специалист", "responsible": "Диденко И.В."},
    "Молодая семья": {"position": "Специалист 1 категории", "responsible": "Бабалаева К.В."},
    "Серпантин": {"position": "Специалист на общественных началах", "responsible": "Красноруцкий А.В."},
    "Точка опоры": {"position": "Ведущий специалист", "responsible": "Бутковская С.А."},
    "Траверс": {"position": "Ведущий специалист", "responsible": "Пластамак А.К."},
    "Черноморские дварфы": {"position": "Специалист на общественных началах", "responsible": "Головишин С.С."},
    "Амплитуда": {"position": "Ведущий специалист", "responsible": "Семенцова Л.Р."},
    "ФКФ Сочи": {"position": "Специалист на общественных началах", "responsible": "Сошинский Ю.А."}
} 

# Новые направления молодежной политики
YOUTH_POLICY_DIRECTIONS = {
    "direction_1": "Воспитание гражданственности, патриотизма, преемственности традиций, уважения к отечественной истории, историческим, национальным и иным традициям народов Российской Федерации",
    "direction_2": "Обеспечение межнационального (межэтнического) и межконфессионального согласия в молодежной среде, профилактика и предупреждение проявлений экстремизма в деятельности молодежных объединений",
    "direction_3": "Поддержка молодых граждан, оказавшихся в трудной жизненной ситуации, инвалидов из числа молодых граждан, а также лиц из числа детей-сирот и детей, оставшихся без попечения родителей",
    "direction_4": "Поддержка инициатив молодежи",
    "direction_5": "Содействие общественной деятельности, направленной на поддержку молодежи",
    "direction_6": "Организация досуга, отдыха, оздоровления молодежи, формирование условий для занятий физической культурой, спортом, содействие здоровому образу жизни молодежи",
    "direction_7": "Предоставление социальных услуг молодежи",
    "direction_8": "Содействие решению жилищных проблем молодежи, молодых семей",
    "direction_9": "Поддержка молодых семей",
    "direction_10": "Содействие образованию молодежи, научной, научно-технической деятельности молодежи",
    "direction_11": "Организация подготовки специалистов по работе с молодежью",
    "direction_12": "Выявление, сопровождение и поддержка молодежи, проявившей одаренность",
    "direction_13": "Развитие института наставничества",
    "direction_14": "Обеспечение гарантий в сфере труда и занятоимости молодежи, содействие трудоустройству молодых граждан, в том числе посредством студенческих отрядов, профессиональному развитию молодых специалистов",
    "direction_15": "Поддержка и содействие предпринимательской деятельности молодежи",
    "direction_16": "Поддержка деятельности молодежных общественных объединений",
    "direction_17": "Содействие участию молодежи в добровольческой (волонтерской) деятельности",
    "direction_18": "Содействие международному и межрегиональному сотрудничеству в сфере молодежной политики",
    "direction_19": "Предупреждение правонарушений и антиобщественных действий молодежи",
    "direction_20": "Поддержка деятельности по созданию и распространению, в том числе в информационно-телекоммуникационной сети \"Интернет\", в средствах массовой информации произведений науки, искусства, литературы и других произведений, направленных на укрепление гражданской идентичности и духовно-нравственных ценностей молодежи",
    "direction_21": "Проведение научно-аналитических исследований по вопросам молодежной политики"
}

# Статусы мероприятий
EVENT_STATUSES = [
    "Клубное",
    "Поселенческое", 
    "Муниципальное",
    "Зональное",
    "Краевое",
    "Федеральное",
    "Международное"
]

class ReportStates(StatesGroup):
    waiting_for_month = State()
    waiting_for_club_selection = State()
    waiting_for_event_name = State()
    waiting_for_event_dates = State()
    waiting_for_event_place = State()
    waiting_for_event_time = State()
    waiting_for_description = State()
    waiting_for_link = State()
    waiting_for_age_14_17 = State()
    waiting_for_age_18_29 = State()
    waiting_for_age_30_35 = State()
    waiting_for_more_events = State()
    editing_event = State()

class PlanStates(StatesGroup):
    waiting_for_plan_month = State()
    waiting_for_plan_club_selection = State()
    waiting_for_direction_selection = State()
    waiting_for_event_name = State()
    waiting_for_event_status = State()
    waiting_for_event_dates = State()
    waiting_for_event_place = State()
    waiting_for_participants_count = State()
    waiting_for_info_link = State()
    waiting_for_more_plan_items = State()
    editing_plan_item = State()

class QuarterlyStates(StatesGroup):
    waiting_for_quarter_selection = State()
    waiting_for_quarter_year = State()
    waiting_for_quarter_club_selection = State()

class AdminStates(StatesGroup):
    waiting_for_admin_action = State()
    waiting_for_club_selection = State()
    waiting_for_document_type = State()
    waiting_for_year_selection = State()
    waiting_for_month_selection = State()
    waiting_for_get_docs_type = State()
    waiting_for_get_docs_year = State()
    waiting_for_get_docs_month = State()
    waiting_for_events_period_start = State()
    waiting_for_events_period_end = State()
    waiting_for_consolidated_month = State()  # Новое: выбор месяца для сводного плана
    waiting_for_consolidated_year = State()   # Новое: выбор года для сводного плана

class CreativeReportStates(StatesGroup):
    waiting_for_club             = State()
    waiting_for_template         = State()
    waiting_for_month            = State()
    waiting_for_overwrite_confirm = State()   # Подтверждение перезаписи существующего отчёта
    waiting_for_place            = State()
    waiting_for_time             = State()
    waiting_for_age_14_17        = State()
    waiting_for_age_18_29        = State()
    waiting_for_age_30_35        = State()
    waiting_for_content          = State()
    waiting_for_links            = State()
    waiting_for_photos           = State()
    # Редактирование готового отчёта
    waiting_for_draft_choice     = State()   # Выбор: продолжить или начать заново
    editing_field                = State()   # Ввод нового значения поля при редактировании

def get_documents_by_period(document_type: str, year: int, month: str = None) -> List[Tuple]:
    """Получение всех документов за указанный период"""
    conn = sqlite3.connect('reports_plans.db')
    cursor = conn.cursor()
    
    if document_type == 'reports':
        table_name = 'reports'
    elif document_type == 'plans':
        table_name = 'plans'
    else:
        return []
    
    if month and month != 'all':
        # Для конкретного месяца
        month_year_str = f"{month} {year}"
        cursor.execute(f'''
            SELECT club_name, file_path, 
                   {'events_count' if document_type == 'reports' else 'plan_items_count'} as items_count,
                   total_participants, created_at 
            FROM {table_name} 
            WHERE month_year = ?
            ORDER BY club_name, created_at DESC
        ''', (month_year_str,))
    else:
        # За весь год
        cursor.execute(f'''
            SELECT club_name, file_path, 
                   {'events_count' if document_type == 'reports' else 'plan_items_count'} as items_count,
                   total_participants, created_at 
            FROM {table_name} 
            WHERE month_year LIKE ?
            ORDER BY club_name, created_at DESC
        ''', (f'%{year}%',))
    
    documents = cursor.fetchall()
    conn.close()
    return documents

# Добавить новую клавиатуру для выбора типа документов в админке
def get_admin_get_docs_type_keyboard() -> InlineKeyboardMarkup:
    """Создает клавиатуру для выбора типа документов при запросе"""
    builder = InlineKeyboardBuilder()
    
    builder.add(
        InlineKeyboardButton(text="📊 Все отчеты", callback_data="get_docs_reports"),
        InlineKeyboardButton(text="📋 Все планы", callback_data="get_docs_plans"),
        InlineKeyboardButton(text="📚 Все документы", callback_data="get_docs_all")
    )
    builder.adjust(1)
    return builder.as_markup()

# Инициализация базы данных
def init_database():
    """Инициализация базы данных"""
    conn = sqlite3.connect('reports_plans.db')
    cursor = conn.cursor()
    
    # Таблица для отчетов
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS reports (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            club_name TEXT,
            month_year TEXT,
            file_path TEXT,
            events_count INTEGER,
            total_participants INTEGER,
            events_json TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    # Добавляем колонку events_json если её нет (для существующих БД)
    try:
        cursor.execute("ALTER TABLE reports ADD COLUMN events_json TEXT")
    except Exception:
        pass

    # Таблица для планов
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS plans (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            club_name TEXT,
            month_year TEXT,
            file_path TEXT,
            plan_items_count INTEGER,
            total_participants INTEGER,
            plan_items_json TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    # Добавляем колонку plan_items_json если её нет (для существующих БД)
    try:
        cursor.execute("ALTER TABLE plans ADD COLUMN plan_items_json TEXT")
    except Exception:
        pass
    
    # Новая таблица для квартальных отчетов
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS quarterly_reports (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            club_name TEXT,
            quarter_year TEXT,
            file_path TEXT,
            events_count INTEGER,
            total_participants INTEGER,
            included_months TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    
    # Таблица для черновиков отчетов
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS report_drafts (
            user_id INTEGER PRIMARY KEY,
            club_name TEXT,
            month_year TEXT,
            events_data TEXT,
            current_event_index INTEGER DEFAULT 0,
            selected_dates TEXT,
            template_path TEXT,
            position TEXT,
            responsible TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    
    # Таблица для черновиков планов
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS plan_drafts (
            user_id INTEGER PRIMARY KEY,
            club_name TEXT,
            month_year TEXT,
            plan_items_data TEXT,
            current_plan_index INTEGER DEFAULT 0,
            position TEXT,
            responsible TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    
    # Таблица для творческих отчётов
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS creative_reports (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            club_name TEXT,
            month_year TEXT,
            template_label TEXT,
            file_path TEXT,
            total_participants INTEGER DEFAULT 0,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')

    # Черновики творческих отчётов
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS creative_drafts (
            user_id INTEGER PRIMARY KEY,
            state_json TEXT NOT NULL,
            current_step TEXT NOT NULL,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')

    # Маппинг Telegram user_id → клуб для отправки уведомлений специалистам.
    # Специалист регистрируется через /register и получает напоминания только по своему клубу.
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS specialist_users (
            user_id INTEGER PRIMARY KEY,
            club_name TEXT NOT NULL,
            username TEXT,
            full_name TEXT,
            notifications_enabled INTEGER DEFAULT 1,
            registered_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')

    # Настройки расписания напоминаний (хранятся как JSON-запись).
    # Позволяет администратору менять дни/время напоминаний прямо из бота.
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS reminder_settings (
            id INTEGER PRIMARY KEY CHECK (id = 1),
            settings_json TEXT NOT NULL,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    # Дефолтные настройки, если записи ещё нет
    cursor.execute('''
        INSERT OR IGNORE INTO reminder_settings (id, settings_json)
        VALUES (1, ?)
    ''', (json.dumps({
        "specialist_day":   20,   # 20-е число — напоминание специалистам
        "specialist_hour":  10,
        "specialist_minute": 0,
        "admin_day":        25,   # 25-е число — итоговое напоминание администратору
        "admin_hour":       10,
        "admin_minute":     0,
    }, ensure_ascii=False),))

    conn.commit()
    conn.close()

# Функции для работы с черновиками отчетов
def save_report_draft(user_id: int, data: Dict):
    """Сохраняет черновик отчета"""
    conn = sqlite3.connect('reports_plans.db')
    cursor = conn.cursor()
    
    cursor.execute('''
        INSERT OR REPLACE INTO report_drafts 
        (user_id, club_name, month_year, events_data, current_event_index, selected_dates, template_path, position, responsible, updated_at)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, CURRENT_TIMESTAMP)
    ''', (
        user_id,
        data.get('club_name'),
        data.get('month_year'),
        json.dumps(data.get('events', [])),
        data.get('current_event_index', 0),
        json.dumps(data.get('selected_dates', [])),
        data.get('template_path'),
        data.get('position'),
        data.get('responsible')
    ))
    
    conn.commit()
    conn.close()

def get_report_draft(user_id: int) -> Optional[Dict]:
    """Получает черновик отчета"""
    conn = sqlite3.connect('reports_plans.db')
    cursor = conn.cursor()
    
    cursor.execute('''
        SELECT club_name, month_year, events_data, current_event_index, selected_dates, template_path, position, responsible
        FROM report_drafts WHERE user_id = ?
    ''', (user_id,))
    
    result = cursor.fetchone()
    conn.close()
    
    if result:
        try:
            events = json.loads(result[2]) if result[2] else []
        except (json.JSONDecodeError, TypeError):
            events = []
        try:
            selected_dates = json.loads(result[4]) if result[4] else []
        except (json.JSONDecodeError, TypeError):
            selected_dates = []
        return {
            'club_name': result[0],
            'month_year': result[1],
            'events': events,
            'current_event_index': result[3],
            'selected_dates': selected_dates,
            'template_path': result[5],
            'position': result[6],
            'responsible': result[7]
        }
    return None

def delete_report_draft(user_id: int):
    """Удаляет черновик отчета"""
    conn = sqlite3.connect('reports_plans.db')
    cursor = conn.cursor()
    
    cursor.execute('DELETE FROM report_drafts WHERE user_id = ?', (user_id,))
    conn.commit()
    conn.close()

# Функции для работы с черновиками планов
def save_plan_draft(user_id: int, data: Dict):
    """Сохраняет черновик плана"""
    conn = sqlite3.connect('reports_plans.db')
    cursor = conn.cursor()
    
    cursor.execute('''
        INSERT OR REPLACE INTO plan_drafts 
        (user_id, club_name, month_year, plan_items_data, current_plan_index, position, responsible, updated_at)
        VALUES (?, ?, ?, ?, ?, ?, ?, CURRENT_TIMESTAMP)
    ''', (
        user_id,
        data.get('club_name'),
        data.get('month_year'),
        json.dumps(data.get('plan_items', [])),
        data.get('current_plan_index', 0),
        data.get('position'),
        data.get('responsible')
    ))
    
    conn.commit()
    conn.close()

def get_plan_draft(user_id: int) -> Optional[Dict]:
    """Получает черновик плана"""
    conn = sqlite3.connect('reports_plans.db')
    cursor = conn.cursor()
    
    cursor.execute('''
        SELECT club_name, month_year, plan_items_data, current_plan_index, position, responsible
        FROM plan_drafts WHERE user_id = ?
    ''', (user_id,))
    
    result = cursor.fetchone()
    conn.close()
    
    if result:
        try:
            plan_items = json.loads(result[2]) if result[2] else []
        except (json.JSONDecodeError, TypeError):
            plan_items = []
        return {
            'club_name': result[0],
            'month_year': result[1],
            'plan_items': plan_items,
            'current_plan_index': result[3],
            'position': result[4],
            'responsible': result[5]
        }
    return None

def delete_plan_draft(user_id: int):
    """Удаляет черновик плана"""
    conn = sqlite3.connect('reports_plans.db')
    cursor = conn.cursor()
    
    cursor.execute('DELETE FROM plan_drafts WHERE user_id = ?', (user_id,))
    conn.commit()
    conn.close()

# ── Функции для творческих отчётов ────────────────────────────────────

def save_creative_report_to_db(
    user_id: int,
    club_name: str,
    month_year: str,
    template_label: str,
    file_path: str,
    total_participants: int,
) -> None:
    """Сохраняет запись о творческом отчёте в БД."""
    conn = sqlite3.connect('reports_plans.db')
    cursor = conn.cursor()
    cursor.execute(
        '''INSERT INTO creative_reports
           (user_id, club_name, month_year, template_label, file_path, total_participants)
           VALUES (?, ?, ?, ?, ?, ?)''',
        (user_id, club_name, month_year, template_label, file_path, total_participants),
    )
    conn.commit()
    conn.close()


def get_existing_creative_report(
    club_name: str,
    template_label: str,
    month_year: str,
) -> Optional[Tuple]:
    """
    Проверяет, существует ли творческий отчёт для данного клуба,
    шаблона и месяца. Возвращает (file_path, created_at) или None.
    """
    conn = sqlite3.connect('reports_plans.db')
    cursor = conn.cursor()
    cursor.execute(
        '''SELECT file_path, created_at FROM creative_reports
           WHERE club_name = ? AND template_label = ? AND month_year = ?
           ORDER BY created_at DESC LIMIT 1''',
        (club_name, template_label, month_year),
    )
    row = cursor.fetchone()
    conn.close()
    return row  # (file_path, created_at) или None


def delete_creative_report_from_db(
    club_name: str,
    template_label: str,
    month_year: str,
) -> None:
    """Удаляет запись о творческом отчёте из БД и файл с диска."""
    conn = sqlite3.connect('reports_plans.db')
    cursor = conn.cursor()
    cursor.execute(
        '''SELECT file_path FROM creative_reports
           WHERE club_name = ? AND template_label = ? AND month_year = ?''',
        (club_name, template_label, month_year),
    )
    rows = cursor.fetchall()
    for (file_path,) in rows:
        if file_path and os.path.exists(file_path):
            try:
                os.remove(file_path)
                logger.info(f"[CR] Deleted file: {file_path}")
            except Exception as exc:
                logger.warning(f"[CR] Could not delete file {file_path}: {exc}")
    cursor.execute(
        '''DELETE FROM creative_reports
           WHERE club_name = ? AND template_label = ? AND month_year = ?''',
        (club_name, template_label, month_year),
    )
    conn.commit()
    conn.close()


# Функции черновиков творческих отчётов
def cr_draft_save(user_id: int, state_data: dict, current_step: str) -> None:
    """Сохраняет черновик творческого отчёта в БД.
    Фотографии (байты) не сериализуются — пользователь будет переотправлять их заново."""
    # Фильтруем несериализуемые поля (байты фото)
    safe_data = {k: v for k, v in state_data.items() if not isinstance(v, (bytes, list))}
    safe_data.pop('photos', None)
    conn = sqlite3.connect('reports_plans.db')
    cursor = conn.cursor()
    cursor.execute(
        '''INSERT INTO creative_drafts (user_id, state_json, current_step, updated_at)
           VALUES (?, ?, ?, CURRENT_TIMESTAMP)
           ON CONFLICT(user_id) DO UPDATE SET
               state_json = excluded.state_json,
               current_step = excluded.current_step,
               updated_at = CURRENT_TIMESTAMP''',
        (user_id, json.dumps(safe_data, ensure_ascii=False), current_step),
    )
    conn.commit()
    conn.close()


def cr_draft_load(user_id: int) -> Optional[Tuple[dict, str]]:
    """Загружает черновик. Возвращает (state_dict, current_step) или None."""
    conn = sqlite3.connect('reports_plans.db')
    cursor = conn.cursor()
    cursor.execute(
        'SELECT state_json, current_step FROM creative_drafts WHERE user_id = ?',
        (user_id,),
    )
    row = cursor.fetchone()
    conn.close()
    if not row:
        return None
    try:
        return json.loads(row[0]), row[1]
    except json.JSONDecodeError:
        return None


def cr_draft_delete(user_id: int) -> None:
    """Удаляет черновик пользователя."""
    conn = sqlite3.connect('reports_plans.db')
    cursor = conn.cursor()
    cursor.execute('DELETE FROM creative_drafts WHERE user_id = ?', (user_id,))
    conn.commit()
    conn.close()


def cr_draft_has(user_id: int) -> bool:
    """True если у пользователя есть незавершённый черновик."""
    conn = sqlite3.connect('reports_plans.db')
    cursor = conn.cursor()
    cursor.execute('SELECT 1 FROM creative_drafts WHERE user_id = ?', (user_id,))
    result = cursor.fetchone()
    conn.close()
    return result is not None


def get_creative_reports_by_period(year: int, month: str = None) -> List[Tuple]:
    """
    Возвращает список творческих отчётов за период.
    Каждая строка: (club_name, file_path, total_participants, created_at, template_label)
    """
    conn = sqlite3.connect('reports_plans.db')
    cursor = conn.cursor()
    if month and month != 'all':
        cursor.execute(
            '''SELECT club_name, file_path, total_participants, created_at, template_label
               FROM creative_reports
               WHERE month_year = ?
               ORDER BY club_name, created_at DESC''',
            (f"{month} {year}",),
        )
    else:
        cursor.execute(
            '''SELECT club_name, file_path, total_participants, created_at, template_label
               FROM creative_reports
               WHERE month_year LIKE ?
               ORDER BY club_name, created_at DESC''',
            (f"%{year}%",),
        )
    rows = cursor.fetchall()
    conn.close()
    return rows


# Остальные функции базы данных
def save_report_to_db(user_id: int, club_name: str, month_year: str, file_path: str, events_count: int, total_participants: int, events: list = None):
    """Сохранение отчёта в БД вместе с данными мероприятий для редактирования."""
    conn = sqlite3.connect('reports_plans.db')
    cursor = conn.cursor()
    events_json = json.dumps(events or [], ensure_ascii=False)
    cursor.execute('''
        INSERT INTO reports (user_id, club_name, month_year, file_path, events_count, total_participants, events_json)
        VALUES (?, ?, ?, ?, ?, ?, ?)
    ''', (user_id, club_name, month_year, file_path, events_count, total_participants, events_json))
    conn.commit()
    conn.close()

def save_plan_to_db(user_id: int, club_name: str, month_year: str, file_path: str, plan_items_count: int, total_participants: int, plan_items: list = None):
    """Сохранение плана в БД вместе с данными пунктов для редактирования."""
    conn = sqlite3.connect('reports_plans.db')
    cursor = conn.cursor()
    plan_items_json = json.dumps(plan_items or [], ensure_ascii=False)
    cursor.execute('''
        INSERT INTO plans (user_id, club_name, month_year, file_path, plan_items_count, total_participants, plan_items_json)
        VALUES (?, ?, ?, ?, ?, ?, ?)
    ''', (user_id, club_name, month_year, file_path, plan_items_count, total_participants, plan_items_json))
    conn.commit()
    conn.close()

def save_quarterly_report_to_db(user_id: int, club_name: str, quarter_year: str, file_path: str, events_count: int, total_participants: int, included_months: str):
    """Сохранение квартального отчета в базу данных"""
    conn = sqlite3.connect('reports_plans.db')
    cursor = conn.cursor()
    
    cursor.execute('''
        INSERT INTO quarterly_reports (user_id, club_name, quarter_year, file_path, events_count, total_participants, included_months)
        VALUES (?, ?, ?, ?, ?, ?, ?)
    ''', (user_id, club_name, quarter_year, file_path, events_count, total_participants, included_months))
    
    conn.commit()
    conn.close()

def get_existing_report(user_id: int, club_name: str, month_year: str) -> Optional[Tuple]:
    """Проверяет существование отчёта в базе данных.
    Ищет по клубу и периоду (без привязки к user_id) — любой специалист клуба видит общий отчёт."""
    conn = sqlite3.connect('reports_plans.db')
    cursor = conn.cursor()
    cursor.execute(
        '''SELECT file_path, events_count, total_participants, created_at
           FROM reports WHERE club_name = ? AND month_year = ?
           ORDER BY created_at DESC LIMIT 1''',
        (club_name, month_year),
    )
    result = cursor.fetchone()
    conn.close()
    return result

def get_existing_plan(user_id: int, club_name: str, month_year: str) -> Optional[Tuple]:
    """Проверяет существование плана в базе данных.
    Ищет по клубу и периоду (без привязки к user_id)."""
    conn = sqlite3.connect('reports_plans.db')
    cursor = conn.cursor()
    cursor.execute(
        '''SELECT file_path, plan_items_count, total_participants, created_at
           FROM plans WHERE club_name = ? AND month_year = ?
           ORDER BY created_at DESC LIMIT 1''',
        (club_name, month_year),
    )
    result = cursor.fetchone()
    conn.close()
    return result


def get_existing_report_events(user_id: int, club_name: str, month_year: str) -> Optional[list]:
    """Загружает список мероприятий из сохранённого отчёта (для редактирования)."""
    conn = sqlite3.connect('reports_plans.db')
    cursor = conn.cursor()
    cursor.execute(
        'SELECT events_json FROM reports WHERE club_name=? AND month_year=? ORDER BY created_at DESC LIMIT 1',
        (club_name, month_year),
    )
    row = cursor.fetchone()
    conn.close()
    if not row or not row[0]:
        return None
    try:
        return json.loads(row[0])
    except Exception:
        return None


def get_existing_plan_items(user_id: int, club_name: str, month_year: str) -> Optional[list]:
    """Загружает список пунктов из сохранённого плана (для редактирования)."""
    conn = sqlite3.connect('reports_plans.db')
    cursor = conn.cursor()
    cursor.execute(
        'SELECT plan_items_json FROM plans WHERE club_name=? AND month_year=? ORDER BY created_at DESC LIMIT 1',
        (club_name, month_year),
    )
    row = cursor.fetchone()
    conn.close()
    if not row or not row[0]:
        return None
    try:
        return json.loads(row[0])
    except Exception:
        return None


def get_status_for_club(club_name: str, month_year: str) -> dict:
    """
    Возвращает полный статус документов клуба за период.
    Не зависит от user_id — ищет по клубу, чтобы показывать актуальную картину
    независимо от того кто последний заполнял.

    Возвращает dict:
        report:   (events_count, total_participants, created_at) | None
        plan:     (plan_items_count, total_participants, created_at) | None
        creative: [(template_label, created_at), ...]  — может быть несколько
        draft:    True если есть незавершённый черновик творческого отчёта
    """
    conn = sqlite3.connect('reports_plans.db')
    cursor = conn.cursor()

    cursor.execute(
        '''SELECT events_count, total_participants, created_at FROM reports
           WHERE club_name = ? AND month_year = ?
           ORDER BY created_at DESC LIMIT 1''',
        (club_name, month_year),
    )
    report_row = cursor.fetchone()

    cursor.execute(
        '''SELECT plan_items_count, total_participants, created_at FROM plans
           WHERE club_name = ? AND month_year = ?
           ORDER BY created_at DESC LIMIT 1''',
        (club_name, month_year),
    )
    plan_row = cursor.fetchone()

    cursor.execute(
        '''SELECT template_label, created_at FROM creative_reports
           WHERE club_name = ? AND month_year = ?
           ORDER BY template_label''',
        (club_name, month_year),
    )
    creative_rows = cursor.fetchall()

    conn.close()
    return {
        "report":   report_row,
        "plan":     plan_row,
        "creative": creative_rows,
    }

def delete_report_from_db(user_id: int, club_name: str, month_year: str):
    """Удаляет отчет из базы данных"""
    conn = sqlite3.connect('reports_plans.db')
    cursor = conn.cursor()
    
    # Сначала получаем путь к файлу
    cursor.execute('''
        SELECT file_path FROM reports 
        WHERE user_id = ? AND club_name = ? AND month_year = ?
    ''', (user_id, club_name, month_year))
    
    result = cursor.fetchone()
    if result:
        file_path = result[0]
        # Удаляем файл
        if os.path.exists(file_path):
            try:
                os.remove(file_path)
            except Exception as e:
                logger.error(f"Error deleting file {file_path}: {e}")
    
    # Удаляем запись из базы данных
    cursor.execute('''
        DELETE FROM reports 
        WHERE user_id = ? AND club_name = ? AND month_year = ?
    ''', (user_id, club_name, month_year))
    
    conn.commit()
    conn.close()

def delete_plan_from_db(user_id: int, club_name: str, month_year: str):
    """Удаляет план из базы данных"""
    conn = sqlite3.connect('reports_plans.db')
    cursor = conn.cursor()
    
    # Сначала получаем путь к файлу
    cursor.execute('''
        SELECT file_path FROM plans 
        WHERE user_id = ? AND club_name = ? AND month_year = ?
    ''', (user_id, club_name, month_year))
    
    result = cursor.fetchone()
    if result:
        file_path = result[0]
        # Удаляем файл
        if os.path.exists(file_path):
            try:
                os.remove(file_path)
            except Exception as e:
                logger.error(f"Error deleting file {file_path}: {e}")
    
    # Удаляем запись из базы данных
    cursor.execute('''
        DELETE FROM plans 
        WHERE user_id = ? AND club_name = ? AND month_year = ?
    ''', (user_id, club_name, month_year))
    
    conn.commit()
    conn.close()

def get_club_reports(club_name: str, year: int = None, month: str = None) -> List[Tuple]:
    """Получение всех отчетов клуба"""
    conn = sqlite3.connect('reports_plans.db')
    cursor = conn.cursor()
    
    if year and month:
        cursor.execute('''
            SELECT month_year, file_path, events_count, total_participants, created_at 
            FROM reports 
            WHERE club_name = ? AND month_year LIKE ?
            ORDER BY created_at DESC
        ''', (club_name, f'{month} {year}'))
    elif year:
        cursor.execute('''
            SELECT month_year, file_path, events_count, total_participants, created_at 
            FROM reports 
            WHERE club_name = ? AND month_year LIKE ?
            ORDER BY created_at DESC
        ''', (club_name, f'%{year}%'))
    else:
        cursor.execute('''
            SELECT month_year, file_path, events_count, total_participants, created_at 
            FROM reports 
            WHERE club_name = ? 
            ORDER BY created_at DESC
        ''', (club_name,))
    
    reports = cursor.fetchall()
    conn.close()
    return reports

def get_club_plans(club_name: str, year: int = None, month: str = None) -> List[Tuple]:
    """Получение всех планов клуба"""
    conn = sqlite3.connect('reports_plans.db')
    cursor = conn.cursor()
    
    if year and month:
        cursor.execute('''
            SELECT month_year, file_path, plan_items_count, total_participants, created_at 
            FROM plans 
            WHERE club_name = ? AND month_year LIKE ?
            ORDER BY created_at DESC
        ''', (club_name, f'{month} {year}'))
    elif year:
        cursor.execute('''
            SELECT month_year, file_path, plan_items_count, total_participants, created_at 
            FROM plans 
            WHERE club_name = ? AND month_year LIKE ?
            ORDER BY created_at DESC
        ''', (club_name, f'%{year}%'))
    else:
        cursor.execute('''
            SELECT month_year, file_path, plan_items_count, total_participants, created_at 
            FROM plans 
            WHERE club_name = ? 
            ORDER BY created_at DESC
        ''', (club_name,))
    
    plans = cursor.fetchall()
    conn.close()
    return plans

def get_all_clubs_with_stats(year: int = None, month: str = None) -> List[Tuple]:
    """Получение статистики по всем клубам"""
    conn = sqlite3.connect('reports_plans.db')
    cursor = conn.cursor()
    
    if year and month:
        month_year_str = f"{month} {year}"
        cursor.execute('''
            SELECT club_name, 
                   COUNT(DISTINCT month_year) as reports_count,
                   (SELECT COUNT(DISTINCT month_year) FROM plans WHERE plans.club_name = reports.club_name AND plans.month_year = ?) as plans_count
            FROM reports 
            WHERE month_year = ?
            GROUP BY club_name
            UNION
            SELECT club_name, 
                   0 as reports_count,
                   COUNT(DISTINCT month_year) as plans_count
            FROM plans 
            WHERE club_name NOT IN (SELECT DISTINCT club_name FROM reports WHERE month_year = ?) 
                  AND month_year = ?
            GROUP BY club_name
        ''', (month_year_str, month_year_str, month_year_str, month_year_str))
    elif year:
        cursor.execute('''
            SELECT club_name, 
                   COUNT(DISTINCT month_year) as reports_count,
                   (SELECT COUNT(DISTINCT month_year) FROM plans WHERE plans.club_name = reports.club_name AND plans.month_year LIKE ?) as plans_count
            FROM reports 
            WHERE month_year LIKE ?
            GROUP BY club_name
            UNION
            SELECT club_name, 
                   0 as reports_count,
                   COUNT(DISTINCT month_year) as plans_count
            FROM plans 
            WHERE club_name NOT IN (SELECT DISTINCT club_name FROM reports WHERE month_year LIKE ?) 
                  AND month_year LIKE ?
            GROUP BY club_name
        ''', (f'%{year}%', f'%{year}%', f'%{year}%', f'%{year}%'))
    else:
        cursor.execute('''
            SELECT club_name, 
                   COUNT(DISTINCT month_year) as reports_count,
                   (SELECT COUNT(DISTINCT month_year) FROM plans WHERE plans.club_name = reports.club_name) as plans_count
            FROM reports 
            GROUP BY club_name
            UNION
            SELECT club_name, 
                   0 as reports_count,
                   COUNT(DISTINCT month_year) as plans_count
            FROM plans 
            WHERE club_name NOT IN (SELECT DISTINCT club_name FROM reports)
            GROUP BY club_name
        ''')
    
    stats = cursor.fetchall()
    conn.close()
    return stats

def get_clubs_statistics(year: int, month: str = None) -> Dict:
    """Получает статистику по клубам за указанный год и месяц"""
    conn = sqlite3.connect('reports_plans.db')
    cursor = conn.cursor()
    
    statistics = {}
    all_clubs = list(CLUBS_DATA.keys())
    
    if month:
        month_year_str = f"{month} {year}"
        
        for club_name in all_clubs:
            # Проверяем наличие отчета
            cursor.execute('''
                SELECT COUNT(*) FROM reports 
                WHERE club_name = ? AND month_year = ?
            ''', (club_name, month_year_str))
            has_report = cursor.fetchone()[0] > 0
            
            # Проверяем наличие плана
            cursor.execute('''
                SELECT COUNT(*) FROM plans 
                WHERE club_name = ? AND month_year = ?
            ''', (club_name, month_year_str))
            has_plan = cursor.fetchone()[0] > 0
            
            statistics[club_name] = {
                'report': has_report,
                'plan': has_plan
            }
    else:
        for club_name in all_clubs:
            # Проверяем наличие отчетов за год
            cursor.execute('''
                SELECT COUNT(*) FROM reports 
                WHERE club_name = ? AND month_year LIKE ?
            ''', (club_name, f'%{year}%'))
            reports_count = cursor.fetchone()[0]
            
            # Проверяем наличие планов за год
            cursor.execute('''
                SELECT COUNT(*) FROM plans 
                WHERE club_name = ? AND month_year LIKE ?
            ''', (club_name, f'%{year}%'))
            plans_count = cursor.fetchone()[0]
            
            statistics[club_name] = {
                'reports_count': reports_count,
                'plans_count': plans_count
            }
    
    conn.close()
    return statistics

def get_missing_months(club_name: str, document_type: str, year: int) -> List[str]:
    """Получение месяцев, для которых нет документов"""
    months = ['январь', 'февраль', 'март', 'апрель', 'май', 'июнь', 
              'июль', 'август', 'сентябрь', 'октябрь', 'ноябрь', 'декабрь']
    
    conn = sqlite3.connect('reports_plans.db')
    cursor = conn.cursor()
    
    if document_type == 'reports':
        cursor.execute('''
            SELECT month_year FROM reports WHERE club_name = ? AND month_year LIKE ?
        ''', (club_name, f'%{year}%'))
    else:
        cursor.execute('''
            SELECT month_year FROM plans WHERE club_name = ? AND month_year LIKE ?
        ''', (club_name, f'%{year}%'))
    
    existing_docs = [row[0] for row in cursor.fetchall()]
    conn.close()
    
    missing_months = []
    for month in months:
        month_year = f"{month} {year}"
        if month_year not in existing_docs:
            missing_months.append(month_year)
    
    return missing_months

async def send_document_to_admin(user_id: int, document_path: str, document_type: str, club_name: str, period: str, events_count: int, total_participants: int):
    """Отправляет документ администратору"""
    try:
        if os.path.exists(document_path):
            doc_file = FSInputFile(document_path)
            
            for admin_id in ADMIN_IDS:
                try:
                    await bot.send_document(
                        admin_id,
                        doc_file,
                        caption=(
                            f"📄 Новый документ сформирован!\n"
                            f"👤 Пользователь: {user_id}\n"
                            f"🏢 Клуб: {club_name}\n"
                            f"📅 Период: {period}\n"
                            f"📋 Тип: {document_type}\n"
                            f"📊 Мероприятий/пунктов: {events_count}\n"
                            f"👥 Участников: {total_participants}\n"
                            f"🕐 Время: {datetime.now().strftime('%d.%m.%Y %H:%M')}"
                        )
                    )
                except Exception as e:
                    logger.error(f"Error sending document to admin {admin_id}: {e}")
    except Exception as e:
        logger.error(f"Error in send_document_to_admin: {e}")

# Вспомогательные функции
def get_clubs_keyboard() -> InlineKeyboardMarkup:
    """Создает клавиатуру для выбора клуба"""
    builder = InlineKeyboardBuilder()

    for club_name in CLUBS_DATA.keys():
        builder.add(InlineKeyboardButton(text=club_name, callback_data=f"club_{club_name}"))

    builder.adjust(2)
    builder.row(InlineKeyboardButton(text="⬅️ Назад к выбору месяца", callback_data="back_to_month_selection"))
    return builder.as_markup()

def get_admin_clubs_keyboard() -> InlineKeyboardMarkup:
    """Создает клавиатуру для выбора клуба в админке"""
    builder = InlineKeyboardBuilder()

    for club_name in CLUBS_DATA.keys():
        builder.add(InlineKeyboardButton(text=club_name, callback_data=f"admin_club_{club_name}"))

    builder.adjust(2)
    builder.row(InlineKeyboardButton(text="⬅️ Назад", callback_data="back_to_admin_main"))
    return builder.as_markup()

def get_admin_document_type_keyboard() -> InlineKeyboardMarkup:
    """Создает клавиатуру для выбора типа документа в админке"""
    builder = InlineKeyboardBuilder()
    
    builder.add(
        InlineKeyboardButton(text="📊 Отчеты", callback_data="admin_doc_reports"),
        InlineKeyboardButton(text="📋 Планы", callback_data="admin_doc_plans")
    )
    builder.adjust(2)
    return builder.as_markup()

def get_admin_year_keyboard() -> InlineKeyboardMarkup:
    """Создает клавиатуру для выбора года"""
    builder = InlineKeyboardBuilder()

    current_year = datetime.now().year
    builder.add(
        InlineKeyboardButton(text=f"📅 {current_year - 1}", callback_data=f"admin_year_{current_year - 1}"),
        InlineKeyboardButton(text=f"📅 {current_year}", callback_data=f"admin_year_{current_year}"),
        InlineKeyboardButton(text=f"📅 {current_year + 1}", callback_data=f"admin_year_{current_year + 1}"),
        InlineKeyboardButton(text="📅 Все годы", callback_data="admin_year_all"),
        InlineKeyboardButton(text="⬅️ Назад", callback_data="back_to_admin_main"),
    )
    builder.adjust(1)
    return builder.as_markup()

def get_admin_main_keyboard() -> InlineKeyboardMarkup:
    """Создает главную клавиатуру админки"""
    builder = InlineKeyboardBuilder()

    builder.add(
        InlineKeyboardButton(text="📊 Статистика клубов", callback_data="admin_extended_stats"),
        InlineKeyboardButton(text="📁 Документы клуба", callback_data="admin_view_docs"),
        InlineKeyboardButton(text="📦 Скачать документы за период", callback_data="admin_get_docs"),
        InlineKeyboardButton(text="📅 Мероприятия за период", callback_data="admin_show_events"),
        InlineKeyboardButton(text="📋 Сформировать сводный план", callback_data="admin_consolidated_plan"),
        InlineKeyboardButton(text="🔔 Уведомления", callback_data="admin_notifications"),
        InlineKeyboardButton(text="🏠 Главное меню", callback_data="admin_back_to_start"),
    )
    builder.adjust(1)
    return builder.as_markup()

@dp.callback_query(F.data == "admin_back_to_start")
async def admin_back_to_start(callback: types.CallbackQuery, state: FSMContext):
    """Возврат из админки в главное меню"""
    await state.clear()
    await callback.message.delete()
    await cmd_start(callback.message, state)
    await callback.answer()

@dp.callback_query(F.data == "admin_consolidated_plan")
async def process_admin_consolidated_plan(callback: types.CallbackQuery, state: FSMContext):
    """Обработка формирования сводного плана"""
    user_id = callback.from_user.id
    
    if user_id not in ADMIN_IDS:
        await callback.answer("❌ Доступ запрещен")
        return
    
    current_year = datetime.now().year
    
    await callback.message.edit_text(
        "📋 Формирование сводного месячного плана\n\n"
        "Введите год (например: {0}):".format(current_year),
        reply_markup=InlineKeyboardMarkup(inline_keyboard=[[
            InlineKeyboardButton(text="⬅️ Назад", callback_data="back_to_admin_main")
        ]])
    )
    await state.set_state(AdminStates.waiting_for_consolidated_year)
    await callback.answer()

@dp.message(AdminStates.waiting_for_consolidated_year)
async def process_consolidated_year(message: types.Message, state: FSMContext):
    """Обработка ввода года для сводного плана"""
    user_id = message.from_user.id
    
    if user_id not in ADMIN_IDS:
        await message.answer("❌ Доступ запрещен")
        return
    
    try:
        year = int(message.text.strip())
        if year < 2020 or year > 2030:
            await message.answer("❌ Введите корректный год (2020-2030):")
            return
        
        await state.update_data(consolidated_year=year)
        
        await message.answer(
            f"✅ Год {year} сохранен!\n\n"
            f"📅 Теперь введите месяц (например: март, апрель, май):",
            reply_markup=InlineKeyboardMarkup(inline_keyboard=[[
                InlineKeyboardButton(text="⬅️ Назад", callback_data="back_to_admin_main")
            ]])
        )
        await state.set_state(AdminStates.waiting_for_consolidated_month)
        
    except ValueError:
        await message.answer("❌ Пожалуйста, введите корректный год:")

def collect_plans_for_month(month: str, year: int) -> List[Dict]:
    """Собирает все пункты планов из всех клубов за указанный месяц"""
    conn = sqlite3.connect('reports_plans.db')
    cursor = conn.cursor()
    
    month_year = f"{month} {year}"
    all_plan_items = []
    
    # Получаем все планы за указанный месяц
    cursor.execute('''
        SELECT club_name, file_path FROM plans 
        WHERE month_year = ?
        ORDER BY club_name
    ''', (month_year,))
    
    plans = cursor.fetchall()
    conn.close()
    
    for club_name, file_path in plans:
        file_path = resolve_file_path(file_path) or file_path
        if os.path.exists(file_path):
            try:
                # Парсим пункты плана из файла
                plan_items = parse_plan_items_from_file(file_path, club_name)
                all_plan_items.extend(plan_items)
                logger.info(f"Parsed {len(plan_items)} items from {club_name}")
            except Exception as e:
                logger.error(f"Error parsing plan for {club_name}: {e}")
    
    return all_plan_items

def parse_plan_items_from_file(file_path: str, club_name: str) -> List[Dict]:
    """Парсит пункты плана из файла для сводного плана"""
    plan_items = []
    
    try:
        doc = Document(file_path)
        
        # Ищем таблицу в документе
        for table in doc.tables:
            # Проверяем, что это нужная таблица (должна иметь заголовки)
            if len(table.rows) > 1:
                # Проверяем заголовки первой строки
                header_row = table.rows[0]
                header_text = ' '.join([cell.text for cell in header_row.cells])
                
                # Если это похоже на таблицу с планом
                if any(keyword in header_text.lower() for keyword in ['статус', 'наименование', 'мероприятие', 'участник']):
                    
                    # Пропускаем заголовки таблицы
                    for i, row in enumerate(table.rows):
                        if i < 1:  # Пропускаем первую строку с заголовками
                            continue
                        
                        cells = row.cells
                        if len(cells) >= 6:
                            item_data = parse_plan_row_for_consolidated(cells, i, club_name)
                            if item_data:
                                plan_items.append(item_data)
        
        # Если не нашли таблицу с планом, пробуем другой подход
        if not plan_items:
            # Просто парсим все строки таблиц
            for table in doc.tables:
                for i, row in enumerate(table.rows):
                    if i < 1:  # Пропускаем первую строку
                        continue
                    
                    cells = row.cells
                    if len(cells) >= 6:
                        item_data = parse_plan_row_for_consolidated(cells, i, club_name)
                        if item_data:
                            plan_items.append(item_data)
        
    except Exception as e:
        logger.error(f"Error parsing plan file {file_path}: {e}")
    
    return plan_items

def parse_plan_row_for_consolidated(cells, row_index: int, club_name: str) -> Optional[Dict]:
    """Парсит строку плана для сводного плана"""
    try:
        # Проверяем, что это валидная строка с мероприятием
        if len(cells) < 6:
            return None
            
        # Проверяем номер строки (должен быть числом)
        cell_0_text = cells[0].text.strip()
        if not cell_0_text or not cell_0_text.replace('.', '').replace(' ', '').isdigit():
            return None
        
        # Извлекаем данные в правильном порядке
        status = cells[3].text.strip() if len(cells) > 3 else ""  # Статус мероприятия
        event_name = cells[2].text.strip() if len(cells) > 2 else ""  # Форма и наименование
        date_place = cells[4].text.strip() if len(cells) > 4 else ""  # Дата и место
        participants = cells[5].text.strip() if len(cells) > 5 else ""  # Количество участников
        
        # Проверяем, что это не пустая строка
        if not event_name and not status:
            return None
        
        # Форматируем количество участников
        participants_text = participants
        if participants and not participants.endswith('чел.'):
            # Если есть число, добавляем "чел."
            participants_match = re.search(r'\d+', participants)
            if participants_match:
                participants_text = f"{participants_match.group()} чел."
        
        return {
            'club_name': club_name,
            'status': status,
            'event_name': event_name,
            'date_place': date_place,
            'participants': participants_text,
            'row_index': row_index
        }
        
    except Exception as e:
        logger.error(f"Error parsing plan row {row_index}: {e}")
        return None

def extract_leader_info(doc: Document, club_name: str) -> Optional[str]:
    """Извлекает информацию о руководителе из документа в нужном формате"""
    try:
        # Ищем параграфы в конце документа
        leader_text = None
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            # Ищем строки с подписью руководителя
            if text and any(keyword in text.lower() for keyword in ['ведущий', 'специалист', 'руководитель', 'начальник', 'отдел']):
                # Проверяем, что это не просто ФИО, а полная должность
                if len(text) > 10 and not text.startswith('_'):
                    leader_text = text
                    break
        
        if leader_text:
            # Формируем строку в нужном формате
            # Убираем лишние подчеркивания и пробелы
            leader_text = leader_text.replace('_', '').strip()
            # Добавляем информацию о клубе
            return f"{leader_text}, руководитель молодежного клуба \"{club_name}\""
        
        # Если не нашли, формируем из данных клуба
        if club_name in CLUBS_DATA:
            club_info = CLUBS_DATA[club_name]
            return f"{club_info['position']}, руководитель молодежного клуба \"{club_name}\" {club_info['responsible']}"
        
        return None
    except Exception as e:
        logger.error(f"Error extracting leader info: {e}")
        return None

def fill_consolidated_plan_template(template_path: str, month: str, year: int, plan_items: List[Dict]) -> str:
    """Заполнение шаблона сводного плана данными из всех клубов"""
    doc = Document(template_path)
    
    # Формируем текст для замены МЕСЯЦиГОД (с предлогом "на" и словом "года")
    # Например: "на март 2026 года"
    month_year_text = f"на {month} {year} года"
    
    for paragraph in doc.paragraphs:
        if "МЕСЯЦиГОД" in paragraph.text:
            # Заменяем маркер на нужный текст
            paragraph.text = paragraph.text.replace("МЕСЯЦиГОД", month_year_text)
            # Устанавливаем стиль (без жирного)
            set_paragraph_style(paragraph, 14, False, WD_PARAGRAPH_ALIGNMENT.CENTER)
    
    # Собираем информацию о руководителях для каждого клуба
    leader_info_cache = {}
    
    # Заполняем таблицу мероприятиями
    for table in doc.tables:
        total_rows = len(table.rows)
        items_count = len(plan_items)
        
        # Удаляем лишние строки если нужно
        if items_count < total_rows - 1 and total_rows > items_count + 1:
            # Удаляем с конца, чтобы индексы не сбивались
            for i in range(total_rows - 1, items_count, -1):
                if i < len(table.rows):
                    try:
                        tbl = table._tbl
                        tr = table.rows[i]._tr
                        tbl.remove(tr)
                    except Exception as e:
                        logger.error(f"Error removing row {i}: {e}")
        
        # Заполняем данные
        for i, item in enumerate(plan_items):
            row_index = i + 1  # +1 потому что первая строка - заголовок
            if row_index < len(table.rows):
                cells = table.rows[row_index].cells
                
                if len(cells) < 6:
                    continue
                
                # № п/п
                cells[0].text = str(i + 1)
                set_cell_style(cells[0], 12)
                
                # Статус мероприятия (1.1.1, 1.2.1, etc)
                cells[1].text = item.get('status', '')
                set_cell_style(cells[1], 12)
                
                # Форма и наименование мероприятия (1.1.3, 1.2.3, etc)
                cells[2].text = item.get('event_name', '')
                set_cell_style(cells[2], 12)
                
                # Дата и место проведения (1.1.4, 1.2.4, etc)
                cells[3].text = item.get('date_place', '')
                set_cell_style(cells[3], 12)
                
                # Ответственный специалист (1.1.5, 1.2.5, etc)
                club_name = item.get('club_name', '')
                
                # Используем кэш для хранения информации о руководителях
                if club_name not in leader_info_cache:
                    # Формируем информацию о руководителе в нужном формате
                    if club_name in CLUBS_DATA:
                        club_info = CLUBS_DATA[club_name]
                        # Формат: "Должность, руководитель молодежного клуба "Название клуба" ФИО"
                        leader_info_cache[club_name] = f"{club_info['position']}, руководитель молодежного клуба \"{club_name}\" {club_info['responsible']}"
                    else:
                        leader_info_cache[club_name] = club_name
                
                cells[4].text = leader_info_cache[club_name]
                set_cell_style(cells[4], 12)
                
                # Предполагаемое количество участников (1.1.6, 1.2.6, etc)
                cells[5].text = item.get('participants', '')
                set_cell_style(cells[5], 12)
    
    # Сохраняем файл
    safe_month = "".join(c for c in month if c.isalnum())
    filename = f"Сводный план на {safe_month} {year}.docx"
    output_path = os.path.join(TEMP_DIR, filename)
    
    doc.save(output_path)
    return output_path

@dp.message(AdminStates.waiting_for_consolidated_month)
async def process_consolidated_month(message: types.Message, state: FSMContext):
    """Обработка ввода месяца и формирование сводного плана"""
    user_id = message.from_user.id
    
    if user_id not in ADMIN_IDS:
        await message.answer("❌ Доступ запрещен")
        return
    
    month = message.text.strip().lower()
    
    # Проверяем корректность месяца
    valid_months = ['январь', 'февраль', 'март', 'апрель', 'май', 'июнь',
                    'июль', 'август', 'сентябрь', 'октябрь', 'ноябрь', 'декабрь']
    
    if month not in valid_months:
        await message.answer(
            "❌ Неверный месяц. Введите месяц в именительном падеже:\n"
            "например: январь, февраль, март и т.д."
        )
        return
    
    data = await state.get_data()
    year = data.get('consolidated_year')
    
    if not year:
        await message.answer("❌ Ошибка: не найден год")
        await state.clear()
        return
    
    await message.answer(f"📋 Собираю данные из планов за {month} {year} года...")
    
    try:
        # Собираем все пункты планов
        plan_items = collect_plans_for_month(month, year)
        
        if not plan_items:
            await message.answer(
                f"❌ Не найдено планов за {month} {year} года",
                reply_markup=get_admin_main_keyboard()
            )
            await state.set_state(AdminStates.waiting_for_admin_action)
            return
        
        # Показываем статистику
        clubs_represented = set(item['club_name'] for item in plan_items)
        stats_text = (
            f"📊 Найдено данных:\n"
            f"• Клубов: {len(clubs_represented)}\n"
            f"• Мероприятий: {len(plan_items)}\n\n"
            f"Участвуют клубы:\n"
        )
        for club in sorted(clubs_represented):
            stats_text += f"• {club}\n"
        
        await message.answer(stats_text)
        await message.answer("🔄 Формирую сводный план...")
        
        # Формируем сводный план
        template_path = "Шаблон сводного плана.docx"
        if not os.path.exists(template_path):
            await message.answer("❌ Шаблон сводного плана не найден")
            await state.set_state(AdminStates.waiting_for_admin_action)
            return
        
        output_path = fill_consolidated_plan_template(template_path, month, year, plan_items)
        
        # Отправляем файл
        doc_file = FSInputFile(output_path)
        await message.answer_document(
            doc_file,
            caption=(
                f"✅ Сводный план сформирован!\n"
                f"📅 Период: {month} {year}\n"
                f"🏢 Клубов: {len(clubs_represented)}\n"
                f"📋 Мероприятий: {len(plan_items)}"
            )
        )
        
        # Возвращаем в главное меню
        await message.answer(
            "Выберите дальнейшее действие:",
            reply_markup=get_admin_main_keyboard()
        )
        await state.set_state(AdminStates.waiting_for_admin_action)
        
    except Exception as e:
        logger.error(f"Error generating consolidated plan: {e}")
        await message.answer(
            f"❌ Произошла ошибка: {str(e)}",
            reply_markup=get_admin_main_keyboard()
        )
        await state.set_state(AdminStates.waiting_for_admin_action)

def normalize_text(text: str) -> str:
    """Нормализует текст для вставки в документ"""
    if not text:
        return ""
    # Удаляем лишние пробелы и переносы строк
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    return '\n'.join(lines)

@dp.callback_query(F.data == "admin_show_events")
async def process_admin_show_events(callback: types.CallbackQuery, state: FSMContext):
    """Обработка показа мероприятий за период"""
    user_id = callback.from_user.id
    
    if user_id not in ADMIN_IDS:
        await callback.answer("❌ Доступ запрещен")
        return
    
    await callback.message.edit_text(
        "📅 Показать мероприятия за период\n\n"
        "Введите начальную дату в формате ДД.ММ.ГГГГ\n"
        "Например: 02.02.2026",
        reply_markup=InlineKeyboardMarkup(inline_keyboard=[[
            InlineKeyboardButton(text="⬅️ Назад", callback_data="back_to_admin_main")
        ]])
    )
    await state.set_state(AdminStates.waiting_for_events_period_start)
    await callback.answer()

@dp.callback_query(F.data == "back_to_admin_main")
async def back_to_admin_main(callback: types.CallbackQuery, state: FSMContext):
    """Возврат в главное меню админки"""
    await callback.message.edit_text(
        "👋 Добро пожаловать в админ-панель!\n\n"
        "Выберите действие:",
        reply_markup=get_admin_main_keyboard()
    )
    await state.set_state(AdminStates.waiting_for_admin_action)
    await callback.answer()

def get_events_from_plans_by_period(start_date: datetime, end_date: datetime) -> List[Dict]:
    """Получает мероприятия из всех планов за указанный период (обновленная для текстовых дат)"""
    conn = sqlite3.connect('reports_plans.db')
    cursor = conn.cursor()
    
    # Получаем все планы
    cursor.execute('SELECT club_name, month_year, file_path FROM plans ORDER BY created_at DESC')
    all_plans = cursor.fetchall()
    
    events_by_club = {}
    
    for club_name, month_year, file_path in all_plans:
        try:
            file_path = resolve_file_path(file_path) or file_path
            if not os.path.exists(file_path):
                continue
                
            # Парсим мероприятия из плана
            plan_events = parse_events_from_plan_file(file_path, club_name, month_year)
            
            # Фильтруем мероприятия по дате
            for event in plan_events:
                event_dates = event.get('dates_obj_list', [])
                
                # Проверяем каждую дату мероприятия
                for event_date in event_dates:
                    if start_date <= event_date <= end_date:
                        if club_name not in events_by_club:
                            events_by_club[club_name] = []
                        
                        # Клонируем событие для этой даты
                        event_copy = event.copy()
                        event_copy['date'] = event_date.strftime('%d.%m.%Y')
                        event_copy['date_obj'] = event_date
                        event_copy['club_name'] = club_name
                        event_copy['month_year'] = month_year
                        
                        # Добавляем в список
                        events_by_club[club_name].append(event_copy)
                        
        except Exception as e:
            logger.error(f"Error processing plan for {club_name}, {month_year}: {e}")
            continue
    
    conn.close()
    
    # Преобразуем в плоский список и сортируем по дате
    all_events = []
    for club_name, events in events_by_club.items():
        all_events.extend(events)
    
    # Сортируем по дате
    all_events.sort(key=lambda x: x.get('date_obj', datetime.min))
    
    return all_events

def debug_date_parsing(text: str):
    """Отладочная функция для проверки парсинга дат"""
    logger.info(f"Parsing text: {text}")
    
    # Пробуем все методы парсинга
    dates1 = _extract_dates_from_line(text)
    dates2 = _extract_dates_other_formats(text)
    dates3 = extract_all_dates_from_plan_text(text)
    
    logger.info(f"Method 1 found: {[d.strftime('%d.%m.%Y') for d in dates1]}")
    logger.info(f"Method 2 found: {[d.strftime('%d.%m.%Y') for d in dates2]}")
    logger.info(f"Method 3 found: {[d.strftime('%d.%m.%Y') for d in dates3]}")
    
    place = _extract_place_from_text_excluding_dates(text)
    logger.info(f"Extracted place: {place}")

def parse_events_from_plan_file(file_path: str, club_name: str, month_year: str) -> List[Dict]:
    """Парсит мероприятия из файла плана"""
    events = []
    
    try:
        doc = Document(file_path)
        
        for table in doc.tables:
            # Пропускаем заголовки таблицы
            for i, row in enumerate(table.rows):
                if i < 1:  # Пропускаем первую строку с заголовками
                    continue
                
                cells = row.cells
                if len(cells) >= 7:
                    event_data = _parse_plan_event_row(cells, i)
                    if event_data:
                        events.append(event_data)
        
    except Exception as e:
        logger.error(f"Error parsing plan file {file_path}: {e}")
    
    return events


def _parse_plan_event_row(cells, row_index: int) -> Optional[Dict]:
    """Парсит строку с мероприятием из плана (обновленная для текстовых дат)"""
    try:
        # Проверяем, что это валидная строка с мероприятием
        if len(cells) < 7 or not cells[0].text.strip().replace('.', '').isdigit():
            return None
        
        # Извлекаем данные
        direction = cells[1].text.strip()
        event_name = cells[2].text.strip()
        status = cells[3].text.strip()
        date_place = cells[4].text.strip()
        participants_text = cells[5].text.strip()
        info_link = cells[6].text.strip()
        
        # Извлекаем даты с помощью улучшенной функции
        date_objs = extract_all_dates_from_plan_text(date_place)
        
        # Парсим количество участников
        participants_match = re.search(r'(\d+)', participants_text)
        participants = int(participants_match.group(1)) if participants_match else 0
        
        # Извлекаем место проведения (все, что не дата)
        place = _extract_place_from_text_excluding_dates(date_place)
        
        if not event_name:
            return None
        
        # Преобразуем datetime в строку для отображения
        dates_str = [date.strftime('%d.%m.%Y') for date in date_objs]
        
        return {
            'event_name': event_name,
            'date': dates_str[0] if dates_str else None,
            'date_obj': date_objs[0] if date_objs else None,
            'dates_list': dates_str,
            'dates_obj_list': date_objs,
            'place': place,
            'participants': participants,
            'direction': direction,
            'status': status,
            'info_link': info_link,
            'raw_date_text': date_place  # Сохраняем исходный текст для отладки
        }
        
    except Exception as e:
        logger.error(f"Error parsing plan row {row_index}: {e}")
        logger.error(f"Cell contents: {[cell.text for cell in cells]}")
        return None

def _extract_place_from_text_excluding_dates(text: str) -> str:
    """Извлекает место проведения, исключая строки с датами"""
    lines = text.split('\n')
    place_lines = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Проверяем, содержит ли строка дату
        has_date = False
        
        # Проверяем основные форматы дат
        patterns_to_check = [
            r'\d{1,2}\s+[а-я]+\s+\d{4}(?:\s+года|\s+г\.?)?',  # "2 февраля 2026 года"
            r'\d{1,2}\.[а-я]+\.?\s+\d{4}',  # "2 февр. 2026"
            r'\d{1,2}\.\d{1,2}\.\d{4}',  # "02.02.2026"
            r'\d{1,2}\.\d{1,2}\.\d{2}',  # "02.02.26"
            r'\d{1,2}-?\d{0,2}\s+[а-я]+\s+\d{4}',  # "2-4 февраля 2026"
        ]
        
        for pattern in patterns_to_check:
            if re.search(pattern, line, re.IGNORECASE):
                has_date = True
                break
        
        # Если строка не содержит дату и не пустая, добавляем как место
        if not has_date and line:
            place_lines.append(line)
    
    # Объединяем все строки места
    place = ' '.join(place_lines).strip()
    
    # Если не нашли явного места, проверяем последнюю строку
    if not place and lines:
        last_line = lines[-1].strip()
        # Если последняя строка короткая и похожа на адрес
        if len(last_line) < 100 and not any(keyword in last_line.lower() for keyword in ['янв', 'фев', 'мар', 'апр', 'май', 'июн', 'июл', 'авг', 'сен', 'окт', 'ноя', 'дек']):
            place = last_line
    
    return place if place else "Не указано"

def _extract_dates_from_text(text: str) -> List[str]:
    """Извлекает даты из текста"""
    dates = []
    lines = text.split('\n')
    
    for line in lines:
        line = line.strip()
        # Ищем даты в формате DD.MM.YYYY
        date_matches = re.findall(r'\b\d{1,2}\.\d{1,2}\.\d{4}\b', line)
        dates.extend(date_matches)
        
        # Если не нашли, ищем в формате DD.MM.YY
        if not date_matches:
            date_matches_yy = re.findall(r'\b\d{1,2}\.\d{1,2}\.\d{2}\b', line)
            for date_str in date_matches_yy:
                try:
                    day, month, year = date_str.split('.')
                    full_year = f"20{year}" if int(year) < 100 else year
                    standardized = f"{int(day):02d}.{int(month):02d}.{full_year}"
                    dates.append(standardized)
                except:
                    dates.append(date_str)
    
    return dates


def _extract_place_from_text(text: str) -> str:
    """Извлекает место проведения из текста"""
    lines = text.split('\n')
    
    # Ищем строку с местом (обычно последняя не-дата строка)
    for line in reversed(lines):
        line = line.strip()
        if not line:
            continue
        
        # Если это не дата, возвращаем как место
        if not re.search(r'\b\d{1,2}\.\d{1,2}\.\d{4}\b', line) and \
           not re.search(r'\b\d{1,2}\.\d{1,2}\.\d{2}\b', line):
            return line
    
    return "Не указано"

@dp.message(AdminStates.waiting_for_events_period_start)
async def process_events_period_start(message: types.Message, state: FSMContext):
    """Обработка ввода начальной даты периода"""
    user_id = message.from_user.id
    
    if user_id not in ADMIN_IDS:
        await message.answer("❌ Доступ запрещен")
        return
    
    try:
        # Парсим дату
        start_date = datetime.strptime(message.text.strip(), '%d.%m.%Y')
        
        await state.update_data(events_start_date=start_date)
        
        await message.answer(
            f"✅ Начальная дата: {start_date.strftime('%d.%m.%Y')}\n\n"
            f"Теперь введите конечную дату в формате ДД.ММ.ГГГГ\n"
            f"Например: 08.02.2026",
            reply_markup=InlineKeyboardMarkup(inline_keyboard=[[
                InlineKeyboardButton(text="⬅️ Назад", callback_data="back_to_admin_main")
            ]])
        )
        await state.set_state(AdminStates.waiting_for_events_period_end)
        
    except ValueError:
        await message.answer(
            "❌ Неверный формат даты. Введите в формате ДД.ММ.ГГГГ\n"
            "Например: 02.02.2026",
            reply_markup=InlineKeyboardMarkup(inline_keyboard=[[
                InlineKeyboardButton(text="⬅️ Назад", callback_data="back_to_admin_main")
            ]])
        )


@dp.message(AdminStates.waiting_for_events_period_end)
async def process_events_period_end(message: types.Message, state: FSMContext):
    """Обработка ввода конечной даты периода и показ мероприятий"""
    user_id = message.from_user.id
    
    if user_id not in ADMIN_IDS:
        await message.answer("❌ Доступ запрещен")
        return
    
    try:
        # Парсим дату
        end_date = datetime.strptime(message.text.strip(), '%d.%m.%Y')
        
        data = await state.get_data()
        start_date = data.get('events_start_date')
        
        if not start_date:
            await message.answer("❌ Ошибка: не найдена начальная дата")
            await state.clear()
            return
        
        # Проверяем, что конечная дата позже начальной
        if end_date < start_date:
            await message.answer("❌ Конечная дата должна быть позже начальной")
            return
        
        await message.answer(f"🔍 Ищу мероприятия с {start_date.strftime('%d.%m.%Y')} по {end_date.strftime('%d.%m.%Y')}...")
        
        # Получаем мероприятия за период
        events = get_events_from_plans_by_period(start_date, end_date)
        
        if not events:
            await message.answer(
                f"📭 Не найдено мероприятий за период с {start_date.strftime('%d.%m.%Y')} по {end_date.strftime('%d.%m.%Y')}",
                reply_markup=get_admin_main_keyboard()
            )
            await state.set_state(AdminStates.waiting_for_admin_action)
            return
        
        # Группируем мероприятия по дате для лучшего отображения
        events_by_date = {}
        for event in events:
            event_date_str = event.get('date')
            if event_date_str:
                if event_date_str not in events_by_date:
                    events_by_date[event_date_str] = []
                events_by_date[event_date_str].append(event)
        
        # Сортируем даты
        sorted_dates = sorted(events_by_date.keys(), 
                             key=lambda x: datetime.strptime(x, '%d.%m.%Y'))
        
        total_events = len(events)
        total_participants = sum(event.get('participants', 0) for event in events)
        
        # Формируем общую статистику
        report_text = (
            f"📅 Мероприятия с {start_date.strftime('%d.%m.%Y')} по {end_date.strftime('%d.%m.%Y')}\n"
            f"📊 Всего мероприятий: {total_events}\n"
            f"👥 Общий предполагаемый охват: {total_participants} чел.\n"
            f"{'='*40}\n"
        )
        
        # Отправляем первое сообщение со статистикой
        await message.answer(report_text)
        
        # Отправляем мероприятия по датам
        for date_str in sorted_dates:
            date_events = events_by_date[date_str]
            
            # Форматируем дату для отображения
            day, month, year = date_str.split('.')
            months_ru = {
                '01': 'января', '02': 'февраля', '03': 'марта', '04': 'апреля',
                '05': 'мая', '06': 'июня', '07': 'июля', '08': 'августа',
                '09': 'сентября', '10': 'октября', '11': 'ноября', '12': 'декабря'
            }
            formatted_date = f"{int(day)} {months_ru.get(month, month)} {year}"
            
            date_text = f"📅 <b>{formatted_date}:</b>\n\n"
            
            for i, event in enumerate(date_events, 1):
                # Экранируем все текстовые поля
                club_name_escaped = html.escape(event.get('club_name', 'Не указан'))
                event_name_escaped = html.escape(event.get('event_name', 'Без названия'))
                place_escaped = html.escape(event.get('place', 'Не указано'))
                status_escaped = html.escape(event.get('status', 'Не указан'))
                participants = event.get('participants', 0)
                
                date_text += (
                    f"<b>{i}. 🏢 {club_name_escaped}</b>\n\n"
                    f"   📝 <b>{event_name_escaped}</b>\n"
                    f"   📍 <b>Место:</b> {place_escaped}\n"
                    f"   👥 <b>Охват:</b> {participants} чел.\n"
                    f"   🏷️ <b>Статус:</b> {status_escaped}\n"
                    f"   ___________________\n\n"
                )
            
            # Отправляем сообщение для этой даты
            try:
                await message.answer(date_text, parse_mode="HTML")
                await asyncio.sleep(0.5)
            except Exception as e:
                logger.error(f"Error sending date message: {e}")
                # Удаляем HTML-теги и отправляем как обычный текст
                import re
                date_text_plain = re.sub(r'<[^>]+>', '', date_text)
                await message.answer(date_text_plain)
                await asyncio.sleep(0.5)
        
        # Итоговая статистика по клубам
        clubs_summary = {}
        for event in events:
            club = event.get('club_name', 'Неизвестно')
            if club not in clubs_summary:
                clubs_summary[club] = {'events': 0, 'participants': 0}
            clubs_summary[club]['events'] += 1
            clubs_summary[club]['participants'] += event.get('participants', 0)
        
        if clubs_summary:
            summary_text = "\n📊 *Итоги по клубам:*\n\n"
            for club, stats in sorted(clubs_summary.items()):
                summary_text += f"*🏢 {club}:*\n"
                summary_text += f"   📝 *Мероприятий:* {stats['events']}\n"
                summary_text += f"   👥 *Охват:* {stats['participants']} чел.\n\n"
            
            try:
                await message.answer(summary_text, parse_mode="Markdown")
            except:
                summary_text_plain = summary_text.replace('*', '')
                await message.answer(summary_text_plain)
        
        # Возвращаем к главному меню
        await message.answer(
            "Выберите дальнейшее действие:",
            reply_markup=get_admin_main_keyboard()
        )
        await state.set_state(AdminStates.waiting_for_admin_action)
        
    except ValueError:
        await message.answer(
            "❌ Неверный формат даты. Введите в формате ДД.ММ.ГГГГ\n"
            "Например: 08.02.2026",
            reply_markup=InlineKeyboardMarkup(inline_keyboard=[[
                InlineKeyboardButton(text="⬅️ Назад", callback_data="back_to_admin_main")
            ]])
        )

def _extract_dates_from_plan_text(text: str) -> List[datetime]:
    """Извлекает даты из текста плана в формате '2 февраля 2026 года'"""
    dates = []
    
    # Разделяем текст на строки
    lines = text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Пытаемся извлечь даты из каждой строки
        date_matches = _extract_dates_from_line(line)
        dates.extend(date_matches)
    
    return dates


def _extract_dates_from_line(line: str) -> List[datetime]:
    """Извлекает даты из строки в формате '2 февраля 2026 года'"""
    dates = []
    
    # Словарь для преобразования русских названий месяцев
    months_ru = {
        'января': 1, 'февраля': 2, 'марта': 3, 'апреля': 4,
        'мая': 5, 'июня': 6, 'июля': 7, 'августа': 8,
        'сентября': 9, 'октября': 10, 'ноября': 11, 'декабря': 12,
        'январь': 1, 'февраль': 2, 'март': 3, 'апрель': 4,
        'май': 5, 'июнь': 6, 'июль': 7, 'август': 8,
        'сентябрь': 9, 'октябрь': 10, 'ноябрь': 11, 'декабрь': 12
    }
    
    # Паттерн для дат в формате "2 февраля 2026 года" или "12 февраля 2026 года"
    # Также учитываем возможные варианты: "2 февраля 2026 г.", "2 февраля 2026"
    patterns = [
        # Основной паттерн: число (1-2 цифры) + месяц (рус) + год (4 цифры) + "года" или "г."
        r'(\d{1,2})\s+([а-я]+)\s+(\d{4})(?:\s+года|\s+г\.?)?',
        # Паттерн для диапазона дат: "2-4 февраля 2026 года"
        r'(\d{1,2})(?:-(\d{1,2}))?\s+([а-я]+)\s+(\d{4})(?:\s+года|\s+г\.?)?',
        # Паттерн с сокращенными месяцами: "2 февр. 2026 года"
        r'(\d{1,2})\s+([а-я]+\.?)\s+(\d{4})(?:\s+года|\s+г\.?)?',
    ]
    
    for pattern in patterns:
        matches = re.findall(pattern, line, re.IGNORECASE)
        for match in matches:
            try:
                if len(match) == 4 and match[0] and match[1] and match[2] and match[3]:
                    # Обработка диапазона дат: "2-4 февраля 2026 года"
                    start_day = int(match[0])
                    end_day = match[1] if match[1] else None
                    month_str = match[2].lower().rstrip('.')
                    year = int(match[3])
                    
                    # Преобразуем название месяца
                    month = months_ru.get(month_str)
                    if not month:
                        # Пробуем найти без точки
                        month_str_no_dot = month_str.rstrip('.')
                        month = months_ru.get(month_str_no_dot)
                    
                    if month:
                        # Добавляем начальную дату
                        start_date = datetime(year, month, start_day)
                        dates.append(start_date)
                        
                        # Добавляем все даты диапазона, если есть конечный день
                        if end_day:
                            end_day_int = int(end_day)
                            for day in range(start_day + 1, end_day_int + 1):
                                try:
                                    range_date = datetime(year, month, day)
                                    dates.append(range_date)
                                except ValueError:
                                    continue
                
                elif len(match) == 3:
                    # Обработка одиночной даты: "2 февраля 2026 года"
                    day = int(match[0])
                    month_str = match[1].lower().rstrip('.')
                    year = int(match[2])
                    
                    # Преобразуем название месяца
                    month = months_ru.get(month_str)
                    if not month:
                        # Пробуем найти без точки
                        month_str_no_dot = month_str.rstrip('.')
                        month = months_ru.get(month_str_no_dot)
                    
                    if month:
                        date = datetime(year, month, day)
                        dates.append(date)
                        
            except (ValueError, TypeError) as e:
                logger.error(f"Error parsing date match {match}: {e}")
                continue
    
    return dates

def extract_all_dates_from_plan_text(text: str) -> List[datetime]:
    """Извлекает все даты из текста плана, включая множественные даты"""
    dates = []
    
    # Разделяем на строки и обрабатываем каждую
    lines = text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Пытаемся найти даты в строке
        line_dates = _extract_dates_from_line(line)
        dates.extend(line_dates)
        
        # Если не нашли дат в основном формате, ищем другие форматы
        if not line_dates:
            # Попробуем найти в других форматах
            other_formats_dates = _extract_dates_other_formats(line)
            dates.extend(other_formats_dates)
    
    # Удаляем дубликаты и сортируем
    unique_dates = []
    seen = set()
    for date in dates:
        if date not in seen:
            seen.add(date)
            unique_dates.append(date)
    
    unique_dates.sort()
    return unique_dates


def _extract_dates_other_formats(line: str) -> List[datetime]:
    """Извлекает даты в других форматах"""
    dates = []
    
    # 1. Формат DD.MM.YYYY
    matches = re.findall(r'\b\d{1,2}\.\d{1,2}\.\d{4}\b', line)
    for match in matches:
        try:
            date = datetime.strptime(match, '%d.%m.%Y')
            dates.append(date)
        except ValueError:
            continue
    
    # 2. Формат DD.MM.YY
    matches = re.findall(r'\b\d{1,2}\.\d{1,2}\.\d{2}\b', line)
    for match in matches:
        try:
            date = datetime.strptime(match, '%d.%m.%y')
            dates.append(date)
        except ValueError:
            continue
    
    # 3. Формат "число месяц" (предполагаем текущий год)
    months_ru = {
        'января': 1, 'февраля': 2, 'марта': 3, 'апреля': 4,
        'мая': 5, 'июня': 6, 'июля': 7, 'августа': 8,
        'сентября': 9, 'октября': 10, 'ноября': 11, 'декабря': 12,
        'январь': 1, 'февраль': 2, 'март': 3, 'апрель': 4,
        'май': 5, 'июнь': 6, 'июль': 7, 'август': 8,
        'сентябрь': 9, 'октябрь': 10, 'ноябрь': 11, 'декабрь': 12
    }
    
    pattern = r'(\d{1,2})\s+([а-я]+)(?:\.|\s|$)'
    matches = re.findall(pattern, line, re.IGNORECASE)
    for day_str, month_str in matches:
        try:
            day = int(day_str)
            month_str_lower = month_str.lower().rstrip('.')
            month = months_ru.get(month_str_lower)
            
            if month:
                # Предполагаем текущий год
                current_year = datetime.now().year
                date = datetime(current_year, month, day)
                dates.append(date)
        except ValueError:
            continue
    
    return dates

@dp.callback_query(F.data == "admin_get_docs")
async def process_admin_get_docs(callback: types.CallbackQuery, state: FSMContext):
    """Обработка получения документов за период"""
    user_id = callback.from_user.id
    
    if user_id not in ADMIN_IDS:
        await callback.answer("❌ Доступ запрещен")
        return
    
    await callback.message.edit_text(
        "📦 Получение документов за период\n\n"
        "📋 Выберите тип документов:",
        reply_markup=get_admin_get_docs_type_keyboard()
    )
    await state.set_state(AdminStates.waiting_for_get_docs_type)
    await callback.answer()

@dp.callback_query(F.data.startswith("get_docs_"))
async def process_get_docs_type(callback: types.CallbackQuery, state: FSMContext):
    """Обработка выбора типа документов для получения"""
    user_id = callback.from_user.id
    
    if user_id not in ADMIN_IDS:
        await callback.answer("❌ Доступ запрещен")
        return
    
    docs_type = callback.data.replace("get_docs_", "")
    await state.update_data(get_docs_type=docs_type)
    
    await callback.message.edit_text(
        "📅 Выберите год:",
        reply_markup=get_admin_year_keyboard()
    )
    await state.set_state(AdminStates.waiting_for_get_docs_year)
    await callback.answer()

@dp.callback_query(F.data.startswith("admin_year_"), AdminStates.waiting_for_get_docs_year)
async def process_get_docs_year(callback: types.CallbackQuery, state: FSMContext):
    """Обработка выбора года для получения документов"""
    user_id = callback.from_user.id
    
    if user_id not in ADMIN_IDS:
        await callback.answer("❌ Доступ запрещен")
        return
    
    year_str = callback.data.replace("admin_year_", "")
    
    if year_str == "all":
        await callback.answer("❌ Для получения документов нужно выбрать конкретный год")
        return
    
    try:
        year = int(year_str)
        await state.update_data(get_docs_year=year)
    except ValueError:
        await callback.answer("❌ Ошибка выбора года")
        return
    
    await callback.message.edit_text(
        f"📅 Выберите месяц для {year} года:",
        reply_markup=get_admin_month_keyboard(year)
    )
    await state.set_state(AdminStates.waiting_for_get_docs_month)
    await callback.answer()

def build_documents_zip(
    month_display: str,
    reports: List[Tuple],
    plans: List[Tuple],
    creative_reports: List[Tuple],
) -> Tuple[bytes, int, List[str]]:
    """
    Собирает все документы в один ZIP-архив в памяти.

    Структура архива:
        Отчеты/            ← обычные месячные отчёты
        Планы/             ← планы
        Творческие отчеты/ ← только если такие есть

    Возвращает: (zip_bytes, кол-во_добавленных_файлов, список_ошибок)
    """
    buf = io.BytesIO()
    added = 0
    errors: List[str] = []

    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:

        for club_name, file_path, *_ in reports:
            file_path = resolve_file_path(file_path) or file_path
            if os.path.exists(file_path):
                arcname = f"Отчеты/{os.path.basename(file_path)}"
                zf.write(file_path, arcname)
                added += 1
                logger.info(f"[ZIP] Added report: {arcname}")
            else:
                errors.append(f"📊 Отчёт не найден: {club_name}")
                logger.warning(f"[ZIP] Missing report: {file_path}")

        for club_name, file_path, *_ in plans:
            file_path = resolve_file_path(file_path) or file_path
            if os.path.exists(file_path):
                arcname = f"Планы/{os.path.basename(file_path)}"
                zf.write(file_path, arcname)
                added += 1
                logger.info(f"[ZIP] Added plan: {arcname}")
            else:
                errors.append(f"📋 План не найден: {club_name}")
                logger.warning(f"[ZIP] Missing plan: {file_path}")

        for club_name, file_path, *_ in creative_reports:
            file_path = resolve_file_path(file_path) or file_path
            if os.path.exists(file_path):
                arcname = f"Творческие отчеты/{os.path.basename(file_path)}"
                zf.write(file_path, arcname)
                added += 1
                logger.info(f"[ZIP] Added creative: {arcname}")
            else:
                errors.append(f"🎨 Творческий отчёт не найден: {club_name}")
                logger.warning(f"[ZIP] Missing creative: {file_path}")

    return buf.getvalue(), added, errors


@dp.callback_query(F.data.startswith("admin_month_"), AdminStates.waiting_for_get_docs_month)
async def process_get_docs_month(callback: types.CallbackQuery, state: FSMContext):
    """Собирает документы за период и отправляет одним ZIP-архивом."""
    user_id = callback.from_user.id

    if user_id not in ADMIN_IDS:
        await callback.answer("❌ Доступ запрещен")
        return

    month_year_str = callback.data.replace("admin_month_", "")
    parts = month_year_str.split('_')

    if len(parts) != 2:
        await callback.answer("❌ Ошибка формата")
        return

    month, year_str = parts[0], parts[1]

    try:
        year = int(year_str)
    except ValueError:
        await callback.answer("❌ Ошибка года")
        return

    data = await state.get_data()
    docs_type = data.get('get_docs_type')

    month_display = f"{month} {year}" if month != "all" else f"весь {year} год"
    await callback.message.edit_text(f"🔄 Собираю документы за {month_display}...")

    all_clubs = set(CLUBS_DATA.keys())

    # ── Собираем нужные документы из БД ─────────────────────────────
    reports:          List[Tuple] = []
    plans:            List[Tuple] = []
    creative_reports: List[Tuple] = []

    clubs_with_reports: set = set()
    clubs_with_plans:   set = set()

    if docs_type in ('reports', 'all'):
        reports = get_documents_by_period('reports', year, month)
        clubs_with_reports = {r[0] for r in reports}
        # Творческие отчёты всегда идут вместе с обычными
        creative_reports = get_creative_reports_by_period(year, month)

    if docs_type in ('plans', 'all'):
        plans = get_documents_by_period('plans', year, month)
        clubs_with_plans = {p[0] for p in plans}

    total_found = len(reports) + len(plans) + len(creative_reports)

    if total_found == 0:
        await callback.message.answer(
            f"📭 Не найдено документов за {month_display}",
        )
        await callback.message.answer(
            "Выберите дальнейшее действие:",
            reply_markup=get_admin_main_keyboard(),
        )
        await state.set_state(AdminStates.waiting_for_admin_action)
        await callback.answer()
        return

    # ── Строим ZIP в памяти ──────────────────────────────────────────
    zip_bytes, added, errors = build_documents_zip(
        month_display=month_display,
        reports=reports,
        plans=plans,
        creative_reports=creative_reports,
    )

    # ── Отправляем ZIP ───────────────────────────────────────────────
    safe_month = "".join(c for c in month_display if c.isalnum() or c in (" ", "_")).strip().replace(" ", "_")
    zip_filename = f"Документы_{safe_month}.zip"

    zip_buf = io.BytesIO(zip_bytes)
    zip_buf.name = zip_filename

    # Формируем подпись ZIP-файла
    caption_lines = [
        f"📦 Документы за {month_display}",
        f"",
        f"📊 Отчётов: {len(reports)}",
        f"📋 Планов: {len(plans)}",
    ]
    if creative_reports:
        caption_lines.append(f"🎨 Творческих отчётов: {len(creative_reports)}")
    caption_lines += [
        f"",
        f"✅ Добавлено в архив: {added}",
    ]
    if errors:
        caption_lines.append(f"⚠️ Файлов не найдено: {len(errors)}")

    try:
        zip_input = types.BufferedInputFile(zip_bytes, filename=zip_filename)
        await callback.message.answer_document(
            zip_input,
            caption="\n".join(caption_lines),
        )
    except Exception as exc:
        logger.error(f"[ZIP] Error sending zip: {exc}")
        await callback.message.answer(f"❌ Ошибка при отправке архива: {exc}")

    # ── Статистика по отсутствующим документам ───────────────────────
    summary_lines: List[str] = []

    if docs_type == 'reports':
        missing = sorted(all_clubs - clubs_with_reports)
        if missing:
            summary_lines.append(f"📊 Отчёты отсутствуют у {len(missing)} клубов:")
            summary_lines += [f"  ❌ {c}" for c in missing]
        else:
            summary_lines.append("✅ Все клубы предоставили отчёты!")

    elif docs_type == 'plans':
        missing = sorted(all_clubs - clubs_with_plans)
        if missing:
            summary_lines.append(f"📋 Планы отсутствуют у {len(missing)} клубов:")
            summary_lines += [f"  ❌ {c}" for c in missing]
        else:
            summary_lines.append("✅ Все клубы предоставили планы!")

    elif docs_type == 'all':
        missing_reports = all_clubs - clubs_with_reports
        missing_plans   = all_clubs - clubs_with_plans
        missing_both    = missing_reports & missing_plans

        total_clubs      = len(all_clubs)
        clubs_with_both  = len(all_clubs - missing_reports - missing_plans)
        completion_rate  = (clubs_with_both * 100 // total_clubs) if total_clubs else 0

        summary_lines.append("📈 Статистика по отсутствующим документам:")

        if missing_both:
            summary_lines.append(f"  ❌ Нет ни отчёта, ни плана ({len(missing_both)}):")
            summary_lines += [f"    • {c}" for c in sorted(missing_both)]

        only_no_report = missing_reports - missing_plans
        if only_no_report:
            summary_lines.append(f"  ⚠️ Нет только отчёта ({len(only_no_report)}):")
            summary_lines += [f"    • {c}" for c in sorted(only_no_report)]

        only_no_plan = missing_plans - missing_reports
        if only_no_plan:
            summary_lines.append(f"  ⚠️ Нет только плана ({len(only_no_plan)}):")
            summary_lines += [f"    • {c}" for c in sorted(only_no_plan)]

        if not missing_reports and not missing_plans:
            summary_lines.append("  ✅ Все клубы предоставили все документы!")

        summary_lines += [
            "",
            f"🏢 Всего клубов: {total_clubs}",
            f"✅ С обоими документами: {clubs_with_both}",
            f"📈 Выполнение: {completion_rate}%",
        ]

    if errors:
        summary_lines += ["", "⚠️ Файлы не найдены на диске:"]
        summary_lines += [f"  {e}" for e in errors]

    if summary_lines:
        await callback.message.answer("\n".join(summary_lines))

    await callback.message.answer(
        "Выберите дальнейшее действие:",
        reply_markup=get_admin_main_keyboard(),
    )
    await state.set_state(AdminStates.waiting_for_admin_action)
    await callback.answer()


def get_admin_month_keyboard(year: int) -> InlineKeyboardMarkup:
    """Создает клавиатуру для выбора месяца"""
    builder = InlineKeyboardBuilder()
    
    months = [
        ("Январь", "январь"),
        ("Февраль", "февраль"),
        ("Март", "март"),
        ("Апрель", "апрель"),
        ("Май", "май"),
        ("Июнь", "июнь"),
        ("Июль", "июль"),
        ("Август", "август"),
        ("Сентябрь", "сентябрь"),
        ("Октябрь", "октябрь"),
        ("Ноябрь", "ноябрь"),
        ("Декабрь", "декабрь"),
        ("📈 За весь год", "all")
    ]
    
    for month_name, month_value in months:
        builder.add(InlineKeyboardButton(text=month_name, callback_data=f"admin_month_{month_value}_{year}"))
    
    builder.adjust(3)
    return builder.as_markup()

def get_directions_keyboard() -> InlineKeyboardMarkup:
    """Создает клавиатуру для выбора направления"""
    builder = InlineKeyboardBuilder()
    
    for i in range(1, 22):  # Теперь 21 направление
        builder.add(InlineKeyboardButton(text=f"Направление {i}", callback_data=f"direction_{i}"))
    
    builder.adjust(2)  # 2 кнопки в ряду для лучшего отображения
    return builder.as_markup()

def get_status_keyboard() -> InlineKeyboardMarkup:
    """Создает клавиатуру для выбора статуса мероприятия"""
    builder = InlineKeyboardBuilder()
    
    for status in EVENT_STATUSES:
        builder.add(InlineKeyboardButton(text=status, callback_data=f"status_{status}"))
    
    builder.adjust(2)
    return builder.as_markup()

def get_quarters_keyboard() -> InlineKeyboardMarkup:
    """Создает клавиатуру для выбора квартала"""
    builder = InlineKeyboardBuilder()
    builder.add(
        InlineKeyboardButton(text="1 квартал (январь-март)", callback_data="quarter_1"),
        InlineKeyboardButton(text="2 квартал (апрель-июнь)", callback_data="quarter_2"),
        InlineKeyboardButton(text="3 квартал (июль-сентябрь)", callback_data="quarter_3"),
        InlineKeyboardButton(text="4 квартал (октябрь-декабрь)", callback_data="quarter_4")
    )
    builder.adjust(1)
    return builder.as_markup()

def get_time_keyboard(selected_start_time: str = None) -> InlineKeyboardMarkup:
    """Создает клавиатуру для выбора времени с 9:00 до 22:00 с шагом 1 час"""
    builder = InlineKeyboardBuilder()
    
    # Создаем кнопки для времени с 9:00 до 22:00
    for hour in range(9, 23):  # от 9 до 22 включительно
        time_str = f"{hour:02d}:00"
        
        # Если это выбранное время начала, отмечаем его
        if selected_start_time == time_str:
            builder.add(InlineKeyboardButton(text=f"✅{time_str}", callback_data=f"time_{time_str}"))
        else:
            builder.add(InlineKeyboardButton(text=time_str, callback_data=f"time_{time_str}"))
    
    builder.adjust(4)  # 4 кнопки в ряду для времени
    
    # Кнопка для ввода своего времени - отдельным рядом
    builder.row(InlineKeyboardButton(text="✏️ Ввести свое время", callback_data="custom_time"))
    
    # Если уже выбрано время начала, добавляем кнопку отмены выбора
    if selected_start_time:
        builder.row(InlineKeyboardButton(text="🗑️ Отменить выбор начала", callback_data="cancel_start_time"))
    
    return builder.as_markup()

@dp.callback_query(F.data == "cancel_start_time")
async def process_cancel_start_time(callback: types.CallbackQuery, state: FSMContext):
    """Обработка отмены выбора времени начала"""
    user_id = callback.from_user.id
    
    draft = get_report_draft(user_id)
    if not draft or not draft.get('events'):
        await callback.answer("❌ Нет активного черновика")
        return
    
    current_event_index = draft.get('current_event_index', 0)
    
    # Очищаем временные данные
    draft['events'][current_event_index]['time_data'] = {}
    save_report_draft(user_id, draft)
    
    await callback.message.edit_text(
        "✅ Выбор времени начала отменен\n\n"
        "⏰ Теперь выберите время начала мероприятия:",
        reply_markup=get_time_keyboard()
    )
    await callback.answer()

def get_standard_links_keyboard(for_plan: bool = True) -> InlineKeyboardMarkup:
    """Создает клавиатуру для выбора стандартных ссылок"""
    builder = InlineKeyboardBuilder()
    
    if for_plan:
        builder.add(
            InlineKeyboardButton(text="📢 Молодежь Сочи", callback_data="link_telegram_molod"),
            InlineKeyboardButton(text="📱 ЦРМ Сочи", callback_data="link_telegram_crm"),
            InlineKeyboardButton(text="➖ Поставить прочерк", callback_data="link_dash"),
            InlineKeyboardButton(text="✏️ Ввести свою ссылку", callback_data="link_custom")
        )
        builder.adjust(2)  # 2 кнопки в ряду для планов
    else:
        # Для отчетов - возможность ввести свою ссылку или поставить прочерк
        builder.add(
            InlineKeyboardButton(text="➖ Поставить прочерк", callback_data="link_dash"),
            InlineKeyboardButton(text="✏️ Ввести ссылку", callback_data="link_custom")
        )
        builder.adjust(1)  # 1 кнопка в ряду для отчетов
    
    return builder.as_markup()

def get_edit_report_keyboard(event_index: int, events_count: int) -> InlineKeyboardMarkup:
    """Создает клавиатуру для редактирования отчета"""
    builder = InlineKeyboardBuilder()
    
    builder.add(
        InlineKeyboardButton(text="✏️ Название", callback_data=f"edit_report_name_{event_index}"),
        InlineKeyboardButton(text="📅 Даты", callback_data=f"edit_report_dates_{event_index}"),
        InlineKeyboardButton(text="📍 Место", callback_data=f"edit_report_place_{event_index}"),
        InlineKeyboardButton(text="⏰ Время", callback_data=f"edit_report_time_{event_index}"),
        InlineKeyboardButton(text="📝 Описание", callback_data=f"edit_report_desc_{event_index}"),
        InlineKeyboardButton(text="🔗 Ссылка", callback_data=f"edit_report_link_{event_index}"),
        InlineKeyboardButton(text="👥 Участники 14-17", callback_data=f"edit_report_age14_{event_index}"),
        InlineKeyboardButton(text="👥 Участники 18-29", callback_data=f"edit_report_age18_{event_index}"),
        InlineKeyboardButton(text="👥 Участники 30-35", callback_data=f"edit_report_age30_{event_index}")
    )
    
    if events_count > 1:
        builder.add(InlineKeyboardButton(text="🗑️ Удалить мероприятие", callback_data=f"confirm_delete_report_event_{event_index}"))
    
    builder.add(
        InlineKeyboardButton(text="◀️ Назад к списку", callback_data="back_to_events_list"),
        InlineKeyboardButton(text="✅ Завершить редактирование", callback_data="finish_editing")
    )
    
    builder.adjust(2)
    return builder.as_markup()

def get_edit_plan_keyboard(item_index: int, items_count: int) -> InlineKeyboardMarkup:
    """Создает клавиатуру для редактирования плана"""
    builder = InlineKeyboardBuilder()
    
    builder.add(
        InlineKeyboardButton(text="🎯 Направление", callback_data=f"edit_plan_direction_{item_index}"),
        InlineKeyboardButton(text="✏️ Название", callback_data=f"edit_plan_name_{item_index}"),
        InlineKeyboardButton(text="🏷️ Статус", callback_data=f"edit_plan_status_{item_index}"),
        InlineKeyboardButton(text="📅 Даты", callback_data=f"edit_plan_dates_{item_index}"),
        InlineKeyboardButton(text="📍 Место", callback_data=f"edit_plan_place_{item_index}"),
        InlineKeyboardButton(text="👥 Участники", callback_data=f"edit_plan_participants_{item_index}"),
        InlineKeyboardButton(text="🔗 Ссылка", callback_data=f"edit_plan_link_{item_index}")
    )
    
    if items_count > 1:
        builder.add(InlineKeyboardButton(text="🗑️ Удалить пункт", callback_data=f"confirm_delete_plan_item_{item_index}"))
    
    builder.add(
        InlineKeyboardButton(text="◀️ Назад к списку", callback_data="back_to_plan_list"),
        InlineKeyboardButton(text="✅ Завершить редактирование", callback_data="finish_plan_editing")
    )
    
    builder.adjust(2)
    return builder.as_markup()

def get_events_list_keyboard(events: List[Dict], document_type: str = "report") -> InlineKeyboardMarkup:
    """Создает клавиатуру со списком мероприятий/пунктов для редактирования"""
    builder = InlineKeyboardBuilder()
    
    for i, event in enumerate(events):
        if document_type == "report":
            name = event.get('event_name', 'Без названия')[:30]
            builder.add(InlineKeyboardButton(text=f"{i+1}. {name}", callback_data=f"edit_event_{i}"))
        else:
            name = event.get('event_name', 'Без названия')[:30]
            builder.add(InlineKeyboardButton(text=f"{i+1}. {name}", callback_data=f"edit_plan_item_{i}"))
    
    builder.add(
        InlineKeyboardButton(text="➕ Добавить еще", callback_data="add_event" if document_type == "report" else "add_plan_item"),
        InlineKeyboardButton(text="📄 Сформировать документ", callback_data="generate_report" if document_type == "report" else "generate_plan")
    )
    builder.adjust(1)
    return builder.as_markup()

def get_final_keyboard(document_type: str = "report") -> InlineKeyboardMarkup:
    """Создает клавиатуру для завершения или добавления мероприятий"""
    builder = InlineKeyboardBuilder()
    if document_type == "report":
        builder.add(
            InlineKeyboardButton(text="➕ Добавить мероприятие", callback_data="add_event"),
            InlineKeyboardButton(text="✏️ Редактировать мероприятия", callback_data="edit_events"),
            InlineKeyboardButton(text="📄 Сформировать отчет", callback_data="generate_report")
        )
    else:
        builder.add(
            InlineKeyboardButton(text="➕ Добавить пункт плана", callback_data="add_plan_item"),
            InlineKeyboardButton(text="✏️ Редактировать пункты", callback_data="edit_plan_items"),
            InlineKeyboardButton(text="📄 Сформировать план", callback_data="generate_plan")
        )
    builder.adjust(1)
    return builder.as_markup()

def get_overwrite_keyboard(document_type: str) -> InlineKeyboardMarkup:
    """Клавиатура при обнаружении существующего документа: редактировать / перезаполнить / оставить."""
    builder = InlineKeyboardBuilder()
    builder.add(
        InlineKeyboardButton(
            text="✏️ Изменить мероприятия/пункты",
            callback_data=f"edit_existing_{document_type}",
        ),
        InlineKeyboardButton(
            text="🔄 Перезаполнить полностью",
            callback_data=f"overwrite_{document_type}",
        ),
        InlineKeyboardButton(
            text="✅ Оставить как есть",
            callback_data="cancel_overwrite",
        ),
    )
    builder.adjust(1)
    return builder.as_markup()

def get_months_keyboard(year: int = None) -> InlineKeyboardMarkup:
    """Создает клавиатуру для выбора месяца"""
    builder = InlineKeyboardBuilder()

    if year is None:
        year = datetime.now().year

    months = [
        "Январь", "Февраль", "Март", "Апрель", "Май", "Июнь",
        "Июль", "Август", "Сентябрь", "Октябрь", "Ноябрь", "Декабрь"
    ]

    for i, month in enumerate(months, 1):
        builder.add(InlineKeyboardButton(text=month, callback_data=f"month_{i}_{year}"))

    builder.adjust(3)

    current_year = datetime.now().year
    prev_year = current_year - 1
    next_year = current_year + 1

    year_buttons = []
    if year != prev_year:
        year_buttons.append(InlineKeyboardButton(text=f"◀ {prev_year}", callback_data=f"select_year_{prev_year}"))
    year_buttons.append(InlineKeyboardButton(text=f"📅 {year}", callback_data="ignore"))
    if year != next_year:
        year_buttons.append(InlineKeyboardButton(text=f"{next_year} ▶", callback_data=f"select_year_{next_year}"))
    builder.row(*year_buttons)

    builder.row(InlineKeyboardButton(text="❌ Отмена", callback_data="cancel_to_start"))
    return builder.as_markup()

def get_place_keyboard() -> InlineKeyboardMarkup:
    """Создает клавиатуру для быстрого выбора места проведения"""
    builder = InlineKeyboardBuilder()
    
    builder.add(
        InlineKeyboardButton(text="🏠 Дом молодежи", callback_data="place_dom_molodezhi"),
        InlineKeyboardButton(text="✏️ Ввести свое место", callback_data="place_custom")
    )
    
    builder.adjust(1)
    return builder.as_markup()

def get_month_calendar(year: int, month: int, selected_dates: List[str]) -> InlineKeyboardMarkup:
    """Создает календарь для выбора дат с возможностью выбора дней недели"""
    builder = InlineKeyboardBuilder()
    
    # Названия дней недели (сделаем их кликабельными)
    week_days = ["Пн", "Вт", "Ср", "Чт", "Пт", "Сб", "Вс"]
    
    # Добавляем заголовок с днями недели (теперь кликабельные)
    for day_index, day in enumerate(week_days):
        builder.add(InlineKeyboardButton(text=day, callback_data=f"weekday_{day_index}"))
    builder.adjust(7)
    
    # Получаем календарь месяца
    month_cal = monthcalendar(year, month)
    
    # Создаем кнопки для каждого дня
    for week in month_cal:
        for day in week:
            if day == 0:
                builder.add(InlineKeyboardButton(text=" ", callback_data="ignore"))
            else:
                date_str = f"{day:02d}.{month:02d}.{year}"
                if date_str in selected_dates:
                    builder.add(InlineKeyboardButton(text=f"✅{day}", callback_data=f"date_{date_str}"))
                else:
                    builder.add(InlineKeyboardButton(text=str(day), callback_data=f"date_{date_str}"))
        builder.adjust(7)
    
    # Кнопки управления
    builder.row(
        InlineKeyboardButton(text="🗑️ Очистить все", callback_data="clear_all"),
        InlineKeyboardButton(text="✅ Завершить выбор", callback_data="finish_dates")
    )
    
    return builder.as_markup()

def get_weekday_dates(year: int, month: int, weekday_index: int) -> List[str]:
    """Возвращает все даты указанного дня недели в месяце"""
    dates = []
    month_cal = monthcalendar(year, month)
    
    for week in month_cal:
        day = week[weekday_index]  # weekday_index: 0=Пн, 1=Вт, ..., 6=Вс
        if day != 0:  # 0 означает день другого месяца
            date_str = f"{day:02d}.{month:02d}.{year}"
            dates.append(date_str)
    
    return dates

def parse_month_year(month_year: str) -> Tuple[int, int]:
    """Парсит строку с месяцем и годом"""
    months = {
        'январь': 1, 'февраль': 2, 'март': 3, 'апрель': 4,
        'май': 5, 'июнь': 6, 'июль': 7, 'август': 8,
        'сентябрь': 9, 'октябрь': 10, 'ноябрь': 11, 'декабрь': 12
    }
    
    for month_name, month_num in months.items():
        if month_name in month_year.lower():
            year_match = re.search(r'\d{4}', month_year)
            if year_match:
                year = int(year_match.group())
                return month_num, year
    
    current_year = datetime.now().year
    for month_name, month_num in months.items():
        if month_name in month_year.lower():
            return month_num, current_year
    
    return datetime.now().month, datetime.now().year

def format_date_for_display(date_str: str) -> str:
    """Форматирует дату для отображения"""
    day, month, year = date_str.split('.')
    months = {
        '01': 'января', '02': 'февраля', '03': 'марта', '04': 'апреля',
        '05': 'мая', '06': 'июня', '07': 'июля', '08': 'августа',
        '09': 'сентября', '10': 'октября', '11': 'ноября', '12': 'декабря'
    }
    return f"{int(day)} {months[month]} {year} года"

def format_dates_for_cell(dates: List[str], place: str, time: str) -> str:
    """Форматирует даты, место и время для ячейки"""
    formatted_dates = "\n".join([format_date_for_display(date) for date in sorted(dates)])
    return f"{formatted_dates}\n{place}\n{time}"

def format_dates_for_plan(dates: List[str], place: str) -> str:
    """Форматирует даты и место для плана"""
    formatted_dates = "\n".join([format_date_for_display(date) for date in sorted(dates)])
    return f"{formatted_dates}\n{place}"

def is_valid_url(url: str) -> bool:
    """Проверяет, является ли строка валидной URL или прочерком"""
    if url == "-":
        return True
    
    # Проверяем несколько распространенных форматов URL
    patterns = [
        r'^https?://[^\s/$.?#].[^\s]*$',  # http/https URL
        r'^t\.me/[a-zA-Z0-9_]+$',  # Telegram
        r'^@[a-zA-Z0-9_]+$',  # Telegram username
        r'^vk\.com/[a-zA-Z0-9_]+$',  # VK
    ]
    
    for pattern in patterns:
        if re.match(pattern, url, re.IGNORECASE):
            return True
    
    return False

def set_cell_style(cell, size=12, bold=False):
    """Устанавливает стиль для ячейки"""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(size)
            run.font.bold = bold
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def set_paragraph_style(paragraph, size=12, bold=False, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT):
    """Устанавливает стиль для параграфа"""
    paragraph.alignment = alignment
    # Очищаем существующие runs
    for run in paragraph.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        run.font.bold = bold
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def fill_report_template(template_path: str, user_data: Dict, events: List[Dict]) -> str:
    """Заполнение шаблона отчета данными пользователя"""
    doc = Document(template_path)
    
    # Заполняем шапку
    for paragraph in doc.paragraphs:
        if "ДАТАиГОД" in paragraph.text:
            paragraph.text = paragraph.text.replace("ДАТАиГОД", user_data['month_year'])
            set_paragraph_style(paragraph, 14, True, WD_PARAGRAPH_ALIGNMENT.CENTER)
        
        if "НАЗВАНИЕКЛУБА" in paragraph.text:
            paragraph.text = paragraph.text.replace("НАЗВАНИЕКЛУБА", user_data['club_name'])
            set_paragraph_style(paragraph, 14, True, WD_PARAGRAPH_ALIGNMENT.CENTER)
    
    # Заполняем таблицу мероприятиями
    for table in doc.tables:
        total_rows = len(table.rows)
        events_count = len(events)
        
        if events_count < 15 and total_rows > events_count + 2:
            for i in range(total_rows - 1, events_count + 1, -1):
                tbl = table._tbl
                tr = table.rows[i]._tr
                tbl.remove(tr)
        
        for i, event in enumerate(events):
            row_index = i + 2
            if row_index < len(table.rows):
                cells = table.rows[row_index].cells
                
                if len(cells) < 10:
                    continue
                
                cells[0].text = str(i + 1)
                set_cell_style(cells[0], 12)
                
                cells[1].text = event['event_name']
                set_cell_style(cells[1], 12)
                
                date_time_place = format_dates_for_cell(
                    event['dates'], 
                    event['place'], 
                    event['time']
                )
                cells[2].text = date_time_place
                set_cell_style(cells[2], 12)
                
                cells[3].text = event['description']
                set_cell_style(cells[3], 12)
                
                # Ответственный берется из данных клуба
                cells[4].text = user_data['responsible']
                set_cell_style(cells[4], 12)
                
                # Ссылка - если "-", то ставим прочерк, иначе все ссылки через перенос строки
                link_text = "-" if event['link'] == "-" else event['link']
                cells[5].text = link_text
                set_cell_style(cells[5], 12)
                
                total = event['age_14_17'] + event['age_18_29'] + event['age_30_35']
                cells[6].text = str(total)
                set_cell_style(cells[6], 12)
                
                cells[7].text = str(event['age_14_17'])
                set_cell_style(cells[7], 12)
                
                cells[8].text = str(event['age_18_29'])
                set_cell_style(cells[8], 12)
                
                cells[9].text = str(event['age_30_35'])
                set_cell_style(cells[9], 12)
    
    # Заполняем подпись
    for paragraph in doc.paragraphs:
        if "ДОЛЖНОСТЬ" in paragraph.text:
            paragraph.text = paragraph.text.replace("ДОЛЖНОСТЬ", user_data['position'])
            set_paragraph_style(paragraph, 14, False, WD_PARAGRAPH_ALIGNMENT.LEFT)
        if "ФИО" in paragraph.text:
            paragraph.text = paragraph.text.replace("ФИО", user_data['responsible'])
            set_paragraph_style(paragraph, 14, False, WD_PARAGRAPH_ALIGNMENT.LEFT)
    
    # Формируем название файла
    safe_club_name = "".join(c for c in user_data['club_name'] if c.isalnum() or c in (' ', '-', '_')).rstrip()
    safe_month_year = "".join(c for c in user_data['month_year'] if c.isalnum() or c in (' ', '-', '_')).rstrip()
    filename = f"Отчет за {safe_month_year} ({safe_club_name}).docx"
    output_path = os.path.join(TEMP_DIR, filename)
    
    doc.save(output_path)
    
    return output_path

def fill_plan_template(template_path: str, user_data: Dict, plan_items: List[Dict]) -> str:
    """Заполнение шаблона плана данными пользователя"""
    doc = Document(template_path)
    
    # Заполняем шапку
    for paragraph in doc.paragraphs:
        if "ДАТАиГОД" in paragraph.text:
            paragraph.text = paragraph.text.replace("ДАТАиГОД", user_data['month_year'])
            set_paragraph_style(paragraph, 14, True, WD_PARAGRAPH_ALIGNMENT.CENTER)
        
        if "НАЗВАНИЕКЛУБА" in paragraph.text:
            paragraph.text = paragraph.text.replace("НАЗВАНИЕКЛУБА", user_data['club_name'])
            set_paragraph_style(paragraph, 14, True, WD_PARAGRAPH_ALIGNMENT.CENTER)
    
    # Заполняем таблицу пунктами плана
    for table in doc.tables:
        total_rows = len(table.rows)
        plan_count = len(plan_items)
        
        # Удаляем лишние строки если нужно
        if plan_count < 15 and total_rows > plan_count + 1:
            for i in range(total_rows - 1, plan_count, -1):
                tbl = table._tbl
                tr = table.rows[i]._tr
                tbl.remove(tr)
        
        # Заполняем данные
        for i, item in enumerate(plan_items):
            row_index = i + 1
            if row_index < len(table.rows):
                cells = table.rows[row_index].cells
                
                if len(cells) < 7:
                    continue
                
                # № п/п
                cells[0].text = str(i + 1)
                set_cell_style(cells[0], 12)
                
                # Направление реализации
                cells[1].text = item['direction']
                set_cell_style(cells[1], 12)
                
                # Форма и наименование мероприятия
                cells[2].text = item['event_name']
                set_cell_style(cells[2], 12)
                
                # Статус мероприятия
                cells[3].text = item['status']
                set_cell_style(cells[3], 12)
                
                # Дата и место проведения
                date_place = format_dates_for_plan(item['dates'], item['place'])
                cells[4].text = date_place
                set_cell_style(cells[4], 12)
                
                # Предполагаемое количество участников
                cells[5].text = f"{item['participants']} чел."
                set_cell_style(cells[5], 12)
                
                # Информационное освещение - если "-", то ставим прочерк, иначе ссылку
                info_link_text = "-" if item['info_link'] == "-" else item['info_link']
                cells[6].text = info_link_text
                set_cell_style(cells[6], 12)
    
    # Заполняем подпись
    for paragraph in doc.paragraphs:
        if "ДОЛЖНОСТЬ" in paragraph.text:
            paragraph.text = paragraph.text.replace("ДОЛЖНОСТЬ", user_data['position'])
            set_paragraph_style(paragraph, 14, False, WD_PARAGRAPH_ALIGNMENT.LEFT)
        if "ФИО" in paragraph.text:
            paragraph.text = paragraph.text.replace("ФИО", user_data['responsible'])
            set_paragraph_style(paragraph, 14, False, WD_PARAGRAPH_ALIGNMENT.LEFT)
    
    # Формируем название файла
    safe_club_name = "".join(c for c in user_data['club_name'] if c.isalnum() or c in (' ', '-', '_')).rstrip()
    safe_month_year = "".join(c for c in user_data['month_year'] if c.isalnum() or c in (' ', '-', '_')).rstrip()
    filename = f"План на {safe_month_year} ({safe_club_name}).docx"
    output_path = os.path.join(TEMP_DIR, filename)
    
    doc.save(output_path)
    
    return output_path

# Класс для генерации квартальных отчетов
class QuarterlyReportGenerator:
    def __init__(self):
        self.months_data = {}
        self.quarter_metrics = {
            'total_events': 0,
            'total_participants': 0,
            'age_14_17': 0,
            'age_18_29': 0,
            'age_30_35': 0
        }
    
    def collect_monthly_reports(self, club_name: str, year: int, quarter: int) -> Dict[str, List]:
        """Сбор ежемесячных отчетов за квартал с группировкой по месяцам"""
        months_mapping = {
            1: ['январь', 'февраль', 'март'],
            2: ['апрель', 'май', 'июнь'], 
            3: ['июль', 'август', 'сентябрь'],
            4: ['октябрь', 'ноябрь', 'декабрь']
        }
        
        quarter_months = months_mapping.get(quarter, [])
        self.months_data = {}
        
        for month in quarter_months:
            month_year = f"{month} {year}"
            reports = get_club_reports(club_name, year)
            
            # Ищем отчет за конкретный месяц
            monthly_report = None
            for report in reports:
                report_month_year = report[0]  # month_year из базы данных
                if month_year.lower() in report_month_year.lower():
                    monthly_report = report
                    break
            
            logger.info(f"Looking for {month_year}: found {monthly_report}")
            
            if monthly_report and os.path.exists(monthly_report[1]):
                events = self.parse_monthly_report(monthly_report[1])
                self.months_data[month] = events
                logger.info(f"Parsed {len(events)} events for {month}")
                
                # Суммируем метрики
                for event in events:
                    self.quarter_metrics['total_events'] += 1
                    self.quarter_metrics['age_14_17'] += event['age_14_17']
                    self.quarter_metrics['age_18_29'] += event['age_18_29'] 
                    self.quarter_metrics['age_30_35'] += event['age_30_35']
        
        self.quarter_metrics['total_participants'] = (
            self.quarter_metrics['age_14_17'] + 
            self.quarter_metrics['age_18_29'] + 
            self.quarter_metrics['age_30_35']
        )
        
        return self.months_data
    
    def parse_monthly_report(self, file_path: str) -> List[Dict]:
        """Парсит мероприятия из месячного отчета"""
        events = []
        try:
            doc = Document(file_path)
            logger.info(f"Parsing document: {file_path}")
            
            for table_index, table in enumerate(doc.tables):
                logger.info(f"Table {table_index} has {len(table.rows)} rows")
                
                # Ищем строки с мероприятиями (пропускаем заголовки)
                for i, row in enumerate(table.rows):
                    if i < 2:  # Пропускаем первые 2 строки - заголовки
                        continue
                    
                    cells = row.cells
                    if len(cells) >= 10:
                        event_data = self._parse_event_row(cells, i)
                        if event_data:
                            events.append(event_data)
                            logger.info(f"Found event: {event_data['event_name'][:50]}...")
            
            logger.info(f"Total events parsed: {len(events)}")
            return events
        except Exception as e:
            logger.error(f"Error parsing monthly report {file_path}: {e}")
            return []
    
    def _parse_event_row(self, cells, row_index: int) -> Optional[Dict]:
        """Парсит строку с мероприятием - ТОЧНОЕ СОВПАДЕНИЕ С ЯЧЕЙКАМИ"""
        try:
            # Проверяем, что это валидная строка с мероприятием (есть номер)
            cell_0_text = cells[0].text.strip()
            if not cell_0_text or not cell_0_text.replace('.', '').isdigit():
                return None
            
            # Извлекаем данные из ячеек БЕЗ ИЗМЕНЕНИЙ
            event_name = cells[1].text.strip()
            
            # Ячейка с датами, временем и местом - берем как есть
            date_time_place = cells[2].text.strip()
            
            description = cells[3].text.strip()
            link = cells[5].text.strip()
            
            # Парсим числовые значения
            age_14_17 = self._safe_int(cells[7].text)
            age_18_29 = self._safe_int(cells[8].text)
            age_30_35 = self._safe_int(cells[9].text)
            
            # Если все поля пустые, пропускаем
            if not event_name and not description and age_14_17 == 0 and age_18_29 == 0 and age_30_35 == 0:
                return None
            
            # Извлекаем только даты из ячейки (первые строки до места и времени)
            dates = self._extract_dates_only(date_time_place)
            
            event = {
                'event_name': event_name,
                'dates': dates,
                'date_time_place_full': date_time_place,  # Сохраняем полный текст
                'description': description,
                'link': link,
                'age_14_17': age_14_17,
                'age_18_29': age_18_29,
                'age_30_35': age_30_35
            }
            
            logger.info(f"Row {row_index}: {event_name}")
            logger.info(f"Full date/time/place: {date_time_place}")
            logger.info(f"Extracted dates: {dates}")
            
            return event
            
        except Exception as e:
            logger.error(f"Error parsing event row {row_index}: {e}")
            logger.error(f"Cell contents: {[cell.text.strip() for cell in cells]}")
            return None

    def _extract_dates_only(self, date_time_place_text: str) -> List[str]:
        """Извлекает только даты из текста ячейки с датами, временем и местом"""
        dates = []
        try:
            lines = date_time_place_text.split('\n')
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                    
                # Останавливаемся, когда находим место или время
                if any(keyword in line.lower() for keyword in ['ул.', 'улица', 'дом', 'д.', 'время', ':', '00']):
                    break
                    
                # Ищем даты в формате DD.MM.YYYY
                date_matches = re.findall(r'\b\d{1,2}\.\d{1,2}\.\d{4}\b', line)
                dates.extend(date_matches)
                
                # Если не нашли в формате DD.MM.YYYY, ищем в других форматах
                if not date_matches:
                    # Формат DD.MM.YY
                    date_matches_yy = re.findall(r'\b\d{1,2}\.\d{1,2}\.\d{2}\b', line)
                    for date_str in date_matches_yy:
                        try:
                            day, month, year = date_str.split('.')
                            full_year = f"20{year}" if int(year) < 100 else year
                            standardized = f"{int(day):02d}.{int(month):02d}.{full_year}"
                            dates.append(standardized)
                        except:
                            dates.append(date_str)
            
            # Если даты не найдены, но есть текст - возвращаем первую строку как дату
            if not dates and lines:
                first_line = lines[0].strip()
                if first_line and not any(keyword in first_line.lower() for keyword in ['ул.', 'улица', 'дом', 'д.', 'время']):
                    dates.append(first_line)
            
            return dates if dates else ["Дата не указана"]
            
        except Exception as e:
            logger.error(f"Error extracting dates: {e}")
            return ["Ошибка извлечения дат"]
    
    def _safe_int(self, text: str) -> int:
        """Безопасное преобразование в int"""
        try:
            # Удаляем все нецифровые символы, кроме минуса
            cleaned = re.sub(r'[^\d-]', '', text.strip())
            return int(cleaned) if cleaned else 0
        except ValueError:
            return 0
    
    def generate_quarterly_report(self, template_path: str, user_data: Dict) -> str:
        """Генерирует квартальный отчет по специальному шаблону"""
        try:
            doc = Document(template_path)
            logger.info("Starting quarterly report generation")
            
            # Заполняем шапку
            for paragraph in doc.paragraphs:
                if "КВАРТАЛиГОД" in paragraph.text:
                    quarter_text = self._get_quarter_text(user_data['quarter'], user_data['year'])
                    paragraph.text = paragraph.text.replace("КВАРТАЛиГОД", quarter_text)
                    set_paragraph_style(paragraph, 14, True, WD_PARAGRAPH_ALIGNMENT.CENTER)
                    logger.info(f"Replaced КВАРТАЛиГОД with: {quarter_text}")
                
                if "НАЗВАНИЕКЛУБА" in paragraph.text:
                    paragraph.text = paragraph.text.replace("НАЗВАНИЕКЛУБА", user_data['club_name'])
                    set_paragraph_style(paragraph, 14, True, WD_PARAGRAPH_ALIGNMENT.CENTER)
                    logger.info(f"Replaced НАЗВАНИЕКЛУБА with: {user_data['club_name']}")
            
            # Заполняем таблицу с группировкой по месяцам
            self._fill_quarterly_table(doc, user_data)
            
            # Заполняем подпись
            for paragraph in doc.paragraphs:
                if "ДОЛЖНОСТЬ" in paragraph.text:
                    paragraph.text = paragraph.text.replace("ДОЛЖНОСТЬ", user_data['position'])
                    set_paragraph_style(paragraph, 14, False, WD_PARAGRAPH_ALIGNMENT.LEFT)
                if "ФИО" in paragraph.text:
                    paragraph.text = paragraph.text.replace("ФИО", user_data['responsible'])
                    set_paragraph_style(paragraph, 14, False, WD_PARAGRAPH_ALIGNMENT.LEFT)
            
            # Сохраняем файл
            safe_club_name = "".join(c for c in user_data['club_name'] if c.isalnum() or c in (' ', '-', '_')).rstrip()
            filename = f"Квартальный отчет {user_data['quarter']} квартал {user_data['year']} ({safe_club_name}).docx"
            output_path = os.path.join(TEMP_DIR, filename)
            
            doc.save(output_path)
            logger.info(f"Quarterly report saved to: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error generating quarterly report: {e}")
            raise
    
    def _fill_quarterly_table(self, doc: Document, user_data: Dict):
        """Заполняет таблицу квартального отчета с группировкой по месяцам и удаляет пустые строки"""
        try:
            for table_index, table in enumerate(doc.tables):
                logger.info(f"Processing table {table_index} with {len(table.rows)} rows")
                
                # Собираем индексы строк для удаления (пустые строки между месяцами)
                rows_to_remove = []
                events_added = 0
                
                # Проходим по всем месяцам квартала
                for month_num, (month_name, events) in enumerate(self.months_data.items(), 1):
                    if not events:
                        logger.info(f"No events for {month_name}, skipping")
                        continue
                    
                    logger.info(f"Processing {month_name} with {len(events)} events")
                    
                    # Находим строку с заголовком месяца (МЕСЯЦ1, МЕСЯЦ2, МЕСЯЦ3)
                    month_header_row = self._find_month_header_row(table, month_num)
                    if month_header_row is None:
                        logger.warning(f"Month header for МЕСЯЦ{month_num} not found")
                        continue
                    
                    # Обновляем заголовок месяца
                    month_cell = table.rows[month_header_row].cells[0]
                    month_cell.text = month_name.upper()
                    set_cell_style(month_cell, 12, True)
                    logger.info(f"Updated month header for {month_name} at row {month_header_row}")
                    
                    # Заполняем мероприятия для этого месяца
                    start_row = month_header_row + 1
                    events_filled = 0
                    
                    for i, event in enumerate(events):
                        if start_row + i >= len(table.rows):
                            logger.warning(f"Not enough rows in table for all events. Needed: {start_row + i}, available: {len(table.rows)}")
                            break
                        
                        # Заполняем строку мероприятием
                        self._fill_event_row(table.rows[start_row + i], event, events_filled + 1, user_data)
                        events_filled += 1
                        events_added += 1
                    
                    # Помечаем оставшиеся пустые строки после мероприятий для удаления
                    next_month_start = self._find_next_month_header_row(table, month_num + 1)
                    if next_month_start is None:
                        next_month_start = len(table.rows)
                    
                    # Удаляем пустые строки между последним мероприятием и следующим месяцем
                    empty_rows_after = next_month_start - (start_row + events_filled)
                    if empty_rows_after > 0:
                        for i in range(empty_rows_after):
                            row_index_to_remove = start_row + events_filled + i
                            if row_index_to_remove < len(table.rows):
                                rows_to_remove.append(row_index_to_remove)
                    
                    logger.info(f"Added {events_filled} events for {month_name}")
                    logger.info(f"Marked {empty_rows_after} empty rows after {month_name} for removal")
                
                # Удаляем помеченные пустые строки (в обратном порядке чтобы индексы не сдвигались)
                if rows_to_remove:
                    rows_to_remove.sort(reverse=True)  # Удаляем с конца чтобы индексы не менялись
                    logger.info(f"Removing {len(rows_to_remove)} empty rows: {rows_to_remove}")
                    
                    for row_index in rows_to_remove:
                        if row_index < len(table.rows):
                            try:
                                tbl = table._tbl
                                tr = table.rows[row_index]._tr
                                tbl.remove(tr)
                                logger.info(f"Removed empty row {row_index}")
                            except Exception as e:
                                logger.error(f"Error removing row {row_index}: {e}")
                
                logger.info(f"Total events added to table: {events_added}")
                logger.info(f"Final table has {len(table.rows)} rows")
                
        except Exception as e:
            logger.error(f"Error filling quarterly table: {e}")
            raise

    def _find_next_month_header_row(self, table, next_month_num: int) -> Optional[int]:
        """Находит строку с заголовком следующего месяца"""
        if next_month_num > 3:  # Максимум 3 месяца в квартале
            return None
        
        for i, row in enumerate(table.rows):
            if len(row.cells) > 0:
                cell_text = row.cells[0].text.strip()
                if f"МЕСЯЦ{next_month_num}" in cell_text:
                    logger.info(f"Found next month header for МЕСЯЦ{next_month_num} at row {i}")
                    return i
        return None
    
    def _find_month_header_row(self, table, month_num: int) -> Optional[int]:
        """Находит строку с заголовком месяца по шаблону"""
        for i, row in enumerate(table.rows):
            if len(row.cells) > 0:
                cell_text = row.cells[0].text.strip()
                if f"МЕСЯЦ{month_num}" in cell_text:
                    logger.info(f"Found month header for МЕСЯЦ{month_num} at row {i}")
                    return i
        logger.warning(f"Month header for МЕСЯЦ{month_num} not found")
        return None
    
    def _fill_event_row(self, row, event: Dict, event_num: int, user_data: Dict):
        """Заполняет строку с мероприятием"""
        try:
            cells = row.cells
            if len(cells) < 10:
                logger.warning(f"Not enough cells in row: {len(cells)}")
                return
            
            # Проверяем, не является ли строка уже заполненной (заголовком месяца)
            first_cell_text = cells[0].text.strip()
            if any(month_header in first_cell_text for month_header in ['МЕСЯЦ1', 'МЕСЯЦ2', 'МЕСЯЦ3']):
                logger.info(f"Skipping month header row: {first_cell_text}")
                return
            
            # № п/п
            cells[0].text = str(event_num)
            set_cell_style(cells[0], 12)
            
            # Наименование мероприятия
            cells[1].text = event['event_name'] if event['event_name'] else " "
            set_cell_style(cells[1], 12)
            
            # Дата, время и место - берем ИЗНАЧАЛЬНЫЙ ТЕКСТ из месячного отчета
            cells[2].text = event['date_time_place_full'] if event['date_time_place_full'] else " "
            set_cell_style(cells[2], 12)
            
            # Краткое содержание
            cells[3].text = event['description'] if event['description'] else " "
            set_cell_style(cells[3], 12)
            
            # Ответственный
            cells[4].text = user_data['responsible']
            set_cell_style(cells[4], 12)
            
            # Ссылка
            cells[5].text = event['link'] if event['link'] else " "
            set_cell_style(cells[5], 12)
            
            # Охват молодежи
            total = event['age_14_17'] + event['age_18_29'] + event['age_30_35']
            cells[6].text = str(total) if total > 0 else " "
            set_cell_style(cells[6], 12)
            
            cells[7].text = str(event['age_14_17']) if event['age_14_17'] > 0 else " "
            set_cell_style(cells[7], 12)
            
            cells[8].text = str(event['age_18_29']) if event['age_18_29'] > 0 else " "
            set_cell_style(cells[8], 12)
            
            cells[9].text = str(event['age_30_35']) if event['age_30_35'] > 0 else " "
            set_cell_style(cells[9], 12)
            
            logger.info(f"Filled event row {event_num}: {event['event_name'][:30]}...")
            
        except Exception as e:
            logger.error(f"Error filling event row: {e}")
    
    def _get_quarter_text(self, quarter: int, year: int) -> str:
        """Возвращает текстовое представление квартала"""
        quarters = {
            1: "1 квартал",
            2: "2 квартал", 
            3: "3 квартал",
            4: "4 квартал"
        }
        return f"{quarters.get(quarter, '')} {year} года"

# Обработчики callback-ов для выбора месяца
@dp.callback_query(F.data.startswith("select_year_"))
async def process_year_switch(callback: types.CallbackQuery, state: FSMContext):
    """Переключение года в клавиатуре месяцев"""
    year = int(callback.data.replace("select_year_", ""))
    current_state = await state.get_state()

    if current_state == ReportStates.waiting_for_month.state:
        await callback.message.edit_text(
            f"1️⃣ Выберите месяц ({year} год):",
            reply_markup=get_months_keyboard(year)
        )
    elif current_state == PlanStates.waiting_for_plan_month.state:
        await callback.message.edit_text(
            f"1️⃣ Выберите месяц ({year} год):",
            reply_markup=get_months_keyboard(year)
        )
    await callback.answer()

@dp.callback_query(F.data == "cancel_to_start")
async def cancel_to_start(callback: types.CallbackQuery, state: FSMContext):
    """Отмена и возврат к главному меню"""
    user_id = callback.from_user.id
    delete_report_draft(user_id)
    delete_plan_draft(user_id)
    await state.clear()
    await callback.message.delete()
    await cmd_start(callback.message, state)
    await callback.answer()

@dp.callback_query(F.data == "back_to_month_selection")
async def back_to_month_selection(callback: types.CallbackQuery, state: FSMContext):
    """Возврат к выбору месяца"""
    current_state = await state.get_state()
    current_year = datetime.now().year

    if current_state in (ReportStates.waiting_for_club_selection.state, ReportStates.waiting_for_month.state):
        await callback.message.edit_text(
            "1️⃣ Выберите месяц:",
            reply_markup=get_months_keyboard(current_year)
        )
        await state.set_state(ReportStates.waiting_for_month)
    elif current_state in (PlanStates.waiting_for_plan_club_selection.state, PlanStates.waiting_for_plan_month.state):
        await callback.message.edit_text(
            "1️⃣ Выберите месяц:",
            reply_markup=get_months_keyboard(current_year)
        )
        await state.set_state(PlanStates.waiting_for_plan_month)
    await callback.answer()

@dp.callback_query(F.data.startswith("month_"))
async def process_month_selection(callback: types.CallbackQuery, state: FSMContext):
    """Обработка выбора месяца"""
    user_id = callback.from_user.id
    parts = callback.data.replace("month_", "").split("_")

    months_ru = [
        "январь", "февраль", "март", "апрель", "май", "июнь",
        "июль", "август", "сентябрь", "октябрь", "ноябрь", "декабрь"
    ]

    if len(parts) == 2:
        month_num = int(parts[0])
        year = int(parts[1])
    else:
        month_num = int(parts[0])
        year = datetime.now().year

    selected_month = months_ru[month_num - 1]
    month_year = f"{selected_month} {year}"

    current_state = await state.get_state()

    if current_state == ReportStates.waiting_for_month.state:
        draft = get_report_draft(user_id) or {}
        draft['month_year'] = month_year
        save_report_draft(user_id, draft)

        await callback.message.edit_text(
            f"✅ Период: {month_year}\n\n"
            "2️⃣ Выберите молодежный клуб:",
            reply_markup=get_clubs_keyboard()
        )
        await state.set_state(ReportStates.waiting_for_club_selection)

    elif current_state == PlanStates.waiting_for_plan_month.state:
        draft = get_plan_draft(user_id) or {}
        draft['month_year'] = month_year
        save_plan_draft(user_id, draft)

        await callback.message.edit_text(
            f"✅ Период: {month_year}\n\n"
            "2️⃣ Выберите молодежный клуб:",
            reply_markup=get_clubs_keyboard()
        )
        await state.set_state(PlanStates.waiting_for_plan_club_selection)

    await callback.answer()

# Обработчики callback-ов для отчета
@dp.callback_query(F.data.startswith("club_"))
async def process_club_selection(callback: types.CallbackQuery, state: FSMContext):
    """Обработка выбора клуба"""
    user_id = callback.from_user.id
    club_name = callback.data.replace("club_", "")
    
    current_state = await state.get_state()
    
    if club_name in CLUBS_DATA:
        if current_state == ReportStates.waiting_for_club_selection.state:
            # Сохраняем данные в черновик
            draft = get_report_draft(user_id) or {}
            draft.update({
                'club_name': club_name,
                'position': CLUBS_DATA[club_name]['position'],
                'responsible': CLUBS_DATA[club_name]['responsible']
            })
            save_report_draft(user_id, draft)
            
            # Проверяем существование отчета
            existing_report = get_existing_report(user_id, club_name, draft['month_year'])
            if existing_report:
                file_path, events_count, total_participants, created_at = existing_report
                
                real_path = resolve_file_path(file_path)
                date_str = created_at[:10] if created_at else "—"
                if real_path and os.path.exists(real_path):
                    doc_file = FSInputFile(real_path)
                    await callback.message.answer_document(
                        doc_file,
                        caption=(
                            f"📊 Существующий отчёт\n"
                            f"🏢 {club_name} | 📅 {draft['month_year']}\n"
                            f"📋 Мероприятий: {events_count} | 👥 Участников: {total_participants}\n"
                            f"🕐 {date_str}"
                        )
                    )
                else:
                    await callback.message.answer(
                        f"⚠️ Файл отчёта не найден на диске, но запись в базе есть.\n"
                        f"Можно перезаполнить или отредактировать."
                    )

                await callback.message.answer(
                    f"📊 <b>Отчёт за {draft['month_year']} — клуб «{club_name}»</b>\n\n"
                    f"📋 Мероприятий: {events_count}\n"
                    f"👥 Участников: {total_participants}\n"
                    f"🕐 Создан: {date_str}\n\n"
                    "Что хотите сделать?",
                    reply_markup=get_overwrite_keyboard("report"),
                    parse_mode="HTML",
                )
                return
            
            await callback.message.edit_text(
                f"✅ Выбран клуб: {club_name}\n"
                f"📝 Должность: {CLUBS_DATA[club_name]['position']}\n"
                f"👤 Ответственный: {CLUBS_DATA[club_name]['responsible']}\n\n"
                f"Теперь приступим к заполнению данных о мероприятиях.\n\n"
                f"5️⃣ Введите наименование мероприятия:"
            )
            await state.set_state(ReportStates.waiting_for_event_name)
            
        elif current_state == PlanStates.waiting_for_plan_club_selection.state:
            # Сохраняем данные в черновик плана
            draft = get_plan_draft(user_id) or {}
            draft.update({
                'club_name': club_name,
                'position': CLUBS_DATA[club_name]['position'],
                'responsible': CLUBS_DATA[club_name]['responsible']
            })
            save_plan_draft(user_id, draft)
            
            # Проверяем существование плана
            existing_plan = get_existing_plan(user_id, club_name, draft['month_year'])
            if existing_plan:
                file_path, plan_items_count, total_participants, created_at = existing_plan
                
                real_path = resolve_file_path(file_path)
                date_str = created_at[:10] if created_at else "—"
                if real_path and os.path.exists(real_path):
                    doc_file = FSInputFile(real_path)
                    await callback.message.answer_document(
                        doc_file,
                        caption=(
                            f"📋 Существующий план\n"
                            f"🏢 {club_name} | 📅 {draft['month_year']}\n"
                            f"📋 Пунктов: {plan_items_count} | 👥 Участников: {total_participants}\n"
                            f"🕐 {date_str}"
                        )
                    )
                else:
                    await callback.message.answer(
                        f"⚠️ Файл плана не найден на диске, но запись в базе есть.\n"
                        f"Можно перезаполнить или отредактировать."
                    )

                await callback.message.answer(
                    f"📋 <b>План за {draft['month_year']} — клуб «{club_name}»</b>\n\n"
                    f"📋 Пунктов: {plan_items_count}\n"
                    f"👥 Участников: {total_participants}\n"
                    f"🕐 Создан: {date_str}\n\n"
                    "Что хотите сделать?",
                    reply_markup=get_overwrite_keyboard("plan"),
                    parse_mode="HTML",
                )
                return
            
            # Показываем список направлений с описанием
            directions_text = get_directions_full_text()
            
            await callback.message.edit_text(
                f"✅ Выбран клуб: {club_name}\n"
                f"📝 Должность: {CLUBS_DATA[club_name]['position']}\n"
                f"👤 Ответственный: {CLUBS_DATA[club_name]['responsible']}\n\n"
                f"Теперь приступим к заполнению плана.\n\n"
                f"{directions_text}\n\n"
                f"Выберите направление реализации государственной молодежной политики:",
                reply_markup=get_directions_keyboard()
            )
            await state.set_state(PlanStates.waiting_for_direction_selection)
            
        elif current_state == QuarterlyStates.waiting_for_quarter_club_selection.state:
            # Обработка для квартальных отчетов
            data = await state.get_data()
            quarter = data['quarter']
            year = data['year']
            
            # Собираем данные клуба
            user_quarter_data = {
                'club_name': club_name,
                'position': CLUBS_DATA[club_name]['position'],
                'responsible': CLUBS_DATA[club_name]['responsible'],
                'quarter': quarter,
                'year': year
            }
            
            await callback.message.edit_text("📊 Собираю данные за квартал...")
            
            try:
                # Генерируем квартальный отчет
                generator = QuarterlyReportGenerator()
                months_data = generator.collect_monthly_reports(club_name, year, quarter)
                
                if not months_data:
                    await callback.message.answer(
                        f"❌ Не найдено месячных отчетов за {quarter} квартал {year} года для клуба '{club_name}'\n\n"
                        f"Убедитесь, что все месячные отчеты за квартал заполнены и сохранены в системе."
                    )
                    await state.clear()
                    return
                
                # Показываем статистику по найденным данным
                stats_text = f"📊 Найдены отчеты за {quarter} квартал {year}:\n\n"
                for month_name, events in months_data.items():
                    month_participants = sum(e['age_14_17'] + e['age_18_29'] + e['age_30_35'] for e in events)
                    stats_text += f"• {month_name}: {len(events)} мероприятий, {month_participants} участников\n"
                
                stats_text += f"\n📈 Итого за квартал: {generator.quarter_metrics['total_events']} мероприятий, {generator.quarter_metrics['total_participants']} участников"
                
                await callback.message.answer(stats_text)
                await callback.message.answer("🔄 Формирую квартальный отчет...")
                
                # Генерируем отчет
                template_path = "Шаблон квартального отчета.docx"
                output_path = generator.generate_quarterly_report(template_path, user_quarter_data)
                
                # Сохраняем в базу данных
                included_months = ", ".join(months_data.keys())
                save_quarterly_report_to_db(
                    user_id, 
                    club_name, 
                    f"{quarter} квартал {year}",
                    output_path,
                    generator.quarter_metrics['total_events'],
                    generator.quarter_metrics['total_participants'],
                    included_months
                )
                
                # Отправляем файл
                doc_file = FSInputFile(output_path)
                await callback.message.answer_document(
                    doc_file,
                    caption=(
                        f"✅ Квартальный отчет сформирован!\n"
                        f"🏢 Клуб: {club_name}\n"
                        f"📅 Период: {quarter} квартал {year}\n"
                        f"📊 Мероприятий: {generator.quarter_metrics['total_events']}\n"
                        f"👥 Участников: {generator.quarter_metrics['total_participants']}\n"
                        f"📝 Ответственный: {user_quarter_data['position']} {user_quarter_data['responsible']}"
                    )
                )
                
                # Отправляем документ администратору
                await send_document_to_admin(
                    user_id,
                    output_path,
                    f"Квартальный отчет ({quarter} квартал {year})",
                    club_name,
                    f"{quarter} квартал {year}",
                    generator.quarter_metrics['total_events'],
                    generator.quarter_metrics['total_participants']
                )
                
                await state.clear()
                await callback.answer("Квартальный отчет успешно сформирован!")
                
            except Exception as e:
                logger.error(f"Error generating quarterly report: {e}")
                await callback.message.answer(f"❌ Произошла ошибка: {str(e)}\nПопробуйте снова: /new_quarterly")
                await state.clear()
                await callback.answer("Ошибка при формировании отчета")
                
    else:
        await callback.answer("❌ Ошибка выбора клуба")
    
    await callback.answer()

@dp.callback_query(F.data.startswith("overwrite_"))
async def process_overwrite(callback: types.CallbackQuery, state: FSMContext):
    """Обработка подтверждения перезаписи"""
    user_id = callback.from_user.id
    document_type = callback.data.replace("overwrite_", "")
    
    if document_type == "report":
        draft = get_report_draft(user_id)
        if draft and 'club_name' in draft and 'month_year' in draft:
            # Удаляем существующий отчет
            delete_report_from_db(user_id, draft['club_name'], draft['month_year'])
            
            await callback.message.edit_text(
                f"✅ Старый отчет удален!\n\n"
                f"Теперь приступим к заполнению нового отчета.\n\n"
                f"5️⃣ Введите наименование мероприятия:"
            )
            await state.set_state(ReportStates.waiting_for_event_name)
        else:
            await callback.answer("❌ Ошибка: данные черновика неполные")
    
    elif document_type == "plan":
        draft = get_plan_draft(user_id)
        if draft and 'club_name' in draft and 'month_year' in draft:
            # Удаляем существующий план
            delete_plan_from_db(user_id, draft['club_name'], draft['month_year'])
            
            # Показываем список направлений с описанием
            directions_text = get_directions_full_text()
            
            await callback.message.edit_text(
                f"✅ Старый план удален!\n\n"
                f"Теперь приступим к заполнению нового плана.\n\n"
                f"{directions_text}\n\n"
                f"Выберите направление реализации государственной молодежной политики:",
                reply_markup=get_directions_keyboard()
            )
            await state.set_state(PlanStates.waiting_for_direction_selection)
        else:
            await callback.answer("❌ Ошибка: данные черновика неполные")
    
    await callback.answer()

@dp.callback_query(F.data.startswith("edit_existing_"))
async def process_edit_existing(callback: types.CallbackQuery, state: FSMContext):
    """
    Пользователь хочет изменить существующий документ без полного перезаполнения.
    Загружает сохранённые данные в черновик и продолжает с режима редактирования событий/пунктов.
    Если данные не сохранены (старый документ) — предлагает перезаполнить полностью.
    """
    user_id = callback.from_user.id
    document_type = callback.data.replace("edit_existing_", "")

    if document_type == "report":
        draft = get_report_draft(user_id)
        if not draft or 'club_name' not in draft or 'month_year' not in draft:
            await callback.answer("❌ Черновик не найден", show_alert=True)
            return

        events = get_existing_report_events(user_id, draft['club_name'], draft['month_year'])
        if not events:
            # Данных нет — предупреждаем и предлагаем перезаполнить
            await callback.message.edit_text(
                "⚠️ Данные этого отчёта недоступны для редактирования "
                "(документ создан до обновления бота).\n\n"
                "Нажмите «🔄 Перезаполнить полностью» чтобы создать новый отчёт.",
                reply_markup=get_overwrite_keyboard("report"),
            )
            await callback.answer()
            return

        # Восстанавливаем черновик с существующими мероприятиями
        # Убеждаемся что все обязательные поля присутствуют
        draft.setdefault('selected_dates', [])
        draft.setdefault('template_path', 'Шаблон отчета.docx')
        draft.setdefault('position', CLUBS_DATA.get(draft.get('club_name',''), {}).get('position', ''))
        draft.setdefault('responsible', CLUBS_DATA.get(draft.get('club_name',''), {}).get('responsible', ''))
        draft['events'] = events
        draft['current_event_index'] = len(events)
        save_report_draft(user_id, draft)

        # Показываем список мероприятий для редактирования
        events_text = "\n".join(
            f"{i+1}. {e.get('event_name', '—')}"
            for i, e in enumerate(events)
        )
        builder = InlineKeyboardBuilder()
        builder.add(
            InlineKeyboardButton(text="➕ Добавить мероприятие", callback_data="add_event"),
            InlineKeyboardButton(text="✏️ Редактировать мероприятие", callback_data="edit_event_list"),
            InlineKeyboardButton(text="✅ Сформировать отчёт", callback_data="generate_report"),
        )
        builder.adjust(1)
        await callback.message.edit_text(
            f"✏️ <b>Редактирование отчёта</b>\n\n"
            f"Текущие мероприятия ({len(events)}):\n{events_text}\n\n"
            "Выберите действие:",
            reply_markup=builder.as_markup(),
            parse_mode="HTML",
        )
        await state.set_state(ReportStates.waiting_for_more_events)

    elif document_type == "plan":
        draft = get_plan_draft(user_id)
        if not draft or 'club_name' not in draft or 'month_year' not in draft:
            await callback.answer("❌ Черновик не найден", show_alert=True)
            return

        plan_items = get_existing_plan_items(user_id, draft['club_name'], draft['month_year'])
        if not plan_items:
            await callback.message.edit_text(
                "⚠️ Данные этого плана недоступны для редактирования "
                "(документ создан до обновления бота).\n\n"
                "Нажмите «🔄 Перезаполнить полностью» чтобы создать новый план.",
                reply_markup=get_overwrite_keyboard("plan"),
            )
            await callback.answer()
            return

        # Восстанавливаем черновик
        draft.setdefault('position', CLUBS_DATA.get(draft.get('club_name',''), {}).get('position', ''))
        draft.setdefault('responsible', CLUBS_DATA.get(draft.get('club_name',''), {}).get('responsible', ''))
        draft.setdefault('template_path', 'Шаблон плана.docx')
        draft['plan_items'] = plan_items
        draft['current_plan_index'] = len(plan_items)
        save_plan_draft(user_id, draft)

        items_text = "\n".join(
            f"{i+1}. {item.get('event_name', '—')}"
            for i, item in enumerate(plan_items)
        )
        builder = InlineKeyboardBuilder()
        builder.add(
            InlineKeyboardButton(text="➕ Добавить пункт", callback_data="add_plan_item"),
            InlineKeyboardButton(text="✏️ Редактировать пункт", callback_data="edit_plan_item_list"),
            InlineKeyboardButton(text="✅ Сформировать план", callback_data="generate_plan"),
        )
        builder.adjust(1)
        await callback.message.edit_text(
            f"✏️ <b>Редактирование плана</b>\n\n"
            f"Текущие пункты ({len(plan_items)}):\n{items_text}\n\n"
            "Выберите действие:",
            reply_markup=builder.as_markup(),
            parse_mode="HTML",
        )
        await state.set_state(PlanStates.waiting_for_more_plan_items)

    await callback.answer()


@dp.callback_query(F.data == "edit_event_list")
async def process_edit_event_list(callback: types.CallbackQuery, state: FSMContext):
    """Показывает список мероприятий с кнопками для выбора нужного."""
    user_id = callback.from_user.id
    draft = get_report_draft(user_id)
    if not draft or not draft.get('events'):
        await callback.answer("❌ Нет мероприятий", show_alert=True)
        return
    builder = InlineKeyboardBuilder()
    for i, event in enumerate(draft['events']):
        builder.add(InlineKeyboardButton(
            text=f"✏️ {i+1}. {event.get('event_name', '—')[:30]}",
            callback_data=f"edit_event_{i}",
        ))
    builder.add(InlineKeyboardButton(text="⬅️ Назад", callback_data="back_to_report_edit_menu"))
    builder.adjust(1)
    await callback.message.edit_text(
        "Выберите мероприятие для редактирования:",
        reply_markup=builder.as_markup(),
    )
    await callback.answer()


@dp.callback_query(F.data == "edit_plan_item_list")
async def process_edit_plan_item_list(callback: types.CallbackQuery, state: FSMContext):
    """Показывает список пунктов плана с кнопками для выбора нужного."""
    user_id = callback.from_user.id
    draft = get_plan_draft(user_id)
    if not draft or not draft.get('plan_items'):
        await callback.answer("❌ Нет пунктов", show_alert=True)
        return
    builder = InlineKeyboardBuilder()
    for i, item in enumerate(draft['plan_items']):
        builder.add(InlineKeyboardButton(
            text=f"✏️ {i+1}. {item.get('event_name', '—')[:30]}",
            callback_data=f"edit_plan_item_{i}",
        ))
    builder.add(InlineKeyboardButton(text="⬅️ Назад", callback_data="back_to_plan_edit_menu"))
    builder.adjust(1)
    await callback.message.edit_text(
        "Выберите пункт для редактирования:",
        reply_markup=builder.as_markup(),
    )
    await callback.answer()


@dp.callback_query(F.data == "back_to_report_edit_menu")
async def back_to_report_edit_menu(callback: types.CallbackQuery, state: FSMContext):
    user_id = callback.from_user.id
    draft = get_report_draft(user_id)
    events = draft.get('events', []) if draft else []
    events_text = "\n".join(f"{i+1}. {e.get('event_name','—')}" for i, e in enumerate(events))
    builder = InlineKeyboardBuilder()
    builder.add(
        InlineKeyboardButton(text="➕ Добавить мероприятие", callback_data="add_event"),
        InlineKeyboardButton(text="✏️ Редактировать мероприятие", callback_data="edit_event_list"),
        InlineKeyboardButton(text="✅ Сформировать отчёт", callback_data="generate_report"),
    )
    builder.adjust(1)
    await callback.message.edit_text(
        f"✏️ <b>Редактирование отчёта</b>\n\nМероприятия ({len(events)}):\n{events_text}\n\nВыберите действие:",
        reply_markup=builder.as_markup(),
        parse_mode="HTML",
    )
    await state.set_state(ReportStates.waiting_for_more_events)
    await callback.answer()


@dp.callback_query(F.data == "back_to_plan_edit_menu")
async def back_to_plan_edit_menu(callback: types.CallbackQuery, state: FSMContext):
    user_id = callback.from_user.id
    draft = get_plan_draft(user_id)
    items = draft.get('plan_items', []) if draft else []
    items_text = "\n".join(f"{i+1}. {it.get('event_name','—')}" for i, it in enumerate(items))
    builder = InlineKeyboardBuilder()
    builder.add(
        InlineKeyboardButton(text="➕ Добавить пункт", callback_data="add_plan_item"),
        InlineKeyboardButton(text="✏️ Редактировать пункт", callback_data="edit_plan_item_list"),
        InlineKeyboardButton(text="✅ Сформировать план", callback_data="generate_plan"),
    )
    builder.adjust(1)
    await callback.message.edit_text(
        f"✏️ <b>Редактирование плана</b>\n\nПункты ({len(items)}):\n{items_text}\n\nВыберите действие:",
        reply_markup=builder.as_markup(),
        parse_mode="HTML",
    )
    await state.set_state(PlanStates.waiting_for_more_plan_items)
    await callback.answer()


@dp.callback_query(F.data == "cancel_overwrite")
async def process_cancel_overwrite(callback: types.CallbackQuery, state: FSMContext):
    """Пользователь выбрал оставить документ без изменений."""
    await callback.message.edit_text(
        "✅ Документ оставлен без изменений.\n\n"
        "Используйте /start чтобы вернуться в меню."
    )
    await state.clear()
    await callback.answer()

def get_directions_full_text() -> str:
    """Возвращает полный текст всех направлений для отображения"""
    return (
        "📋 Направления реализации государственной молодежной политики:\n\n"
        "1) Воспитание гражданственности, патриотизма, преемственности традиций, уважения к отечественной истории, историческим, национальным и иным традициям народов Российской Федерации;\n\n"
        "2) Обеспечение межнационального (межэтнического) и межконфессионального согласия в молодежной среде, профилактика и предупреждение проявлений экстремизма в деятельности молодежных объединений;\n\n"
        "3) Поддержка молодых граждан, оказавшихся в трудной жизненной ситуации, инвалидов из числа молодых граждан, а также лиц из числа детей-сирот и детей, оставшихся без попечения родителей;\n\n"
        "4) Поддержка инициатив молодежи;\n\n"
        "5) Содействие общественной деятельности, направленной на поддержку молодежи;\n\n"
        "6) Организация досуга, отдыха, оздоровления молодежи, формирование условий для занятий физической культурой, спортом, содействие здоровому образу жизни молодежи;\n\n"
        "7) Предоставление социальных услуг молодежи;\n\n"
        "8) Содействие решению жилищных проблем молодежи, молодых семей;\n\n"
        "9) Поддержка молодых семей;\n\n"
        "10) Содействие образованию молодежи, научной, научно-технической деятельности молодежи;\n\n"
        "11) Организация подготовки специалистов по работе с молодежью;\n\n"
        "12) Выявление, сопровождение и поддержка молодежи, проявившей одаренность;\n\n"
        "13) Развитие института наставничества;\n\n"
        "14) Обеспечение гарантий в сфере труда и занятости молодежи, содействие трудоустройству молодых граждан, в том числе посредством студенческих отрядов, профессиональному развитию молодых специалистов;\n\n"
        "15) Поддержка и содействие предпринимательской деятельности молодежи;\n\n"
        "16) Поддержка деятельности молодежных общественных объединений;\n\n"
        "17) Содействие участию молодежи в добровольческой (волонтерской) деятельности;\n\n"
        "18) Содействие международному и межрегиональному сотрудничеству в сфере молодежной политики;\n\n"
        "19) Предупреждение правонарушений и антиобщественных действий молодежи;\n\n"
        "20) Поддержка деятельности по созданию и распространению, в том числе в информационно-телекоммуникационной сети \"Интернет\", в средствах массовой информации произведений науки, искусства, литературы и других произведений, направленных на укрепление гражданской идентичности и духовно-нравственных ценностей молодежи;\n\n"
        "21) Проведение научно-аналитических исследований по вопросам молодежной политики."
    )

# Обработчики callback-ов для выбора времени
@dp.callback_query(F.data.startswith("time_"))
async def process_time_selection(callback: types.CallbackQuery, state: FSMContext):
    """Обработка выбора времени"""
    user_id = callback.from_user.id
    selected_time = callback.data.replace("time_", "")
    
    current_state = await state.get_state()
    
    if current_state == ReportStates.waiting_for_event_time.state:
        draft = get_report_draft(user_id)
        if not draft or not draft.get('events'):
            await callback.answer("❌ Ошибка: нет активного мероприятия")
            return
        
        current_event_index = draft.get('current_event_index', 0)
        current_time_data = draft['events'][current_event_index].get('time_data', {})
        
        # Если время начала еще не выбрано, сохраняем его
        if 'start_time' not in current_time_data:
            current_time_data['start_time'] = selected_time
            draft['events'][current_event_index]['time_data'] = current_time_data
            save_report_draft(user_id, draft)
            
            await callback.message.edit_text(
                f"✅ Выбрано время начала: {selected_time}\n\n"
                f"Теперь выберите время окончания:",
                reply_markup=get_time_keyboard(selected_start_time=selected_time)
            )
        else:
            # Время начала уже выбрано, сохраняем время окончания
            start_time = current_time_data['start_time']
            end_time = selected_time
            
            # Проверяем, что время окончания позже времени начала
            start_hour = int(start_time.split(':')[0])
            end_hour = int(end_time.split(':')[0])
            
            if end_hour <= start_hour:
                await callback.answer("❌ Время окончания должно быть позже времени начала")
                return
            
            # Формируем итоговое время в формате "9:00-18:00"
            final_time = f"{start_time}-{end_time}"
            draft['events'][current_event_index]['time'] = final_time
            # Очищаем временные данные
            draft['events'][current_event_index]['time_data'] = {}
            save_report_draft(user_id, draft)
            
            await callback.message.edit_text(
                f"✅ Время сохранено: {final_time}\n\n"
                f"9️⃣ Введите краткое содержание мероприятия:"
            )
            await state.set_state(ReportStates.waiting_for_description)
    
    await callback.answer()

@dp.callback_query(F.data == "custom_time")
async def process_custom_time(callback: types.CallbackQuery, state: FSMContext):
    """Обработка выбора кастомного ввода времени"""
    user_id = callback.from_user.id
    
    draft = get_report_draft(user_id)
    if draft and draft.get('events'):
        current_event_index = draft.get('current_event_index', 0)
        current_time_data = draft['events'][current_event_index].get('time_data', {})
        
        if 'start_time' in current_time_data:
            # Уже выбрано время начала, предлагаем ввести время окончания
            start_time = current_time_data['start_time']
            await callback.message.edit_text(
                f"⏰ Вы уже выбрали время начала: {start_time}\n\n"
                f"Теперь введите время окончания в формате ЧЧ:MM (например: 18:00):\n\n"
                f"Или выберите из стандартных вариантов:",
                reply_markup=get_time_keyboard(selected_start_time=start_time)
            )
        else:
            # Время начала еще не выбрано
            await callback.message.edit_text(
                "⏰ Введите время начала в формате ЧЧ:MM (например: 14:00):\n\n"
                "Или выберите из стандартных вариантов:",
                reply_markup=get_time_keyboard()
            )
    else:
        await callback.message.edit_text(
            "⏰ Введите время в любом формате:\n\n"
            "• Одно время: 14:00\n"
            "• Диапазон: 14:00-16:00\n"
            "• Текстово: весь день\n\n"
            "Или выберите из стандартных вариантов:",
            reply_markup=get_time_keyboard()
        )
    
    await callback.answer()

@dp.callback_query(F.data == "skip_end_time")
async def process_skip_end_time(callback: types.CallbackQuery, state: FSMContext):
    """Обработка пропуска ввода времени окончания"""
    user_id = callback.from_user.id
    
    draft = get_report_draft(user_id)
    if not draft or not draft.get('events'):
        await callback.answer("❌ Нет активного черновика")
        return
    
    current_event_index = draft.get('current_event_index', 0)
    
    await callback.message.edit_text(
        f"✅ Время сохранено: {draft['events'][current_event_index]['time']}\n\n"
        f"9️⃣ Введите краткое содержание мероприятия:"
    )
    await state.set_state(ReportStates.waiting_for_description)
    await callback.answer()

# Обработчики callback-ов для выбора стандартных ссылок
@dp.callback_query(F.data.startswith("link_"))
async def process_link_selection(callback: types.CallbackQuery, state: FSMContext):
    """Обработка выбора ссылки"""
    user_id = callback.from_user.id
    link_type = callback.data.replace("link_", "")
    
    current_state = await state.get_state()
    
    if current_state == PlanStates.waiting_for_info_link.state:
        draft = get_plan_draft(user_id)
        if draft and draft.get('plan_items'):
            current_plan_index = draft.get('current_plan_index', 0)
            
            if link_type == "custom":
                await callback.message.edit_text(
                    "✏️ Введите свою ссылку (или '-' для прочерка):"
                )
                return
            elif link_type == "dash":
                # Устанавливаем прочерк
                draft['plan_items'][current_plan_index]['info_link'] = "-"
                save_plan_draft(user_id, draft)
                
                plan_count = len(draft['plan_items'])
                current_item = draft['plan_items'][current_plan_index]
                
                await callback.message.edit_text(
                    f"✅ Пункт плана #{plan_count} сохранен!\n\n"
                    f"📋 Направление: {current_item['direction'][:100]}...\n"
                    f"📝 Мероприятие: {current_item['event_name']}\n"
                    f"🏷️ Статус: {current_item['status']}\n"
                    f"📅 Дат выбрано: {len(current_item['dates'])}\n"
                    f"📍 Место: {current_item['place']}\n"
                    f"👥 Участников: {current_item['participants']} чел.\n"
                    f"🔗 Ссылка: -\n\n"
                    f"Выберите дальнейшее действие:",
                    reply_markup=get_final_keyboard("plan")
                )
                await state.set_state(PlanStates.waiting_for_more_plan_items)
                return
            else:
                selected_link = STANDARD_LINKS.get(link_type)
                if selected_link:
                    draft['plan_items'][current_plan_index]['info_link'] = selected_link
                    save_plan_draft(user_id, draft)
                    
                    plan_count = len(draft['plan_items'])
                    current_item = draft['plan_items'][current_plan_index]
                    
                    await callback.message.edit_text(
                        f"✅ Пункт плана #{plan_count} сохранен!\n\n"
                        f"📋 Направление: {current_item['direction'][:100]}...\n"
                        f"📝 Мероприятие: {current_item['event_name']}\n"
                        f"🏷️ Статус: {current_item['status']}\n"
                        f"📅 Дат выбрано: {len(current_item['dates'])}\n"
                        f"📍 Место: {current_item['place']}\n"
                        f"👥 Участников: {current_item['participants']} чел.\n"
                        f"🔗 Ссылка: {selected_link}\n\n"
                        f"Выберите дальнейшее действие:",
                        reply_markup=get_final_keyboard("plan")
                    )
                    await state.set_state(PlanStates.waiting_for_more_plan_items)
    
    elif current_state == ReportStates.waiting_for_link.state:
        draft = get_report_draft(user_id)
        if draft and draft.get('events'):
            current_event_index = draft.get('current_event_index', 0)
            
            if link_type == "custom":
                await callback.message.edit_text(
                    "✏️ Введите свою ссылку (или '-' для прочерка):\n\n"
                    "• Можно ввести несколько ссылок, каждую с новой строки\n"
                    "• Или введите '-' для прочерка\n\n"
                    "Пример:\nhttps://t.me/link1\nhttps://vk.com/link2"
                )
                return
            elif link_type == "dash":
                # Устанавливаем прочерк
                draft['events'][current_event_index]['link'] = "-"
                save_report_draft(user_id, draft)
                
                await callback.message.edit_text(
                    f"✅ Ссылка сохранена: -\n\n"
                    f"👥 Введите количество участников 14-17 лет:\n(если нет — введите 0)"
                )
                await state.set_state(ReportStates.waiting_for_age_14_17)
                return
    
    await callback.answer()

# Обработчики callback-ов для редактирования
@dp.callback_query(F.data == "edit_events")
async def edit_events_callback(callback: types.CallbackQuery, state: FSMContext):
    """Обработка редактирования мероприятий"""
    user_id = callback.from_user.id
    
    draft = get_report_draft(user_id)
    if not draft or not draft.get('events'):
        await callback.answer("❌ Нет мероприятий для редактирования")
        return
    
    events = draft['events']
    await callback.message.edit_text(
        "📋 Выберите мероприятие для редактирования:",
        reply_markup=get_events_list_keyboard(events, "report")
    )
    await callback.answer()

@dp.callback_query(F.data == "edit_plan_items")
async def edit_plan_items_callback(callback: types.CallbackQuery, state: FSMContext):
    """Обработка редактирования пунктов плана"""
    user_id = callback.from_user.id
    
    draft = get_plan_draft(user_id)
    if not draft or not draft.get('plan_items'):
        await callback.answer("❌ Нет пунктов плана для редактирования")
        return
    
    plan_items = draft['plan_items']
    await callback.message.edit_text(
        "📋 Выберите пункт плана для редактирования:",
        reply_markup=get_events_list_keyboard(plan_items, "plan")
    )
    await callback.answer()

@dp.callback_query(F.data.startswith("edit_event_"))
async def edit_specific_event(callback: types.CallbackQuery, state: FSMContext):
    """Обработка редактирования конкретного мероприятия"""
    user_id = callback.from_user.id
    event_index = int(callback.data.replace("edit_event_", ""))
    
    draft = get_report_draft(user_id)
    if not draft or not draft.get('events') or event_index >= len(draft['events']):
        await callback.answer("❌ Ошибка: мероприятие не найдено")
        return
    
    event = draft['events'][event_index]
    events_count = len(draft['events'])
    
    event_text = (
        f"✏️ Редактирование мероприятия #{event_index + 1}\n\n"
        f"📝 Название: {event.get('event_name', 'Не указано')}\n"
        f"📅 Даты: {len(event.get('dates', []))} выбрано\n"
        f"📍 Место: {event.get('place', 'Не указано')}\n"
        f"⏰ Время: {event.get('time', 'Не указано')}\n"
        f"📝 Описание: {event.get('description', 'Не указано')[:50]}...\n"
        f"🔗 Ссылка: {event.get('link', 'Не указана')}\n"
        f"👥 Участники 14-17: {event.get('age_14_17', 0)}\n"
        f"👥 Участники 18-29: {event.get('age_18_29', 0)}\n"
        f"👥 Участники 30-35: {event.get('age_30_35', 0)}\n\n"
        f"Выберите что хотите изменить:"
    )
    
    await callback.message.edit_text(
        event_text,
        reply_markup=get_edit_report_keyboard(event_index, events_count)
    )
    await state.update_data(editing_event_index=event_index)
    await state.set_state(ReportStates.editing_event)
    await callback.answer()

@dp.callback_query(F.data.startswith("edit_plan_item_"))
async def edit_specific_plan_item(callback: types.CallbackQuery, state: FSMContext):
    """Обработка редактирования конкретного пункта плана"""
    user_id = callback.from_user.id
    item_index = int(callback.data.replace("edit_plan_item_", ""))
    
    draft = get_plan_draft(user_id)
    if not draft or not draft.get('plan_items') or item_index >= len(draft['plan_items']):
        await callback.answer("❌ Ошибка: пункт плана не найден")
        return
    
    item = draft['plan_items'][item_index]
    items_count = len(draft['plan_items'])
    
    item_text = (
        f"✏️ Редактирование пункта плана #{item_index + 1}\n\n"
        f"🎯 Направление: {item.get('direction', 'Не указано')[:50]}...\n"
        f"📝 Название: {item.get('event_name', 'Не указано')}\n"
        f"🏷️ Статус: {item.get('status', 'Не указан')}\n"
        f"📅 Даты: {len(item.get('dates', []))} выбрано\n"
        f"📍 Место: {item.get('place', 'Не указано')}\n"
        f"👥 Участники: {item.get('participants', 0)} чел.\n"
        f"🔗 Ссылка: {item.get('info_link', 'Не указана')}\n\n"
        f"Выберите что хотите изменить:"
    )
    
    await callback.message.edit_text(
        item_text,
        reply_markup=get_edit_plan_keyboard(item_index, items_count)
    )
    await state.update_data(editing_plan_index=item_index)
    await state.set_state(PlanStates.editing_plan_item)
    await callback.answer()

@dp.callback_query(F.data.startswith("edit_report_"))
async def edit_report_field(callback: types.CallbackQuery, state: FSMContext):
    """Обработка редактирования конкретного поля отчета"""
    user_id = callback.from_user.id
    field_data = callback.data.replace("edit_report_", "")
    field_parts = field_data.split('_')
    
    if len(field_parts) < 2:
        await callback.answer("❌ Ошибка формата")
        return
    
    field_name = field_parts[0]
    event_index = int(field_parts[1])
    
    await state.update_data(
        editing_field=field_name,
        editing_event_index=event_index
    )
    
    field_prompts = {
        'name': "Введите новое название мероприятия:",
        'dates': "Выберите новые даты проведения:",
        'place': "Выберите или введите новое место проведения:",
        'time': "Выберите новое время проведения:",
        'desc': "Введите новое описание мероприятия:",
        'link': "Выберите или введите новую ссылку:",
        'age14': "Введите новое количество участников 14-17 лет:",
        'age18': "Введите новое количество участников 18-29 лет:",
        'age30': "Введите новое количество участников 30-35 лет:"
    }
    
    prompt = field_prompts.get(field_name)
    if not prompt:
        await callback.answer("❌ Неизвестное поле")
        return
    
    if field_name == 'dates':
        draft = get_report_draft(user_id)
        if draft:
            month_num, year = parse_month_year(draft['month_year'])
            # Получаем текущие выбранные даты для этого мероприятия
            current_dates = draft['events'][event_index].get('dates', [])
            calendar_markup = get_month_calendar(year, month_num, current_dates)
            await callback.message.edit_text(
                f"📅 Выберите новые даты проведения:\n\n{prompt}",
                reply_markup=calendar_markup
            )
            # Сохраняем выбранные даты в состоянии для редактирования
            await state.update_data(editing_dates=current_dates)
            return
    elif field_name == 'time':
        await callback.message.edit_text(
            f"⏰ Выберите время проведения:\n\n{prompt}",
            reply_markup=get_time_keyboard()
        )
        return
    elif field_name == 'link':
        await callback.message.edit_text(
            f"🔗 Выберите или введите ссылку:\n\n{prompt}",
            reply_markup=get_standard_links_keyboard(for_plan=False)
        )
        return
    elif field_name == 'place':
        await callback.message.edit_text(
            f"📍 Выберите или введите место проведения:\n\n{prompt}",
            reply_markup=get_place_keyboard()
        )
        return
    
    await callback.message.edit_text(prompt)
    await callback.answer()

@dp.callback_query(F.data.startswith("edit_plan_"))
async def edit_plan_field(callback: types.CallbackQuery, state: FSMContext):
    """Обработка редактирования конкретного поля плана"""
    user_id = callback.from_user.id
    field_data = callback.data.replace("edit_plan_", "")
    field_parts = field_data.split('_')
    
    if len(field_parts) < 2:
        await callback.answer("❌ Ошибка формата")
        return
    
    field_name = field_parts[0]
    item_index = int(field_parts[1])
    
    await state.update_data(
        editing_plan_field=field_name,
        editing_plan_index=item_index
    )
    
    field_prompts = {
        'direction': "Выберите новое направление:",
        'name': "Введите новое название мероприятия:",
        'status': "Выберите новый статус:",
        'dates': "Выберите новые даты проведения:",
        'place': "Введите новое место проведения:",
        'participants': "Введите новое количество участников:",
        'link': "Выберите или введите новую ссылку:"
    }
    
    prompt = field_prompts.get(field_name)
    if not prompt:
        await callback.answer("❌ Неизвестное поле")
        return
    
    if field_name == 'direction':
        await callback.message.edit_text(
            f"🎯 Выберите направление:\n\n{prompt}",
            reply_markup=get_directions_keyboard()
        )
    elif field_name == 'status':
        await callback.message.edit_text(
            f"🏷️ Выберите статус:\n\n{prompt}",
            reply_markup=get_status_keyboard()
        )
    elif field_name == 'dates':
        draft = get_plan_draft(user_id)
        if draft:
            month_num, year = parse_month_year(draft['month_year'])
            # Получаем текущие выбранные даты для этого пункта плана
            current_dates = draft['plan_items'][item_index].get('dates', [])
            calendar_markup = get_month_calendar(year, month_num, current_dates)
            await callback.message.edit_text(
                f"📅 Выберите новые даты проведения:\n\n{prompt}",
                reply_markup=calendar_markup
            )
            # Сохраняем выбранные даты в состоянии для редактирования
            await state.update_data(editing_dates=current_dates)
    elif field_name == 'link':
        await callback.message.edit_text(
            f"🔗 Выберите или введите ссылку:\n\n{prompt}",
            reply_markup=get_standard_links_keyboard(for_plan=True)
        )
    else:
        await callback.message.edit_text(prompt)
    
    await callback.answer()

# Обработчик для выбора дней недели
@dp.callback_query(F.data.startswith("weekday_"))
async def process_weekday_selection(callback: types.CallbackQuery, state: FSMContext):
    """Обработка выбора дня недели"""
    user_id = callback.from_user.id
    weekday_index = int(callback.data.replace("weekday_", ""))
    
    current_state = await state.get_state()
    
    if current_state in [ReportStates.waiting_for_event_dates.state, PlanStates.waiting_for_event_dates.state, 
                         ReportStates.editing_event.state, PlanStates.editing_plan_item.state]:
        
        draft = None
        if current_state in [ReportStates.waiting_for_event_dates.state, ReportStates.editing_event.state]:
            draft = get_report_draft(user_id)
        else:
            draft = get_plan_draft(user_id)
        
        if not draft:
            await callback.answer("❌ Нет активного черновика")
            return
        
        month_num, year = parse_month_year(draft['month_year'])
        
        # Получаем все даты для выбранного дня недели
        weekday_dates = get_weekday_dates(year, month_num, weekday_index)
        
        if current_state == ReportStates.waiting_for_event_dates.state:
            selected_dates = draft.get('selected_dates', [])
            # Добавляем все даты этого дня недели
            for date_str in weekday_dates:
                if date_str not in selected_dates:
                    selected_dates.append(date_str)
            draft['selected_dates'] = selected_dates
            save_report_draft(user_id, draft)
            
            calendar_markup = get_month_calendar(year, month_num, selected_dates)
            selected_dates_text = "\n".join([f"• {format_date_for_display(d)}" for d in sorted(selected_dates)])
            
            await callback.message.edit_text(
                f"✅ Добавлены все {['Понедельники', 'Вторники', 'Среды', 'Четверги', 'Пятницы', 'Субботы', 'Воскресенья'][weekday_index]} месяца!\n\n"
                f"📅 Выберите даты проведения мероприятия:\n\n"
                f"📋 Выбранные даты:\n{selected_dates_text if selected_dates_text else 'Нет выбранных дат'}\n\n"
                f"🗓️ Календарь {draft['month_year']}:",
                reply_markup=calendar_markup
            )
            
        elif current_state == PlanStates.waiting_for_event_dates.state:
            current_plan_index = draft.get('current_plan_index', 0)
            current_dates = draft['plan_items'][current_plan_index].get('dates', [])
            # Добавляем все даты этого дня недели
            for date_str in weekday_dates:
                if date_str not in current_dates:
                    current_dates.append(date_str)
            draft['plan_items'][current_plan_index]['dates'] = current_dates
            save_plan_draft(user_id, draft)
            
            calendar_markup = get_month_calendar(year, month_num, current_dates)
            selected_dates_text = "\n".join([f"• {format_date_for_display(d)}" for d in sorted(current_dates)])
            
            await callback.message.edit_text(
                f"✅ Добавлены все {['Понедельники', 'Вторники', 'Среды', 'Четверги', 'Пятницы', 'Субботы', 'Воскресенья'][weekday_index]} месяца!\n\n"
                f"📅 Выберите даты проведения мероприятия:\n\n"
                f"📋 Выбранные даты:\n{selected_dates_text if selected_dates_text else 'Нет выбранных дат'}\n\n"
                f"🗓️ Календарь {draft['month_year']}:",
                reply_markup=calendar_markup
            )
        
        elif current_state in [ReportStates.editing_event.state, PlanStates.editing_plan_item.state]:
            data = await state.get_data()
            editing_dates = data.get('editing_dates', [])
            # Добавляем все даты этого дня недели
            for date_str in weekday_dates:
                if date_str not in editing_dates:
                    editing_dates.append(date_str)
            await state.update_data(editing_dates=editing_dates)
            
            calendar_markup = get_month_calendar(year, month_num, editing_dates)
            selected_dates_text = "\n".join([f"• {format_date_for_display(d)}" for d in sorted(editing_dates)])
            
            await callback.message.edit_text(
                f"✅ Добавлены все {['Понедельники', 'Вторники', 'Среды', 'Четверги', 'Пятницы', 'Субботы', 'Воскресенья'][weekday_index]} месяца!\n\n"
                f"📅 Выберите новые даты проведения:\n\n"
                f"📋 Выбранные даты:\n{selected_dates_text if selected_dates_text else 'Нет выбранных дат'}\n\n"
                f"🗓️ Календарь {draft['month_year']}:",
                reply_markup=calendar_markup
            )
    
    await callback.answer()

# Обработчик для быстрого выбора места
@dp.callback_query(F.data == "place_dom_molodezhi")
async def process_place_dom_molodezhi(callback: types.CallbackQuery, state: FSMContext):
    """Обработка выбора Дома молодежи"""
    user_id = callback.from_user.id
    current_state = await state.get_state()
    
    place_text = "г. Сочи, ул. Навагинская, 9"
    
    if current_state == ReportStates.waiting_for_event_place.state:
        draft = get_report_draft(user_id)
        if draft and draft.get('events'):
            current_event_index = draft.get('current_event_index', 0)
            draft['events'][current_event_index]['place'] = place_text
            save_report_draft(user_id, draft)
            
            await callback.message.edit_text(
                f"✅ Место сохранено: {place_text}\n\n"
                "8️⃣ Теперь выберите время проведения мероприятия:",
                reply_markup=get_time_keyboard()
            )
            await state.set_state(ReportStates.waiting_for_event_time)
    
    elif current_state == PlanStates.waiting_for_event_place.state:
        draft = get_plan_draft(user_id)
        if draft and draft.get('plan_items'):
            current_plan_index = draft.get('current_plan_index', 0)
            draft['plan_items'][current_plan_index]['place'] = place_text
            save_plan_draft(user_id, draft)
            
            await callback.message.edit_text(
                f"✅ Место сохранено: {place_text}\n\n"
                "👥 Теперь введите предполагаемое количество участников:"
            )
            await state.set_state(PlanStates.waiting_for_participants_count)
    
    elif current_state in [ReportStates.editing_event.state, PlanStates.editing_plan_item.state]:
        data = await state.get_data()
        if current_state == ReportStates.editing_event.state:
            event_index = data.get('editing_event_index')
            draft = get_report_draft(user_id)
            if draft and draft.get('events') and event_index < len(draft['events']):
                draft['events'][event_index]['place'] = place_text
                save_report_draft(user_id, draft)
                
                # Возвращаем к меню редактирования
                event = draft['events'][event_index]
                events_count = len(draft['events'])
                
                event_text = (
                    f"✅ Место обновлено: {place_text}\n\n"
                    f"✏️ Редактирование мероприятия #{event_index + 1}\n\n"
                    f"📝 Название: {event.get('event_name', 'Не указано')}\n"
                    f"📅 Даты: {len(event.get('dates', []))} выбрано\n"
                    f"📍 Место: {event.get('place', 'Не указано')}\n"
                    f"⏰ Время: {event.get('time', 'Не указано')}\n"
                    f"📝 Описание: {event.get('description', 'Не указано')[:50]}...\n"
                    f"🔗 Ссылка: {event.get('link', 'Не указана')}\n"
                    f"👥 Участники 14-17: {event.get('age_14_17', 0)}\n"
                    f"👥 Участники 18-29: {event.get('age_18_29', 0)}\n"
                    f"👥 Участники 30-35: {event.get('age_30_35', 0)}\n\n"
                    f"Выберите что хотите изменить:"
                )
                
                await callback.message.edit_text(
                    event_text,
                    reply_markup=get_edit_report_keyboard(event_index, events_count)
                )
        
        else:  # PlanStates.editing_plan_item.state
            item_index = data.get('editing_plan_index')
            draft = get_plan_draft(user_id)
            if draft and draft.get('plan_items') and item_index < len(draft['plan_items']):
                draft['plan_items'][item_index]['place'] = place_text
                save_plan_draft(user_id, draft)
                
                # Возвращаем к меню редактирования
                item = draft['plan_items'][item_index]
                items_count = len(draft['plan_items'])
                
                item_text = (
                    f"✅ Место обновлено: {place_text}\n\n"
                    f"✏️ Редактирование пункта плана #{item_index + 1}\n\n"
                    f"🎯 Направление: {item.get('direction', 'Не указано')[:50]}...\n"
                    f"📝 Название: {item.get('event_name', 'Не указано')}\n"
                    f"🏷️ Статус: {item.get('status', 'Не указан')}\n"
                    f"📅 Даты: {len(item.get('dates', []))} выбрано\n"
                    f"📍 Место: {item.get('place', 'Не указано')}\n"
                    f"👥 Участники: {item.get('participants', 0)} чел.\n"
                    f"🔗 Ссылка: {item.get('info_link', 'Не указана')}\n\n"
                    f"Выберите что хотите изменить:"
                )
                
                await callback.message.edit_text(
                    item_text,
                    reply_markup=get_edit_plan_keyboard(item_index, items_count)
                )
    
    await callback.answer()

@dp.callback_query(F.data == "place_custom")
async def process_place_custom(callback: types.CallbackQuery, state: FSMContext):
    """Обработка выбора кастомного ввода места"""
    current_state = await state.get_state()
    
    if current_state == ReportStates.waiting_for_event_place.state:
        await callback.message.edit_text(
            "📍 Введите место проведения мероприятия:"
        )
    elif current_state == PlanStates.waiting_for_event_place.state:
        await callback.message.edit_text(
            "📍 Введите место проведения мероприятия:"
        )
    
    await callback.answer()

# Обработчик для выбора дат при редактировании
@dp.callback_query(ReportStates.editing_event, F.data.startswith("date_"))
async def process_edit_date_selection(callback: types.CallbackQuery, state: FSMContext):
    """Обработка выбора даты при редактировании"""
    user_id = callback.from_user.id
    date_str = callback.data.replace("date_", "")
    
    data = await state.get_data()
    event_index = data.get('editing_event_index')
    editing_dates = data.get('editing_dates', [])
    
    if date_str in editing_dates:
        editing_dates.remove(date_str)
    else:
        editing_dates.append(date_str)
    
    await state.update_data(editing_dates=editing_dates)
    
    draft = get_report_draft(user_id)
    if draft:
        month_num, year = parse_month_year(draft['month_year'])
        calendar_markup = get_month_calendar(year, month_num, editing_dates)
        
        selected_dates_text = "\n".join([f"• {format_date_for_display(d)}" for d in sorted(editing_dates)])
        
        await callback.message.edit_text(
            f"📅 Выберите новые даты проведения мероприятия:\n\n"
            f"📋 Выбранные даты:\n{selected_dates_text if selected_dates_text else 'Нет выбранных дат'}\n\n"
            f"🗓️ Календарь {draft['month_year']}:",
            reply_markup=calendar_markup
        )
    
    await callback.answer()

@dp.callback_query(PlanStates.editing_plan_item, F.data.startswith("date_"))
async def process_edit_plan_date_selection(callback: types.CallbackQuery, state: FSMContext):
    """Обработка выбора даты при редактировании плана"""
    user_id = callback.from_user.id
    date_str = callback.data.replace("date_", "")
    
    data = await state.get_data()
    item_index = data.get('editing_plan_index')
    editing_dates = data.get('editing_dates', [])
    
    if date_str in editing_dates:
        editing_dates.remove(date_str)
    else:
        editing_dates.append(date_str)
    
    await state.update_data(editing_dates=editing_dates)
    
    draft = get_plan_draft(user_id)
    if draft:
        month_num, year = parse_month_year(draft['month_year'])
        calendar_markup = get_month_calendar(year, month_num, editing_dates)
        
        selected_dates_text = "\n".join([f"• {format_date_for_display(d)}" for d in sorted(editing_dates)])
        
        await callback.message.edit_text(
            f"📅 Выберите новые даты проведения мероприятия:\n\n"
            f"📋 Выбранные даты:\n{selected_dates_text if selected_dates_text else 'Нет выбранных дат'}\n\n"
            f"🗓️ Календарь {draft['month_year']}:",
            reply_markup=calendar_markup
        )
    
    await callback.answer()

# Обработчик завершения выбора дат при редактировании
@dp.callback_query(ReportStates.editing_event, F.data == "finish_dates")
async def finish_edit_date_selection(callback: types.CallbackQuery, state: FSMContext):
    """Завершение выбора дат при редактировании"""
    user_id = callback.from_user.id
    
    data = await state.get_data()
    event_index = data.get('editing_event_index')
    editing_dates = data.get('editing_dates', [])
    
    if not editing_dates:
        await callback.answer("❌ Выберите хотя бы одну дату")
        return
    
    draft = get_report_draft(user_id)
    if draft and draft.get('events') and event_index < len(draft['events']):
        draft['events'][event_index]['dates'] = editing_dates.copy()
        save_report_draft(user_id, draft)
        
        selected_dates_count = len(editing_dates)
        
        # Возвращаем к меню редактирования
        event = draft['events'][event_index]
        events_count = len(draft['events'])
        
        event_text = (
            f"✅ Обновлено дат: {selected_dates_count}\n\n"
            f"✏️ Редактирование мероприятия #{event_index + 1}\n\n"
            f"📝 Название: {event.get('event_name', 'Не указано')}\n"
            f"📅 Даты: {len(event.get('dates', []))} выбрано\n"
            f"📍 Место: {event.get('place', 'Не указано')}\n"
            f"⏰ Время: {event.get('time', 'Не указано')}\n"
            f"📝 Описание: {event.get('description', 'Не указано')[:50]}...\n"
            f"🔗 Ссылка: {event.get('link', 'Не указана')}\n"
            f"👥 Участники 14-17: {event.get('age_14_17', 0)}\n"
            f"👥 Участники 18-29: {event.get('age_18_29', 0)}\n"
            f"👥 Участники 30-35: {event.get('age_30_35', 0)}\n\n"
            f"Выберите что хотите изменить:"
        )
        
        await callback.message.edit_text(
            event_text,
            reply_markup=get_edit_report_keyboard(event_index, events_count)
        )
    
    await callback.answer()

@dp.callback_query(PlanStates.editing_plan_item, F.data == "finish_dates")
async def finish_edit_plan_date_selection(callback: types.CallbackQuery, state: FSMContext):
    """Завершение выбора дат при редактировании плана"""
    user_id = callback.from_user.id
    
    data = await state.get_data()
    item_index = data.get('editing_plan_index')
    editing_dates = data.get('editing_dates', [])
    
    if not editing_dates:
        await callback.answer("❌ Выберите хотя бы одну дату")
        return
    
    draft = get_plan_draft(user_id)
    if draft and draft.get('plan_items') and item_index < len(draft['plan_items']):
        draft['plan_items'][item_index]['dates'] = editing_dates.copy()
        save_plan_draft(user_id, draft)
        
        selected_dates_count = len(editing_dates)
        
        # Возвращаем к меню редактирования
        item = draft['plan_items'][item_index]
        items_count = len(draft['plan_items'])
        
        item_text = (
            f"✅ Обновлено дат: {selected_dates_count}\n\n"
            f"✏️ Редактирование пункта плана #{item_index + 1}\n\n"
            f"🎯 Направление: {item.get('direction', 'Не указано')[:50]}...\n"
            f"📝 Название: {item.get('event_name', 'Не указано')}\n"
            f"🏷️ Статус: {item.get('status', 'Не указан')}\n"
            f"📅 Даты: {len(item.get('dates', []))} выбрано\n"
            f"📍 Место: {item.get('place', 'Не указано')}\n"
            f"👥 Участники: {item.get('participants', 0)} чел.\n"
            f"🔗 Ссылка: {item.get('info_link', 'Не указана')}\n\n"
            f"Выберите что хотите изменить:"
        )
        
        await callback.message.edit_text(
            item_text,
            reply_markup=get_edit_plan_keyboard(item_index, items_count)
        )
    
    await callback.answer()

@dp.callback_query(F.data.startswith("confirm_delete_report_event_"))
async def confirm_delete_report_event(callback: types.CallbackQuery, state: FSMContext):
    """Подтверждение удаления мероприятия из отчета"""
    user_id = callback.from_user.id
    event_index = int(callback.data.replace("confirm_delete_report_event_", ""))

    draft = get_report_draft(user_id)
    if not draft or not draft.get('events') or event_index >= len(draft['events']):
        await callback.answer("❌ Ошибка: мероприятие не найдено")
        return

    event_name = draft['events'][event_index].get('event_name', f'Мероприятие #{event_index + 1}')
    builder = InlineKeyboardBuilder()
    builder.add(
        InlineKeyboardButton(text="🗑️ Да, удалить", callback_data=f"delete_report_event_{event_index}"),
        InlineKeyboardButton(text="↩️ Отмена", callback_data=f"edit_event_{event_index}")
    )
    builder.adjust(1)
    await callback.message.edit_text(
        f"⚠️ Удалить мероприятие?\n\n«{event_name[:60]}»\n\nЭто действие нельзя отменить.",
        reply_markup=builder.as_markup()
    )
    await callback.answer()

@dp.callback_query(F.data.startswith("delete_report_event_"))
async def delete_report_event(callback: types.CallbackQuery, state: FSMContext):
    """Удаление мероприятия из отчета"""
    user_id = callback.from_user.id
    event_index = int(callback.data.replace("delete_report_event_", ""))

    draft = get_report_draft(user_id)
    if not draft or not draft.get('events') or event_index >= len(draft['events']):
        await callback.answer("❌ Ошибка: мероприятие не найдено")
        return

    event_name = draft['events'][event_index].get('event_name', f'#{event_index + 1}')
    draft['events'].pop(event_index)

    if draft.get('current_event_index', 0) >= event_index:
        draft['current_event_index'] = max(0, draft.get('current_event_index', 0) - 1)

    save_report_draft(user_id, draft)

    await callback.message.edit_text(
        f"✅ Мероприятие «{event_name[:40]}» удалено!\n\n"
        f"Выберите дальнейшее действие:",
        reply_markup=get_final_keyboard("report")
    )
    await state.set_state(ReportStates.waiting_for_more_events)
    await callback.answer()

@dp.callback_query(F.data.startswith("confirm_delete_plan_item_"))
async def confirm_delete_plan_item(callback: types.CallbackQuery, state: FSMContext):
    """Подтверждение удаления пункта плана"""
    user_id = callback.from_user.id
    item_index = int(callback.data.replace("confirm_delete_plan_item_", ""))

    draft = get_plan_draft(user_id)
    if not draft or not draft.get('plan_items') or item_index >= len(draft['plan_items']):
        await callback.answer("❌ Ошибка: пункт плана не найден")
        return

    item_name = draft['plan_items'][item_index].get('event_name', f'Пункт #{item_index + 1}')
    builder = InlineKeyboardBuilder()
    builder.add(
        InlineKeyboardButton(text="🗑️ Да, удалить", callback_data=f"delete_plan_item_{item_index}"),
        InlineKeyboardButton(text="↩️ Отмена", callback_data=f"edit_plan_item_{item_index}")
    )
    builder.adjust(1)
    await callback.message.edit_text(
        f"⚠️ Удалить пункт плана?\n\n«{item_name[:60]}»\n\nЭто действие нельзя отменить.",
        reply_markup=builder.as_markup()
    )
    await callback.answer()

@dp.callback_query(F.data.startswith("delete_plan_item_"))
async def delete_plan_item(callback: types.CallbackQuery, state: FSMContext):
    """Удаление пункта из плана"""
    user_id = callback.from_user.id
    item_index = int(callback.data.replace("delete_plan_item_", ""))

    draft = get_plan_draft(user_id)
    if not draft or not draft.get('plan_items') or item_index >= len(draft['plan_items']):
        await callback.answer("❌ Ошибка: пункт плана не найден")
        return

    item_name = draft['plan_items'][item_index].get('event_name', f'#{item_index + 1}')
    draft['plan_items'].pop(item_index)

    if draft.get('current_plan_index', 0) >= item_index:
        draft['current_plan_index'] = max(0, draft.get('current_plan_index', 0) - 1)

    save_plan_draft(user_id, draft)

    await callback.message.edit_text(
        f"✅ Пункт «{item_name[:40]}» удален!\n\n"
        f"Выберите дальнейшее действие:",
        reply_markup=get_final_keyboard("plan")
    )
    await state.set_state(PlanStates.waiting_for_more_plan_items)
    await callback.answer()

@dp.callback_query(F.data == "back_to_events_list")
async def back_to_events_list(callback: types.CallbackQuery, state: FSMContext):
    """Возврат к списку мероприятий"""
    user_id = callback.from_user.id
    
    draft = get_report_draft(user_id)
    if not draft or not draft.get('events'):
        await callback.answer("❌ Нет мероприятий")
        return
    
    events = draft['events']
    await callback.message.edit_text(
        "📋 Выберите мероприятие для редактирования:",
        reply_markup=get_events_list_keyboard(events, "report")
    )
    await callback.answer()

@dp.callback_query(F.data == "back_to_plan_list")
async def back_to_plan_list(callback: types.CallbackQuery, state: FSMContext):
    """Возврат к списку пунктов плана"""
    user_id = callback.from_user.id
    
    draft = get_plan_draft(user_id)
    if not draft or not draft.get('plan_items'):
        await callback.answer("❌ Нет пунктов плана")
        return
    
    plan_items = draft['plan_items']
    await callback.message.edit_text(
        "📋 Выберите пункт плана для редактирования:",
        reply_markup=get_events_list_keyboard(plan_items, "plan")
    )
    await callback.answer()

@dp.callback_query(F.data == "finish_editing")
async def finish_editing(callback: types.CallbackQuery, state: FSMContext):
    """Завершение редактирования"""
    user_id = callback.from_user.id
    
    draft = get_report_draft(user_id)
    if not draft or not draft.get('events'):
        await callback.answer("❌ Нет мероприятий")
        return
    
    events_count = len(draft['events'])
    total_participants = sum(
        event['age_14_17'] + event['age_18_29'] + event['age_30_35'] 
        for event in draft['events']
    )
    
    await callback.message.edit_text(
        f"✅ Редактирование завершено!\n\n"
        f"📋 Мероприятий: {events_count}\n"
        f"👥 Общий охват: {total_participants} чел.\n\n"
        f"Выберите дальнейшее действие:",
        reply_markup=get_final_keyboard("report")
    )
    await state.set_state(ReportStates.waiting_for_more_events)
    await callback.answer()

@dp.callback_query(F.data == "finish_plan_editing")
async def finish_plan_editing(callback: types.CallbackQuery, state: FSMContext):
    """Завершение редактирования плана"""
    user_id = callback.from_user.id
    
    draft = get_plan_draft(user_id)
    if not draft or not draft.get('plan_items'):
        await callback.answer("❌ Нет пунктов плана")
        return
    
    plan_count = len(draft['plan_items'])
    total_participants = sum(item['participants'] for item in draft['plan_items'])
    
    await callback.message.edit_text(
        f"✅ Редактирование завершено!\n\n"
        f"📋 Пунктов плана: {plan_count}\n"
        f"👥 Общий охват: {total_participants} чел.\n\n"
        f"Выберите дальнейшее действие:",
        reply_markup=get_final_keyboard("plan")
    )
    await state.set_state(PlanStates.waiting_for_more_plan_items)
    await callback.answer()

# Обработчики сообщений для редактирования
@dp.message(ReportStates.editing_event)
async def process_edit_report_field(message: types.Message, state: FSMContext):
    """Обработка ввода данных при редактировании отчета"""
    user_id = message.from_user.id
    data = await state.get_data()
    field_name = data.get('editing_field')
    event_index = data.get('editing_event_index')
    
    if field_name is None or event_index is None:
        await message.answer("❌ Ошибка редактирования. Начните заново.")
        await state.clear()
        return
    
    draft = get_report_draft(user_id)
    if not draft or not draft.get('events') or event_index >= len(draft['events']):
        await message.answer("❌ Ошибка: мероприятие не найдено")
        await state.clear()
        return
    
    event = draft['events'][event_index]
    
    if field_name in ['age14', 'age18', 'age30']:
        # Обработка числовых полей
        try:
            value = int(message.text)
            if value < 0:
                await message.answer("❌ Количество не может быть отрицательным. Введите снова:")
                return
            
            if field_name == 'age14':
                event['age_14_17'] = value
            elif field_name == 'age18':
                event['age_18_29'] = value
            elif field_name == 'age30':
                event['age_30_35'] = value
                
        except ValueError:
            await message.answer("❌ Пожалуйста, введите корректное число:")
            return
    else:
        # Обработка текстовых полей
        if field_name == 'name':
            event['event_name'] = message.text
        elif field_name == 'place':
            event['place'] = message.text
        elif field_name == 'desc':
            event['description'] = message.text
        elif field_name == 'link':
            if not is_valid_url(message.text):
                await message.answer("❌ Это не корректная ссылка. Введите валидный URL:")
                return
            event['link'] = message.text
    
    save_report_draft(user_id, draft)
    
    # Возвращаем к меню редактирования
    events_count = len(draft['events'])
    event_text = (
        f"✅ Поле обновлено!\n\n"
        f"✏️ Редактирование мероприятия #{event_index + 1}\n\n"
        f"📝 Название: {event.get('event_name', 'Не указано')}\n"
        f"📅 Даты: {len(event.get('dates', []))} выбрано\n"
        f"📍 Место: {event.get('place', 'Не указано')}\n"
        f"⏰ Время: {event.get('time', 'Не указано')}\n"
        f"📝 Описание: {event.get('description', 'Не указано')[:50]}...\n"
        f"🔗 Ссылка: {event.get('link', 'Не указана')}\n"
        f"👥 Участники 14-17: {event.get('age_14_17', 0)}\n"
        f"👥 Участники 18-29: {event.get('age_18_29', 0)}\n"
        f"👥 Участники 30-35: {event.get('age_30_35', 0)}\n\n"
        f"Выберите что хотите изменить:"
    )
    
    await message.answer(
        event_text,
        reply_markup=get_edit_report_keyboard(event_index, events_count)
    )

@dp.message(PlanStates.editing_plan_item)
async def process_edit_plan_field(message: types.Message, state: FSMContext):
    """Обработка ввода данных при редактировании плана"""
    user_id = message.from_user.id
    data = await state.get_data()
    field_name = data.get('editing_plan_field')
    item_index = data.get('editing_plan_index')
    
    if field_name is None or item_index is None:
        await message.answer("❌ Ошибка редактирования. Начните заново.")
        await state.clear()
        return
    
    draft = get_plan_draft(user_id)
    if not draft or not draft.get('plan_items') or item_index >= len(draft['plan_items']):
        await message.answer("❌ Ошибка: пункт плана не найден")
        await state.clear()
        return
    
    item = draft['plan_items'][item_index]
    
    if field_name == 'participants':
        # Обработка числового поля
        try:
            value = int(message.text)
            if value < 0:
                await message.answer("❌ Количество не может быть отрицательным. Введите снова:")
                return
            item['participants'] = value
        except ValueError:
            await message.answer("❌ Пожалуйста, введите корректное число:")
            return
    elif field_name == 'link':
        # Проверка URL
        if not is_valid_url(message.text):
            await message.answer("❌ Это не корректная ссылка. Введите валидный URL:")
            return
        item['info_link'] = message.text
    else:
        # Обработка текстовых полей
        if field_name == 'name':
            item['event_name'] = message.text
        elif field_name == 'place':
            item['place'] = message.text
    
    save_plan_draft(user_id, draft)
    
    # Возвращаем к меню редактирования
    items_count = len(draft['plan_items'])
    item_text = (
        f"✅ Поле обновлено!\n\n"
        f"✏️ Редактирование пункта плана #{item_index + 1}\n\n"
        f"🎯 Направление: {item.get('direction', 'Не указано')[:50]}...\n"
        f"📝 Название: {item.get('event_name', 'Не указано')}\n"
        f"🏷️ Статус: {item.get('status', 'Не указан')}\n"
        f"📅 Даты: {len(item.get('dates', []))} выбрано\n"
        f"📍 Место: {item.get('place', 'Не указано')}\n"
        f"👥 Участники: {item.get('participants', 0)} чел.\n"
        f"🔗 Ссылка: {item.get('info_link', 'Не указана')}\n\n"
        f"Выберите что хотите изменить:"
    )
    
    await message.answer(
        item_text,
        reply_markup=get_edit_plan_keyboard(item_index, items_count)
    )

# Обработчики callback-ов для админской расширенной статистики
@dp.callback_query(F.data == "admin_extended_stats")
async def process_admin_extended_stats(callback: types.CallbackQuery, state: FSMContext):
    """Обработка просмотра расширенной статистики"""
    user_id = callback.from_user.id
    
    if user_id not in ADMIN_IDS:
        await callback.answer("❌ Доступ запрещен")
        return
    
    await callback.message.edit_text(
        "📊 Выберите год для просмотра статистики:",
        reply_markup=get_admin_year_keyboard()
    )
    await state.set_state(AdminStates.waiting_for_year_selection)
    await callback.answer()

@dp.callback_query(F.data.startswith("admin_year_"))
async def process_admin_year_selection(callback: types.CallbackQuery, state: FSMContext):
    """Обработка выбора года для расширенной статистики"""
    user_id = callback.from_user.id
    
    if user_id not in ADMIN_IDS:
        await callback.answer("❌ Доступ запрещен")
        return
    
    year_str = callback.data.replace("admin_year_", "")
    
    if year_str == "all":
        # Статистика за все годы
        current_year = datetime.now().year
        await state.update_data(selected_year=None)
    else:
        try:
            year = int(year_str)
            await state.update_data(selected_year=year)
        except ValueError:
            await callback.answer("❌ Ошибка выбора года")
            return
    
    current_year = datetime.now().year if year_str == "all" else int(year_str)
    
    await callback.message.edit_text(
        f"📅 Выберите месяц для просмотра статистики за {current_year} год:",
        reply_markup=get_admin_month_keyboard(current_year)
    )
    await state.set_state(AdminStates.waiting_for_month_selection)
    await callback.answer()

@dp.callback_query(F.data.startswith("admin_month_"))
async def process_admin_month_selection(callback: types.CallbackQuery, state: FSMContext):
    """Обработка выбора месяца для расширенной статистики"""
    user_id = callback.from_user.id
    
    if user_id not in ADMIN_IDS:
        await callback.answer("❌ Доступ запрещен")
        return
    
    month_year_str = callback.data.replace("admin_month_", "")
    parts = month_year_str.split('_')
    
    if len(parts) != 2:
        await callback.answer("❌ Ошибка формата")
        return
    
    month = parts[0]
    year_str = parts[1]
    
    try:
        year = int(year_str) if year_str != "all" else None
    except ValueError:
        await callback.answer("❌ Ошибка года")
        return
    
    if month == "all":
        # Статистика за весь год
        if year is None:
            # Статистика за все годы
            await show_extended_stats_for_all_years(callback)
        else:
            # Статистика за конкретный год
            await show_extended_stats_for_year(callback, year)
    else:
        # Статистика за конкретный месяц
        if year is None:
            year = datetime.now().year
        
        await show_extended_stats_for_month(callback, year, month)
    
    await callback.answer()

async def show_extended_stats_for_all_years(callback: types.CallbackQuery):
    """Показывает статистику за все годы"""
    conn = sqlite3.connect('reports_plans.db')
    cursor = conn.cursor()
    
    try:
        # Получаем список всех лет из базы данных
        cursor.execute('''
            SELECT DISTINCT SUBSTR(month_year, -4) as year FROM reports
            UNION
            SELECT DISTINCT SUBSTR(month_year, -4) as year FROM plans
            ORDER BY year DESC
        ''')
        
        years = [row[0] for row in cursor.fetchall()]
        
        if not years:
            await callback.message.edit_text(
                "📊 Статистика отсутствует за все годы",
                reply_markup=get_admin_main_keyboard()
            )
            return
        
        stats_text = "📊 Статистика по клубам за все годы:\n\n"
        
        for club_name in CLUBS_DATA.keys():
            stats_text += f"🏢 {club_name}:\n"
            
            for year in years:
                # Отчеты за год
                cursor.execute('''
                    SELECT COUNT(*) FROM reports 
                    WHERE club_name = ? AND month_year LIKE ?
                ''', (club_name, f'%{year}%'))
                reports_count = cursor.fetchone()[0]
                
                # Планы за год
                cursor.execute('''
                    SELECT COUNT(*) FROM plans 
                    WHERE club_name = ? AND month_year LIKE ?
                ''', (club_name, f'%{year}%'))
                plans_count = cursor.fetchone()[0]
                
                if reports_count > 0 or plans_count > 0:
                    stats_text += f"  📅 {year}: "
                    stats_text += f"📊 {reports_count} отч." if reports_count > 0 else "📊 0 отч."
                    stats_text += f", 📋 {plans_count} пл." if plans_count > 0 else ", 📋 0 пл."
                    stats_text += "\n"
            
            # Добавляем разделитель между клубами
            stats_text += "\n"
        
        await callback.message.edit_text(
            stats_text,
            reply_markup=get_admin_main_keyboard()
        )
        
    except Exception as e:
        logger.error(f"Error showing stats for all years: {e}")
        await callback.message.edit_text(
            f"❌ Ошибка при получении статистики: {str(e)}",
            reply_markup=get_admin_main_keyboard()
        )
    finally:
        conn.close()

async def show_extended_stats_for_year(callback: types.CallbackQuery, year: int):
    """Показывает статистику за конкретный год"""
    statistics = get_clubs_statistics(year)
    
    months = [
        "январь", "февраль", "март", "апрель", "май", "июнь",
        "июль", "август", "сентябрь", "октябрь", "ноябрь", "декабрь"
    ]
    
    stats_text = f"📊 Статистика по клубам за {year} год:\n\n"
    
    for club_name, club_stats in statistics.items():
        reports_count = club_stats.get('reports_count', 0)
        plans_count = club_stats.get('plans_count', 0)
        
        stats_text += f"🏢 {club_name}:\n"
        stats_text += f"  📊 Отчетов: {reports_count} ({reports_count * 100 // 12}% от годового плана)\n"
        stats_text += f"  📋 Планов: {plans_count} ({plans_count * 100 // 12}% от годового плана)\n"
        stats_text += f"  📅 По месяцам:\n"
        
        for month in months:
            month_year = f"{month} {year}"
            
            conn = sqlite3.connect('reports_plans.db')
            cursor = conn.cursor()
            
            cursor.execute('''
                SELECT COUNT(*) FROM reports WHERE club_name = ? AND month_year = ?
            ''', (club_name, month_year))
            has_report = cursor.fetchone()[0] > 0
            
            cursor.execute('''
                SELECT COUNT(*) FROM plans WHERE club_name = ? AND month_year = ?
            ''', (club_name, month_year))
            has_plan = cursor.fetchone()[0] > 0
            
            conn.close()
            
            report_symbol = "✅" if has_report else "❌"
            plan_symbol = "✅" if has_plan else "❌"
            
            stats_text += f"    {month[:3]}: {report_symbol}отч {plan_symbol}пл\n"
        
        stats_text += "\n"
    
    await callback.message.edit_text(
        stats_text,
        reply_markup=get_admin_main_keyboard()
    )

async def show_extended_stats_for_month(callback: types.CallbackQuery, year: int, month: str):
    """Показывает статистику за конкретный месяц"""
    statistics = get_clubs_statistics(year, month)
    
    stats_text = f"📊 Статистика по клубам за {month} {year}:\n\n"
    stats_text += "✅ - документ есть\n❌ - документ отсутствует\n\n"
    
    for club_name, club_stats in statistics.items():
        has_report = club_stats.get('report', False)
        has_plan = club_stats.get('plan', False)
        
        report_symbol = "✅" if has_report else "❌"
        plan_symbol = "✅" if has_plan else "❌"
        
        stats_text += f"{report_symbol}отч {plan_symbol}пл - {club_name}\n"
    
    # Добавляем итоговую статистику
    total_clubs = len(statistics)
    clubs_with_report = sum(1 for stats in statistics.values() if stats.get('report', False))
    clubs_with_plan = sum(1 for stats in statistics.values() if stats.get('plan', False))
    clubs_with_both = sum(1 for stats in statistics.values() if stats.get('report', False) and stats.get('plan', False))
    
    stats_text += f"\n📈 Итоги:\n"
    stats_text += f"🏢 Всего клубов: {total_clubs}\n"
    stats_text += f"📊 С отчетом: {clubs_with_report} ({clubs_with_report * 100 // total_clubs if total_clubs > 0 else 0}%)\n"
    stats_text += f"📋 С планом: {clubs_with_plan} ({clubs_with_plan * 100 // total_clubs if total_clubs > 0 else 0}%)\n"
    stats_text += f"✅ С обоими документами: {clubs_with_both} ({clubs_with_both * 100 // total_clubs if total_clubs > 0 else 0}%)\n"
    stats_text += f"❌ Без документов: {total_clubs - clubs_with_both} ({(total_clubs - clubs_with_both) * 100 // total_clubs if total_clubs > 0 else 0}%)"
    
    await callback.message.edit_text(
        stats_text,
        reply_markup=get_admin_main_keyboard()
    )

# Остальные обработчики callback-ов для админки
@dp.callback_query(F.data.startswith("admin_club_"))
async def process_admin_club_selection(callback: types.CallbackQuery, state: FSMContext):
    """Обработка выбора клуба в админке"""
    user_id = callback.from_user.id
    club_name = callback.data.replace("admin_club_", "")
    
    if user_id not in ADMIN_IDS:
        await callback.answer("❌ Доступ запрещен")
        return
    
    await state.update_data(selected_club=club_name)
    
    await callback.message.edit_text(
        f"🏢 Выбран клуб: {club_name}\n\n"
        f"📋 Выберите тип документов для просмотра:",
        reply_markup=get_admin_document_type_keyboard()
    )
    await state.set_state(AdminStates.waiting_for_document_type)
    await callback.answer()

@dp.callback_query(F.data.startswith("admin_doc_"))
async def process_admin_document_type(callback: types.CallbackQuery, state: FSMContext):
    """Обработка выбора типа документа в админке"""
    user_id = callback.from_user.id
    doc_type = callback.data.replace("admin_doc_", "")
    
    if user_id not in ADMIN_IDS:
        await callback.answer("❌ Доступ запрещен")
        return
    
    data = await state.get_data()
    club_name = data.get('selected_club')
    
    if doc_type == 'reports':
        documents = get_club_reports(club_name)
        doc_type_name = "отчеты"
    else:
        documents = get_club_plans(club_name)
        doc_type_name = "планы"
    
    if not documents:
        missing_months = get_missing_months(club_name, doc_type, datetime.now().year)
        missing_text = "\n".join([f"❌ {month}" for month in missing_months])
        
        await callback.message.edit_text(
            f"🏢 Клуб: {club_name}\n"
            f"📋 {doc_type_name.capitalize()}: не найдено\n\n"
            f"📅 Отсутствуют документы за:\n{missing_text}"
        )
    else:
        # Отправляем все документы
        for doc in documents:
            month_year, file_path, items_count, participants_count, created_at = doc
            
            file_path = resolve_file_path(file_path) or file_path
            if os.path.exists(file_path):
                doc_file = FSInputFile(file_path)
                await callback.message.answer_document(
                    doc_file,
                    caption=(
                        f"🏢 Клуб: {club_name}\n"
                        f"📅 Период: {month_year}\n"
                        f"📊 Количество мероприятий: {items_count}\n"
                        f"👥 Участников: {participants_count}\n"
                        f"🕐 Создан: {created_at}"
                    )
                )
            else:
                await callback.message.answer(
                    f"❌ Файл не найден: {month_year}"
                )
        
        # Показываем статистику по отсутствующим месяцам
        missing_months = get_missing_months(club_name, doc_type, datetime.now().year)
        if missing_months:
            missing_text = "\n".join([f"❌ {month}" for month in missing_months])
            await callback.message.answer(
                f"📅 Отсутствуют {doc_type_name} за:\n{missing_text}"
            )
        else:
            await callback.message.answer(
                f"✅ Все {doc_type_name} за текущий год присутствуют!"
            )
    
    await state.set_state(AdminStates.waiting_for_admin_action)
    await callback.message.answer(
        "Выберите дальнейшее действие:",
        reply_markup=get_admin_main_keyboard()
    )
    await callback.answer()

@dp.callback_query(F.data == "admin_stats")
async def process_admin_stats(callback: types.CallbackQuery, state: FSMContext):
    """Обработка просмотра статистики"""
    user_id = callback.from_user.id
    
    if user_id not in ADMIN_IDS:
        await callback.answer("❌ Доступ запрещен")
        return
    
    stats = get_all_clubs_with_stats()
    
    if not stats:
        await callback.message.edit_text("📊 Статистика отсутствует")
        return
    
    stats_text = "📊 Статистика по клубам:\n\n"
    for club_name, reports_count, plans_count in stats:
        stats_text += f"🏢 {club_name}:\n"
        stats_text += f"   📊 Отчетов: {reports_count}\n"
        stats_text += f"   📋 Планов: {plans_count}\n\n"
    
    await callback.message.edit_text(
        stats_text,
        reply_markup=get_admin_main_keyboard()
    )
    await callback.answer()

@dp.callback_query(F.data == "admin_view_docs")
async def process_admin_view_docs(callback: types.CallbackQuery, state: FSMContext):
    """Обработка просмотра документов"""
    user_id = callback.from_user.id
    
    if user_id not in ADMIN_IDS:
        await callback.answer("❌ Доступ запрещен")
        return
    
    await callback.message.edit_text(
        "🏢 Выберите клуб для просмотра документов:",
        reply_markup=get_admin_clubs_keyboard()
    )
    await state.set_state(AdminStates.waiting_for_club_selection)
    await callback.answer()

@dp.callback_query(F.data.startswith("direction_"))
async def process_direction_selection(callback: types.CallbackQuery, state: FSMContext):
    """Обработка выбора направления"""
    user_id = callback.from_user.id
    direction_num = callback.data.replace("direction_", "")
    
    direction_key = f"direction_{direction_num}"
    if direction_key in YOUTH_POLICY_DIRECTIONS:
        direction_text = YOUTH_POLICY_DIRECTIONS[direction_key]
        
        # Сохраняем в черновик плана
        draft = get_plan_draft(user_id) or {}
        plan_items = draft.get('plan_items', [])
        
        plan_items.append({
            'direction': direction_text,
            'event_name': '',
            'status': '',
            'dates': [],
            'place': '',
            'participants': 0,
            'info_link': ''
        })
        
        draft['plan_items'] = plan_items
        draft['current_plan_index'] = len(plan_items) - 1
        save_plan_draft(user_id, draft)
        
        await callback.message.edit_text(
            f"✅ Выбрано направление {direction_num}\n\n"
            f"📋 Теперь введите название мероприятия:"
        )
        await state.set_state(PlanStates.waiting_for_event_name)
    
    await callback.answer()

@dp.callback_query(F.data.startswith("status_"))
async def process_status_selection(callback: types.CallbackQuery, state: FSMContext):
    """Обработка выбора статуса мероприятия"""
    user_id = callback.from_user.id
    status = callback.data.replace("status_", "")
    
    draft = get_plan_draft(user_id)
    if draft and draft.get('plan_items'):
        current_plan_index = draft.get('current_plan_index', 0)
        draft['plan_items'][current_plan_index]['status'] = status
        save_plan_draft(user_id, draft)
        
        month_num, year = parse_month_year(draft['month_year'])
        calendar_markup = get_month_calendar(year, month_num, [])
        
        await callback.message.edit_text(
            f"✅ Статус мероприятия: {status}\n\n"
            f"📅 Теперь выберите даты проведения мероприятия:",
            reply_markup=calendar_markup
        )
        await state.set_state(PlanStates.waiting_for_event_dates)
    
    await callback.answer()

@dp.callback_query(F.data.startswith("date_"))
async def process_date_selection(callback: types.CallbackQuery, state: FSMContext):
    """Обработка выбора даты"""
    user_id = callback.from_user.id
    date_str = callback.data.replace("date_", "")
    
    current_state = await state.get_state()
    
    if current_state == ReportStates.waiting_for_event_dates.state:
        draft = get_report_draft(user_id)
        if draft:
            selected_dates = draft.get('selected_dates', [])
            
            if date_str in selected_dates:
                selected_dates.remove(date_str)
            else:
                selected_dates.append(date_str)
            
            draft['selected_dates'] = selected_dates
            save_report_draft(user_id, draft)
            
            month_num, year = parse_month_year(draft['month_year'])
            calendar_markup = get_month_calendar(year, month_num, selected_dates)
            
            selected_dates_text = "\n".join([f"• {format_date_for_display(d)}" for d in sorted(selected_dates)])
            
            await callback.message.edit_text(
                f"📅 Выберите даты проведения мероприятия:\n\n"
                f"📋 Выбранные даты:\n{selected_dates_text if selected_dates_text else 'Нет выбранных дат'}\n\n"
                f"🗓️ Календарь {draft['month_year']}:",
                reply_markup=calendar_markup
            )
    
    elif current_state == PlanStates.waiting_for_event_dates.state:
        draft = get_plan_draft(user_id)
        if draft and draft.get('plan_items'):
            current_plan_index = draft.get('current_plan_index', 0)
            current_dates = draft['plan_items'][current_plan_index].get('dates', [])
            
            if date_str in current_dates:
                current_dates.remove(date_str)
            else:
                current_dates.append(date_str)
            
            draft['plan_items'][current_plan_index]['dates'] = current_dates
            save_plan_draft(user_id, draft)
            
            month_num, year = parse_month_year(draft['month_year'])
            calendar_markup = get_month_calendar(year, month_num, current_dates)
            
            selected_dates_text = "\n".join([f"• {format_date_for_display(d)}" for d in sorted(current_dates)])
            
            await callback.message.edit_text(
                f"📅 Выберите даты проведения мероприятия:\n\n"
                f"📋 Выбранные даты:\n{selected_dates_text if selected_dates_text else 'Нет выбранных дат'}\n\n"
                f"🗓️ Календарь {draft['month_year']}:",
                reply_markup=calendar_markup
            )
    
    await callback.answer()

@dp.callback_query(F.data == "clear_all")
async def clear_all_dates(callback: types.CallbackQuery, state: FSMContext):
    """Очистка всех выбранных дат"""
    user_id = callback.from_user.id
    
    current_state = await state.get_state()
    
    if current_state == ReportStates.waiting_for_event_dates.state:
        draft = get_report_draft(user_id)
        if draft:
            draft['selected_dates'] = []
            save_report_draft(user_id, draft)
            
            month_num, year = parse_month_year(draft['month_year'])
            calendar_markup = get_month_calendar(year, month_num, [])
            
            await callback.message.edit_text(
                f"📅 Выберите даты проведения мероприятия:\n\n"
                f"📋 Выбранные даты:\nНет выбранных дат\n\n"
                f"🗓️ Календарь {draft['month_year']}:",
                reply_markup=calendar_markup
            )
    elif current_state == PlanStates.waiting_for_event_dates.state:
        draft = get_plan_draft(user_id)
        if draft and draft.get('plan_items'):
            current_plan_index = draft.get('current_plan_index', 0)
            draft['plan_items'][current_plan_index]['dates'] = []
            save_plan_draft(user_id, draft)
            
            month_num, year = parse_month_year(draft['month_year'])
            calendar_markup = get_month_calendar(year, month_num, [])
            
            await callback.message.edit_text(
                f"📅 Выберите даты проведения мероприятия:\n\n"
                f"📋 Выбранные даты:\nНет выбранных дат\n\n"
                f"🗓️ Календарь {draft['month_year']}:",
                reply_markup=calendar_markup
            )
    
    await callback.answer("Все даты очищены")

@dp.callback_query(F.data == "finish_dates")
async def finish_date_selection(callback: types.CallbackQuery, state: FSMContext):
    """Завершение выбора дат"""
    user_id = callback.from_user.id
    
    current_state = await state.get_state()
    
    if current_state == ReportStates.waiting_for_event_dates.state:
        draft = get_report_draft(user_id)
        if not draft or not draft.get('selected_dates'):
            await callback.answer("❌ Выберите хотя бы одну дату")
            return
        
        events = draft.get('events', [])
        current_event_index = draft.get('current_event_index', 0)
        
        if current_event_index < len(events):
            events[current_event_index]['dates'] = draft['selected_dates'].copy()
        else:
            # Создаем новое мероприятие если нужно
            events.append({
                'event_name': '',
                'dates': draft['selected_dates'].copy(),
                'place': '',
                'time': '',
                'description': '',
                'link': '',
                'age_14_17': 0,
                'age_18_29': 0,
                'age_30_35': 0
            })
            draft['current_event_index'] = len(events) - 1
        
        draft['events'] = events
        save_report_draft(user_id, draft)
        
        selected_dates_count = len(draft['selected_dates'])
        
        await callback.message.edit_text(
            f"✅ Выбрано дат: {selected_dates_count}\n\n"
            f"📍 Теперь введите место проведения мероприятия:",
            reply_markup=get_place_keyboard()  # Добавьте эту строку
        )
        
        await state.set_state(ReportStates.waiting_for_event_place)
    
    elif current_state == PlanStates.waiting_for_event_dates.state:
        draft = get_plan_draft(user_id)
        if not draft or not draft.get('plan_items'):
            await callback.answer("❌ Выберите хотя бы одну дату")
            return
        
        current_plan_index = draft.get('current_plan_index', 0)
        current_dates = draft['plan_items'][current_plan_index].get('dates', [])
        
        if not current_dates:
            await callback.answer("❌ Выберите хотя бы одну дату")
            return
        
        selected_dates_count = len(current_dates)
        
        await callback.message.edit_text(
            f"✅ Выбрано дат: {selected_dates_count}\n\n"
            f"📍 Теперь введите место проведения мероприятия:",
            reply_markup=get_place_keyboard()  # Добавьте эту строку
        )
        
        await state.set_state(PlanStates.waiting_for_event_place)
    
    await callback.answer()

@dp.callback_query(F.data == "add_event")
async def add_event_callback(callback: types.CallbackQuery, state: FSMContext):
    """Обработка добавления нового мероприятия"""
    user_id = callback.from_user.id
    
    draft = get_report_draft(user_id)
    if not draft:
        await callback.answer("❌ Нет активного черновика")
        return
    
    events = draft.get('events', [])
    if len(events) >= 15:
        await callback.answer("❌ Достигнут лимит в 15 мероприятий")
        return
    
    draft['current_event_index'] = len(events)
    draft['selected_dates'] = []
    save_report_draft(user_id, draft)
    
    await callback.message.edit_text(
        f"✅ Добавляем мероприятие #{len(events) + 1}\n\n"
        f"Введите наименование мероприятия:"
    )
    await state.set_state(ReportStates.waiting_for_event_name)
    await callback.answer()

@dp.callback_query(F.data == "add_plan_item")
async def add_plan_item_callback(callback: types.CallbackQuery, state: FSMContext):
    """Обработка добавления нового пункта плана"""
    user_id = callback.from_user.id
    
    draft = get_plan_draft(user_id)
    if not draft:
        await callback.answer("❌ Нет активного черновика")
        return
    
    plan_items = draft.get('plan_items', [])
    if len(plan_items) >= 15:
        await callback.answer("❌ Достигнут лимит в 15 пунктов плана")
        return
    
    draft['current_plan_index'] = len(plan_items)
    save_plan_draft(user_id, draft)
    
    await callback.message.edit_text(
        f"✅ Добавляем пункт плана #{len(plan_items) + 1}\n\n"
        f"Выберите направление реализации государственной молодежной политики:",
        reply_markup=get_directions_keyboard()
    )
    await state.set_state(PlanStates.waiting_for_direction_selection)
    await callback.answer()

@dp.callback_query(F.data == "generate_report")
async def generate_report_callback(callback: types.CallbackQuery, state: FSMContext):
    """Обработка формирования отчета"""
    user_id = callback.from_user.id
    
    draft = get_report_draft(user_id)
    if not draft or not draft.get('events'):
        await callback.answer("❌ Нет данных для формирования отчета")
        return
    
    # Проверяем наличие всех необходимых полей
    required_fields = ['club_name', 'position', 'responsible', 'month_year']
    for field in required_fields:
        if field not in draft:
            await callback.answer(f"❌ Отсутствуют необходимые данные: {field}")
            return
    
    await callback.message.edit_text("📊 Формирую отчет...")
    
    try:
        template_path = draft.get('template_path', 'Шаблон отчета.docx')
        user_data = {
            'club_name': draft['club_name'],
            'position': draft['position'],
            'responsible': draft['responsible'],
            'month_year': draft['month_year']
        }
        
        filled_doc_path = fill_report_template(template_path, user_data, draft['events'])
        
        doc_file = FSInputFile(filled_doc_path)
        
        events_count = len(draft['events'])
        total_participants = sum(
            event['age_14_17'] + event['age_18_29'] + event['age_30_35'] 
            for event in draft['events']
        )
        
        filename = os.path.basename(filled_doc_path)
        
        # Сохраняем отчет в базу данных (вместе с мероприятиями для редактирования)
        save_report_to_db(
            user_id, 
            draft['club_name'], 
            draft['month_year'],
            filled_doc_path,
            events_count,
            total_participants,
            draft['events'],
        )
        
        # Удаляем черновик
        delete_report_draft(user_id)
        
        await callback.message.answer_document(
            doc_file,
            caption=(
                f"✅ {filename}\n"
                f"📋 Количество мероприятий: {events_count}\n"
                f"👥 Общий охват: {total_participants} чел.\n"
                f"🏢 Клуб: {draft['club_name']}\n"
                f"📝 Ответственный: {draft['position']} {draft['responsible']}"
            )
        )

        # ── Подтверждения охватов ────────────────────────────────────────
        confirmation_paths: list = []
        if os.path.exists(COVERAGE_TEMPLATE_PATH):
            await callback.message.answer("📋 Формирую подтверждения охватов...")
            try:
                confirmation_paths = generate_coverage_confirmations(
                    events=draft['events'],
                    club_name=draft['club_name'],
                    month_year=draft['month_year'],
                    output_dir=TEMP_DIR,
                    template_path=COVERAGE_TEMPLATE_PATH,
                )
                # Отправляем специалисту по одному файлу
                if confirmation_paths:
                    for conf_path in confirmation_paths:
                        await callback.message.answer_document(
                            FSInputFile(conf_path),
                            caption=f"📋 {os.path.basename(conf_path).replace('.docx', '')}",
                        )
                        await asyncio.sleep(0.3)
            except Exception as cov_exc:
                logger.error(f"[COVERAGE] {cov_exc}", exc_info=True)
                await callback.message.answer(
                    f"⚠️ Ошибка при формировании подтверждений: {cov_exc}"
                )
        else:
            logger.warning(f"[COVERAGE] Template not found: {COVERAGE_TEMPLATE_PATH}")

        # ── ZIP для администратора: отчёт + все подтверждения охватов ───
        try:
            safe_club  = "".join(c for c in draft['club_name'] if c.isalnum() or c in " _-").strip()
            safe_month = draft['month_year'].replace(" ", "_")
            zip_name   = f"Отчет_{safe_club}_{safe_month}.zip"
            zip_path   = os.path.join(TEMP_DIR, zip_name)

            zip_buf = io.BytesIO()
            with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
                # Сам отчёт в корень архива
                zf.write(filled_doc_path, os.path.basename(filled_doc_path))
                # Подтверждения охватов в папку «Охваты»
                for conf_path in confirmation_paths:
                    zf.write(conf_path, os.path.join("Охваты", os.path.basename(conf_path)))

            zip_bytes = zip_buf.getvalue()
            with open(zip_path, "wb") as f:
                f.write(zip_bytes)

            # Отправляем ZIP каждому администратору
            caption_zip = (
                f"📦 Новый отчёт сформирован\n\n"
                f"🏢 Клуб: {draft['club_name']}\n"
                f"📅 Период: {draft['month_year']}\n"
                f"👤 Пользователь: {user_id}\n"
                f"📋 Мероприятий: {events_count}\n"
                f"👥 Участников: {total_participants}\n"
                f"📎 Файлов в архиве: {1 + len(confirmation_paths)}\n"
                f"🕐 {datetime.now().strftime('%d.%m.%Y %H:%M')}"
            )
            for admin_id in ADMIN_IDS:
                try:
                    await bot.send_document(
                        admin_id,
                        types.BufferedInputFile(zip_bytes, filename=zip_name),
                        caption=caption_zip,
                    )
                except Exception as admin_exc:
                    logger.error(f"[ZIP] Cannot send to admin {admin_id}: {admin_exc}")

        except Exception as zip_exc:
            logger.error(f"[ZIP] Error building zip: {zip_exc}", exc_info=True)
        
        await state.clear()
        await callback.answer("Отчет успешно сформирован!")
        
    except Exception as e:
        logger.error(f"Error generating report: {e}")
        await callback.message.answer(f"❌ Произошла ошибка: {str(e)}\nПопробуйте снова: /new_report")
        await state.clear()
        await callback.answer("Ошибка при формировании отчета")

@dp.callback_query(F.data == "generate_plan")
async def generate_plan_callback(callback: types.CallbackQuery, state: FSMContext):
    """Обработка формирования плана"""
    user_id = callback.from_user.id
    
    draft = get_plan_draft(user_id)
    if not draft or not draft.get('plan_items'):
        await callback.answer("❌ Нет данных для формирования плана")
        return
    
    # Проверяем наличие всех необходимых полей
    required_fields = ['club_name', 'position', 'responsible', 'month_year']
    for field in required_fields:
        if field not in draft:
            await callback.answer(f"❌ Отсутствуют необходимые данные: {field}")
            return
    
    await callback.message.edit_text("📋 Формирую план...")
    
    try:
        template_path = "Шаблон плана.docx"
        
        if not os.path.exists(template_path):
            await callback.message.answer("❌ Шаблон плана не найден.")
            return
        
        user_data = {
            'club_name': draft['club_name'],
            'position': draft['position'],
            'responsible': draft['responsible'],
            'month_year': draft['month_year']
        }
        
        filled_doc_path = fill_plan_template(template_path, user_data, draft['plan_items'])
        
        doc_file = FSInputFile(filled_doc_path)
        
        plan_count = len(draft['plan_items'])
        total_participants = sum(item['participants'] for item in draft['plan_items'])
        
        filename = os.path.basename(filled_doc_path)
        
        # Сохраняем план в базу данных (вместе с пунктами для редактирования)
        save_plan_to_db(
            user_id, 
            draft['club_name'], 
            draft['month_year'],
            filled_doc_path,
            plan_count,
            total_participants,
            draft['plan_items'],
        )
        
        # Удаляем черновик
        delete_plan_draft(user_id)
        
        await callback.message.answer_document(
            doc_file,
            caption=(
                f"✅ {filename}\n"
                f"📋 Количество пунктов плана: {plan_count}\n"
                f"👥 Общий охват: {total_participants} чел.\n"
                f"🏢 Клуб: {draft['club_name']}\n"
                f"📝 Ответственный: {draft['position']} {draft['responsible']}"
            )
        )
        
        # Отправляем документ администратору
        await send_document_to_admin(
            user_id,
            filled_doc_path,
            "План",
            draft['club_name'],
            draft['month_year'],
            plan_count,
            total_participants
        )
        
        await state.clear()
        await callback.answer("План успешно сформирован!")
        
    except Exception as e:
        logger.error(f"Error generating plan: {e}")
        await callback.message.answer(f"❌ Произошла ошибка: {str(e)}\nПопробуйте снова: /new_plan")
        await state.clear()
        await callback.answer("Ошибка при формировании плана")

@dp.callback_query(F.data == "ignore")
async def ignore_callback(callback: types.CallbackQuery):
    """Игнорируем пустые кнопки"""
    await callback.answer()

# Обработчики команд для квартальных отчетов
@dp.message(Command("new_quarterly"))
async def cmd_new_quarterly(message: types.Message, state: FSMContext):
    """Начало формирования квартального отчета"""
    user_id = message.from_user.id
    
    template_path = "Шаблон квартального отчета.docx"
    if not os.path.exists(template_path):
        await message.answer("❌ Шаблон квартального отчета не найден.")
        return
    
    await message.answer(
        "📊 Формирование квартального отчета\n\n"
        "📅 Выберите квартал:",
        reply_markup=get_quarters_keyboard()
    )
    await state.set_state(QuarterlyStates.waiting_for_quarter_selection)

@dp.callback_query(F.data.startswith("quarter_"))
async def process_quarter_selection(callback: types.CallbackQuery, state: FSMContext):
    """Обработка выбора квартала"""
    user_id = callback.from_user.id
    quarter = int(callback.data.replace("quarter_", ""))
    
    await state.update_data(quarter=quarter)
    
    current_year = datetime.now().year
    await callback.message.edit_text(
        f"✅ Выбран {quarter} квартал\n\n"
        f"📅 Введите год (например: {current_year}):"
    )
    await state.set_state(QuarterlyStates.waiting_for_quarter_year)
    await callback.answer()

@dp.message(QuarterlyStates.waiting_for_quarter_year)
async def process_quarter_year(message: types.Message, state: FSMContext):
    """Обработка года для квартального отчета"""
    user_id = message.from_user.id
    
    try:
        year = int(message.text)
        if year < 2020 or year > 2030:
            await message.answer("❌ Введите корректный год (2020-2030):")
            return
        
        await state.update_data(year=year)
        
        await message.answer(
            f"✅ Год {year} сохранен!\n\n"
            f"🏢 Теперь выберите молодежный клуб:",
            reply_markup=get_clubs_keyboard()
        )
        await state.set_state(QuarterlyStates.waiting_for_quarter_club_selection)
        
    except ValueError:
        await message.answer("❌ Пожалуйста, введите корректный год:")

# Основные обработчики сообщений
@dp.message(Command("start"))
async def cmd_start(message: types.Message, state: FSMContext):
    """Обработчик команды /start"""
    await state.clear()
    is_admin = message.from_user.id in ADMIN_IDS

    builder = InlineKeyboardBuilder()
    builder.add(
        InlineKeyboardButton(text="📊 Заполнить отчет", callback_data="start_new_report"),
        InlineKeyboardButton(text="📋 Заполнить план", callback_data="start_new_plan"),
        InlineKeyboardButton(text="📈 Квартальный отчет", callback_data="start_new_quarterly"),
        InlineKeyboardButton(text="🎨 Творческий отчет", callback_data="start_creative_report"),
    )
    if is_admin:
        builder.add(InlineKeyboardButton(text="⚙️ Панель администратора", callback_data="start_admin"))
    builder.adjust(1)

    greeting = f"👋 Привет, {message.from_user.first_name}!\n\n"
    text = (
        greeting +
        "Этот бот поможет сформировать документы молодежного клуба.\n\n"
        "Выберите действие:"
    )
    await message.answer(text, reply_markup=builder.as_markup())

@dp.callback_query(F.data == "start_new_report")
async def start_new_report_from_menu(callback: types.CallbackQuery, state: FSMContext):
    user_id = callback.from_user.id
    template_path = "Шаблон отчета.docx"
    if not os.path.exists(template_path):
        await callback.answer("❌ Шаблон отчета не найден.", show_alert=True)
        return
    draft = {'template_path': template_path, 'events': [], 'current_event_index': 0, 'selected_dates': []}
    save_report_draft(user_id, draft)
    current_year = datetime.now().year
    await callback.message.edit_text(
        "📊 Заполнение отчета о работе молодежного клуба\n\n"
        "Для отмены в любой момент нажмите /cancel\n\n"
        "1️⃣ Выберите месяц отчета:",
        reply_markup=get_months_keyboard(current_year)
    )
    await state.set_state(ReportStates.waiting_for_month)
    await callback.answer()

@dp.callback_query(F.data == "start_new_plan")
async def start_new_plan_from_menu(callback: types.CallbackQuery, state: FSMContext):
    user_id = callback.from_user.id
    template_path = "Шаблон плана.docx"
    if not os.path.exists(template_path):
        await callback.answer("❌ Шаблон плана не найден.", show_alert=True)
        return
    draft = {'plan_items': [], 'current_plan_index': 0}
    save_plan_draft(user_id, draft)
    current_year = datetime.now().year
    await callback.message.edit_text(
        "📋 Заполнение плана работы молодежного клуба\n\n"
        "Для отмены в любой момент нажмите /cancel\n\n"
        "1️⃣ Выберите месяц плана:",
        reply_markup=get_months_keyboard(current_year)
    )
    await state.set_state(PlanStates.waiting_for_plan_month)
    await callback.answer()

@dp.callback_query(F.data == "start_new_quarterly")
async def start_new_quarterly_from_menu(callback: types.CallbackQuery, state: FSMContext):
    template_path = "Шаблон квартального отчета.docx"
    if not os.path.exists(template_path):
        await callback.answer("❌ Шаблон квартального отчета не найден.", show_alert=True)
        return
    await callback.message.edit_text(
        "📈 Формирование квартального отчета\n\n"
        "📅 Выберите квартал:",
        reply_markup=get_quarters_keyboard()
    )
    await state.set_state(QuarterlyStates.waiting_for_quarter_selection)
    await callback.answer()

@dp.callback_query(F.data == "start_admin")
async def start_admin_from_menu(callback: types.CallbackQuery, state: FSMContext):
    if callback.from_user.id not in ADMIN_IDS:
        await callback.answer("❌ Доступ запрещен", show_alert=True)
        return
    await callback.message.edit_text(
        "⚙️ Панель администратора\n\nВыберите действие:",
        reply_markup=get_admin_main_keyboard()
    )
    await state.set_state(AdminStates.waiting_for_admin_action)
    await callback.answer()

@dp.message(Command("admin"))
async def cmd_admin(message: types.Message, state: FSMContext):
    """Обработчик команды админки"""
    if message.from_user.id not in ADMIN_IDS:
        await message.answer("❌ Доступ запрещен")
        return
    
    await message.answer(
        "👋 Добро пожаловать в админ-панель!\n\n"
        "Выберите действие:",
        reply_markup=get_admin_main_keyboard()
    )
    await state.set_state(AdminStates.waiting_for_admin_action)

@dp.message(Command("cancel"))
async def cmd_cancel(message: types.Message, state: FSMContext):
    """Обработчик команды отмены"""
    current_state = await state.get_state()
    if current_state is None:
        await message.answer(
            "ℹ️ Нет активных операций.\n\n"
            "Используйте /start чтобы начать."
        )
        return

    user_id = message.from_user.id
    delete_report_draft(user_id)
    delete_plan_draft(user_id)

    await state.clear()
    await message.answer(
        "❌ Заполнение отменено. Черновики удалены.\n\n"
        "Используйте /start чтобы начать снова."
    )


# ── Вспомогательная функция форматирования статуса ────────────

def _format_status_text(club_name: str, month_year: str) -> str:
    """Строит текст статуса документов для club_name за month_year."""
    status = get_status_for_club(club_name, month_year)

    # ── Отчёт ─────────────────────────────────────────────────
    if status["report"]:
        events_count, participants, created_at = status["report"]
        date_str = created_at[:10] if created_at else "—"
        report_line = (
            f"✅ Сдан\n"
            f"     📋 Мероприятий: {events_count} | 👥 Участников: {participants}\n"
            f"     🕐 {date_str}"
        )
    else:
        report_line = "❌ Не сдан"

    # ── План ───────────────────────────────────────────────────
    if status["plan"]:
        items_count, participants, created_at = status["plan"]
        date_str = created_at[:10] if created_at else "—"
        plan_line = (
            f"✅ Сдан\n"
            f"     📋 Пунктов: {items_count} | 👥 Участников: {participants}\n"
            f"     🕐 {date_str}"
        )
    else:
        plan_line = "❌ Не сдан"

    # ── Творческие отчёты ──────────────────────────────────────
    if status["creative"]:
        creative_lines = []
        for label, created_at in status["creative"]:
            date_str = created_at[:10] if created_at else "—"
            creative_lines.append(f"     ✅ {label} ({date_str})")
        creative_block = "✅ Есть:\n" + "\n".join(creative_lines)
    else:
        creative_block = "❌ Нет"

    # ── Считаем % выполнения ───────────────────────────────────
    done = sum([
        bool(status["report"]),
        bool(status["plan"]),
        bool(status["creative"]),
    ])
    bar = "▓" * done + "░" * (3 - done)
    percent = done * 100 // 3

    return (
        f"📊 <b>Статус документов</b>\n"
        f"🏢 Клуб: <b>{club_name}</b>\n"
        f"📅 Период: <b>{month_year}</b>\n"
        f"\n"
        f"[{bar}] {percent}%\n"
        f"\n"
        f"📊 <b>Отчёт:</b> {report_line}\n"
        f"\n"
        f"📋 <b>План:</b> {plan_line}\n"
        f"\n"
        f"🎨 <b>Творческие отчёты:</b> {creative_block}"
    )


def _status_months_keyboard(club_name: str) -> InlineKeyboardMarkup:
    """Клавиатура выбора месяца для /status.
    callback_data: status_m_{МЕСЯЦ_ИНДЕКС}_{ГОД}_{КЛУБ_ИНДЕКС}
    Индексы вместо строк — чтобы не превышать лимит Telegram в 64 байта."""
    now = datetime.now()
    month_names = [
        "январь","февраль","март","апрель","май","июнь",
        "июль","август","сентябрь","октябрь","ноябрь","декабрь",
    ]
    all_clubs = sorted(CLUBS_DATA.keys())
    club_idx = all_clubs.index(club_name) if club_name in all_clubs else 0
    builder = InlineKeyboardBuilder()
    for delta in range(3):
        # Корректный переход на предыдущий месяц
        month_num = now.month - delta
        year = now.year
        while month_num < 1:
            month_num += 12
            year -= 1
        label = f"{month_names[month_num - 1].capitalize()} {year}"
        builder.add(InlineKeyboardButton(
            text=label,
            callback_data=f"status_m_{month_num}_{year}_{club_idx}",
        ))
    builder.adjust(1)
    return builder.as_markup()


@dp.message(Command("status"))
async def cmd_status(message: types.Message):
    """
    Показывает статус сданных документов за текущий месяц.

    Логика определения клуба:
    1. Если пользователь зарегистрирован через /register — берём его клуб.
    2. Если это администратор — показываем выбор клуба.
    3. Иначе — предлагаем зарегистрироваться.
    """
    user_id = message.from_user.id
    now = datetime.now()
    month_names = [
        "январь","февраль","март","апрель","май","июнь",
        "июль","август","сентябрь","октябрь","ноябрь","декабрь",
    ]
    month_year = f"{month_names[now.month - 1]} {now.year}"

    club_name = notif_get_club_for_user(user_id)

    if club_name:
        # Специалист — сразу показываем его клуб + кнопки смены месяца
        text = _format_status_text(club_name, month_year)
        builder = InlineKeyboardBuilder()
        _all_clubs = sorted(CLUBS_DATA.keys())
        _club_idx = _all_clubs.index(club_name) if club_name in _all_clubs else 0
        builder.add(InlineKeyboardButton(
            text="📅 Другой месяц",
            callback_data=f"status_cm_{_club_idx}",
        ))
        builder.adjust(1)
        await message.answer(text, parse_mode="HTML", reply_markup=builder.as_markup())

    elif user_id in ADMIN_IDS:
        # Администратор — предлагает выбрать клуб
        clubs = sorted(CLUBS_DATA.keys())
        builder = InlineKeyboardBuilder()
        for club in clubs:
            builder.add(InlineKeyboardButton(
                text=club,
                callback_data=f"status_club|{club}",
            ))
        builder.adjust(2)
        await message.answer(
            "📊 <b>Статус документов</b>\n\nВыберите клуб:",
            parse_mode="HTML",
            reply_markup=builder.as_markup(),
        )

    else:
        await message.answer(
            "ℹ️ Чтобы видеть статус своего клуба — зарегистрируйтесь:\n\n"
            "/register — привязать клуб к аккаунту"
        )


@dp.callback_query(F.data.startswith("status_club|"))
async def status_select_club(callback: types.CallbackQuery):
    """Администратор выбрал клуб — показываем статус за текущий месяц."""
    club_name = callback.data.removeprefix("status_club|")
    now = datetime.now()
    month_names = [
        "январь","февраль","март","апрель","май","июнь",
        "июль","август","сентябрь","октябрь","ноябрь","декабрь",
    ]
    month_year = f"{month_names[now.month - 1]} {now.year}"
    text = _format_status_text(club_name, month_year)

    builder = InlineKeyboardBuilder()
    _all_clubs2 = sorted(CLUBS_DATA.keys())
    _club_idx2 = _all_clubs2.index(club_name) if club_name in _all_clubs2 else 0
    builder.add(
        InlineKeyboardButton(text="📅 Другой месяц", callback_data=f"status_cm_{_club_idx2}"),
        InlineKeyboardButton(text="⬅️ Другой клуб",  callback_data="status_back_clubs"),
    )
    builder.adjust(1)
    await callback.message.edit_text(text, parse_mode="HTML", reply_markup=builder.as_markup())
    await callback.answer()


@dp.callback_query(F.data.startswith("status_cm_"))
async def status_choose_month(callback: types.CallbackQuery):
    """Показывает выбор месяца для текущего клуба."""
    all_clubs = sorted(CLUBS_DATA.keys())
    try:
        club_idx = int(callback.data.removeprefix("status_cm_"))
        club_name = all_clubs[club_idx]
    except (ValueError, IndexError):
        await callback.answer("❌ Ошибка", show_alert=True)
        return
    await callback.message.edit_text(
        f"📅 Выберите месяц для клуба <b>{club_name}</b>:",
        parse_mode="HTML",
        reply_markup=_status_months_keyboard(club_name),
    )
    await callback.answer()


@dp.callback_query(F.data.startswith("status_m_"))
async def status_select_month(callback: types.CallbackQuery):
    """Показывает статус за выбранный месяц.
    callback_data формат: status_m_{month_num}_{year}_{club_idx}"""
    month_names = ["январь","февраль","март","апрель","май","июнь",
                   "июль","август","сентябрь","октябрь","ноябрь","декабрь"]
    all_clubs = sorted(CLUBS_DATA.keys())
    try:
        parts = callback.data.removeprefix("status_m_").split("_")
        month_num, year, club_idx = int(parts[0]), int(parts[1]), int(parts[2])
        month_year = f"{month_names[month_num - 1]} {year}"
        club_name = all_clubs[club_idx]
    except (ValueError, IndexError):
        await callback.answer("❌ Ошибка формата", show_alert=True)
        return
    text = _format_status_text(club_name, month_year)

    builder = InlineKeyboardBuilder()
    _clubs_list = sorted(CLUBS_DATA.keys())
    _cidx = _clubs_list.index(club_name) if club_name in _clubs_list else 0
    builder.add(
        InlineKeyboardButton(text="📅 Другой месяц", callback_data=f"status_cm_{_cidx}"),
    )
    if callback.from_user.id in ADMIN_IDS:
        builder.add(InlineKeyboardButton(text="⬅️ Другой клуб", callback_data="status_back_clubs"))
    builder.adjust(1)
    await callback.message.edit_text(text, parse_mode="HTML", reply_markup=builder.as_markup())
    await callback.answer()


@dp.callback_query(F.data == "status_back_clubs")
async def status_back_clubs(callback: types.CallbackQuery):
    """Возврат к выбору клуба (для администратора)."""
    clubs = sorted(CLUBS_DATA.keys())
    builder = InlineKeyboardBuilder()
    for club in clubs:
        builder.add(InlineKeyboardButton(
            text=club,
            callback_data=f"status_club|{club}",
        ))
    builder.adjust(2)
    await callback.message.edit_text(
        "📊 <b>Статус документов</b>\n\nВыберите клуб:",
        parse_mode="HTML",
        reply_markup=builder.as_markup(),
    )
    await callback.answer()

@dp.message(Command("new_report"))
async def cmd_new_report(message: types.Message, state: FSMContext):
    """Начало заполнения нового отчета"""
    user_id = message.from_user.id
    
    template_path = "Шаблон отчета.docx"
    
    if not os.path.exists(template_path):
        await message.answer("❌ Шаблон отчета не найден.")
        return
    
    # Создаем начальный черновик
    draft = {
        'template_path': template_path,
        'events': [],
        'current_event_index': 0,
        'selected_dates': []
    }
    save_report_draft(user_id, draft)
    
    current_year = datetime.now().year
    
    await message.answer(
        f"📊 Заполнение отчета о работе молодежного клуба\n\n"
        f"Для отмены в любой момент нажмите /cancel\n\n"
        f"1️⃣ Выберите месяц отчета:",
        reply_markup=get_months_keyboard(current_year)
    )
    await state.set_state(ReportStates.waiting_for_month)

@dp.message(Command("new_plan"))
async def cmd_new_plan(message: types.Message, state: FSMContext):
    """Начало заполнения нового плана"""
    user_id = message.from_user.id
    
    template_path = "Шаблон плана.docx"
    
    if not os.path.exists(template_path):
        await message.answer("❌ Шаблон плана не найден.")
        return
    
    # Создаем начальный черновик
    draft = {
        'plan_items': [],
        'current_plan_index': 0
    }
    save_plan_draft(user_id, draft)
    
    current_year = datetime.now().year
    
    await message.answer(
        f"📋 Заполнение плана работы молодежного клуба\n\n"
        f"Для отмены в любой момент нажмите /cancel\n\n"
        f"1️⃣ Выберите месяц плана:",
        reply_markup=get_months_keyboard(current_year)
    )
    await state.set_state(PlanStates.waiting_for_plan_month)

# Резервные обработчики для текстового ввода месяца (на случай если кнопки не работают)
@dp.message(ReportStates.waiting_for_month)
async def process_month_year_fallback(message: types.Message, state: FSMContext):
    """Резервный обработчик для текстового ввода месяца"""
    user_id = message.from_user.id
    
    draft = get_report_draft(user_id) or {}
    draft['month_year'] = message.text
    save_report_draft(user_id, draft)
    
    await message.answer(
        "✅ Месяц и год сохранены!\n\n"
        "2️⃣ Выберите молодежный клуб:",
        reply_markup=get_clubs_keyboard()
    )
    await state.set_state(ReportStates.waiting_for_club_selection)

@dp.message(PlanStates.waiting_for_plan_month)
async def process_plan_month_year_fallback(message: types.Message, state: FSMContext):
    """Резервный обработчик для текстового ввода месяца для плана"""
    user_id = message.from_user.id
    
    draft = get_plan_draft(user_id) or {}
    draft['month_year'] = message.text
    save_plan_draft(user_id, draft)
    
    await message.answer(
        "✅ Месяц и год сохранены!\n\n"
        "2️⃣ Выберите молодежный клуб:",
        reply_markup=get_clubs_keyboard()
    )
    await state.set_state(PlanStates.waiting_for_plan_club_selection)

# Обработчики для ввода времени с клавиатурой
@dp.message(ReportStates.waiting_for_event_time)
async def process_event_time_fallback(message: types.Message, state: FSMContext):
    """Обработка текстового ввода времени с поддержкой кастомного формата"""
    user_id = message.from_user.id
    
    draft = get_report_draft(user_id)
    if not draft or not draft.get('events'):
        await message.answer("❌ Нет активного черновика. Начните заново: /new_report")
        await state.clear()
        return
    
    current_event_index = draft.get('current_event_index', 0)
    current_time_data = draft['events'][current_event_index].get('time_data', {})
    
    # Проверяем, что введен какой-то текст
    if not message.text.strip():
        await message.answer("❌ Введите время:")
        return
    
    # Проверяем длину текста (не более 50 символов)
    if len(message.text.strip()) > 50:
        await message.answer("❌ Слишком длинный текст. Введите время короче:")
        return
    
    input_text = message.text.strip()
    
    # Если уже выбрано время начала через кнопки
    if 'start_time' in current_time_data:
        start_time = current_time_data['start_time']
        
        # Проверяем формат введенного времени окончания
        time_pattern = re.compile(r'^\d{1,2}:\d{2}$')
        if not time_pattern.match(input_text):
            await message.answer("❌ Введите время окончания в формате ЧЧ:MM (например: 18:00):")
            return
        
        # Проверяем, что время окончания позже времени начала
        start_hour = int(start_time.split(':')[0])
        end_hour = int(input_text.split(':')[0])
        
        if end_hour <= start_hour:
            await message.answer("❌ Время окончания должно быть позже времени начала. Введите снова:")
            return
        
        # Формируем итоговое время
        final_time = f"{start_time}-{input_text}"
        draft['events'][current_event_index]['time'] = final_time
        draft['events'][current_event_index]['time_data'] = {}
        save_report_draft(user_id, draft)
        
        await message.answer(
            f"✅ Время сохранено: {final_time}\n\n"
            f"9️⃣ Введите краткое содержание мероприятия:"
        )
        await state.set_state(ReportStates.waiting_for_description)
        
    else:
        # Проверяем, введен ли сразу диапазон
        if '-' in input_text:
            # Пользователь ввел диапазон сразу
            times = input_text.split('-')
            if len(times) == 2:
                start_time = times[0].strip()
                end_time = times[1].strip()
                
                # Проверяем форматы
                time_pattern = re.compile(r'^\d{1,2}:\d{2}$')
                if time_pattern.match(start_time) and time_pattern.match(end_time):
                    start_hour = int(start_time.split(':')[0])
                    end_hour = int(end_time.split(':')[0])
                    
                    if end_hour > start_hour:
                        draft['events'][current_event_index]['time'] = input_text
                        draft['events'][current_event_index]['time_data'] = {}
                        save_report_draft(user_id, draft)
                        
                        await message.answer(
                            f"✅ Время сохранено: {input_text}\n\n"
                            f"9️⃣ Введите краткое содержание мероприятия:"
                        )
                        await state.set_state(ReportStates.waiting_for_description)
                        return
                    else:
                        await message.answer("❌ Время окончания должно быть позже времени начала. Введите снова:")
                        return
        
        # Если введено одно время или невалидный диапазон, сохраняем как время начала
        time_pattern = re.compile(r'^\d{1,2}:\d{2}$')
        if time_pattern.match(input_text):
            current_time_data['start_time'] = input_text
            draft['events'][current_event_index]['time_data'] = current_time_data
            save_report_draft(user_id, draft)
            
            await message.answer(
                f"✅ Выбрано время начала: {input_text}\n\n"
                f"Теперь введите время окончания в формате ЧЧ:MM:",
                reply_markup=get_time_keyboard(selected_start_time=input_text)
            )
        else:
            # Принимаем любой текст как валидное время (для текстовых описаний)
            draft['events'][current_event_index]['time'] = input_text
            draft['events'][current_event_index]['time_data'] = {}
            save_report_draft(user_id, draft)
            
            await message.answer(
                f"✅ Время сохранено: {input_text}\n\n"
                f"9️⃣ Введите краткое содержание мероприятия:"
            )
            await state.set_state(ReportStates.waiting_for_description)

# Обработчики для ввода ссылок с проверкой
@dp.message(ReportStates.waiting_for_link)
async def process_link_fallback(message: types.Message, state: FSMContext):
    """Обработка ввода ссылки с проверкой и поддержкой множественных ссылок"""
    user_id = message.from_user.id
    
    # Пропускаем прочерк без проверки
    if message.text != "-":
        # Разделяем ссылки по переносам строк
        links = [link.strip() for link in message.text.split('\n') if link.strip()]
        
        # Проверяем каждую ссылку
        for link in links:
            if not is_valid_url(link):
                await message.answer(f"❌ Ссылка '{link}' не является корректной. Введите валидные URL, каждый с новой строки, или '-' для прочерка:")
                return
    
    draft = get_report_draft(user_id)
    if not draft or not draft.get('events'):
        await message.answer("❌ Нет активного черновика. Начните заново: /new_report")
        await state.clear()
        return
    
    current_event_index = draft.get('current_event_index', 0)
    
    # Сохраняем все ссылки как одну строку с разделителем переноса строки
    if message.text == "-":
        draft['events'][current_event_index]['link'] = "-"
    else:
        # Сохраняем все ссылки через перенос строки
        draft['events'][current_event_index]['link'] = '\n'.join([link.strip() for link in message.text.split('\n') if link.strip()])
    
    save_report_draft(user_id, draft)
    
    link_display = "-" if message.text == "-" else f"{len(links)} ссылок"
    
    await message.answer(
        f"✅ Ссылки сохранены: {link_display}\n\n"
        f"👥 Введите количество участников 14-17 лет:\n(если нет — введите 0)"
    )
    await state.set_state(ReportStates.waiting_for_age_14_17)

@dp.message(PlanStates.waiting_for_info_link)
async def process_info_link_fallback(message: types.Message, state: FSMContext):
    """Обработка ввода информационной ссылки с проверкой"""
    user_id = message.from_user.id
    
    # Пропускаем прочерк без проверки
    if message.text != "-" and not is_valid_url(message.text):
        await message.answer("❌ Это не корректная ссылка. Введите валидный URL или '-' для прочерка:")
        return
    
    draft = get_plan_draft(user_id)
    if not draft or not draft.get('plan_items'):
        await message.answer("❌ Нет активного черновика. Начните заново: /new_plan")
        await state.clear()
        return 
    
    current_plan_index = draft.get('current_plan_index', 0)
    draft['plan_items'][current_plan_index]['info_link'] = message.text
    save_plan_draft(user_id, draft)
    
    plan_count = len(draft['plan_items'])
    current_item = draft['plan_items'][current_plan_index]
    
    link_display = "-" if message.text == "-" else message.text
    
    await message.answer(
        f"✅ Пункт плана #{plan_count} сохранен!\n\n"
        f"📋 Направление: {current_item['direction'][:100]}...\n"
        f"📝 Мероприятие: {current_item['event_name']}\n"
        f"🏷️ Статус: {current_item['status']}\n"
        f"📅 Дат выбрано: {len(current_item['dates'])}\n"
        f"📍 Место: {current_item['place']}\n"
        f"👥 Участников: {current_item['participants']} чел.\n"
        f"🔗 Ссылка: {link_display}\n\n"
        f"Выберите дальнейшее действие:",
        reply_markup=get_final_keyboard("plan")
    )
    await state.set_state(PlanStates.waiting_for_more_plan_items)

# Обработчики для ввода места с клавиатурой
@dp.message(ReportStates.waiting_for_event_place)
async def process_event_place(message: types.Message, state: FSMContext):
    """Обработка ввода места проведения"""
    user_id = message.from_user.id
    
    draft = get_report_draft(user_id)
    if not draft or not draft.get('events'):
        await message.answer("❌ Нет активного черновика. Начните заново: /new_report")
        await state.clear()
        return
    
    current_event_index = draft.get('current_event_index', 0)
    draft['events'][current_event_index]['place'] = message.text
    save_report_draft(user_id, draft)
    
    await message.answer(
        f"✅ Место сохранено: {message.text}\n\n"
        "8️⃣ Теперь выберите время начала мероприятия:",
        reply_markup=get_time_keyboard()
    )
    await state.set_state(ReportStates.waiting_for_event_time)

@dp.message(PlanStates.waiting_for_event_place)
async def process_plan_event_place(message: types.Message, state: FSMContext):
    """Обработка места проведения для плана"""
    user_id = message.from_user.id
    
    draft = get_plan_draft(user_id)
    if not draft or not draft.get('plan_items'):
        await message.answer("❌ Нет активного черновика. Начните заново: /new_plan")
        await state.clear()
        return
    
    current_plan_index = draft.get('current_plan_index', 0)
    draft['plan_items'][current_plan_index]['place'] = message.text
    save_plan_draft(user_id, draft)
    
    await message.answer(
        f"✅ Место сохранено: {message.text}\n\n"
        "👥 Теперь введите предполагаемое количество участников:"
    )
    await state.set_state(PlanStates.waiting_for_participants_count)

# Остальные обработчики остаются без изменений
@dp.message(ReportStates.waiting_for_event_name)
async def process_event_name(message: types.Message, state: FSMContext):
    """Обработка наименования мероприятия для отчета"""
    user_id = message.from_user.id
    
    draft = get_report_draft(user_id)
    if not draft:
        await message.answer("❌ Нет активного черновика. Начните заново: /new_report")
        await state.clear()
        return
    
    events = draft.get('events', [])
    current_event_index = draft.get('current_event_index', 0)
    
    if current_event_index < len(events):
        events[current_event_index]['event_name'] = message.text
    else:
        events.append({
            'event_name': message.text,
            'dates': [],
            'place': '',
            'time': '',
            'description': '',
            'link': '',
            'age_14_17': 0,
            'age_18_29': 0,
            'age_30_35': 0
        })
        draft['current_event_index'] = len(events) - 1
    
    draft['events'] = events
    save_report_draft(user_id, draft)
    
    month_num, year = parse_month_year(draft['month_year'])
    calendar_markup = get_month_calendar(year, month_num, draft.get('selected_dates', []))
    
    await message.answer(
        f"✅ Наименование мероприятия сохранено!\n\n"
        f"📅 Выберите даты проведения мероприятия:\n\n"
        f"📋 Выбранные даты:\nНет выбранных дат\n\n"
        f"🗓️ Календарь {draft['month_year']}:",
        reply_markup=calendar_markup
    )
    await state.set_state(ReportStates.waiting_for_event_dates)

@dp.message(PlanStates.waiting_for_event_name)
async def process_plan_event_name(message: types.Message, state: FSMContext):
    """Обработка наименования мероприятия для плана"""
    user_id = message.from_user.id
    
    draft = get_plan_draft(user_id)
    if not draft or not draft.get('plan_items'):
        await message.answer("❌ Нет активного черновика. Начните заново: /new_plan")
        await state.clear()
        return
    
    current_plan_index = draft.get('current_plan_index', 0)
    draft['plan_items'][current_plan_index]['event_name'] = message.text
    save_plan_draft(user_id, draft)
    
    await message.answer(
        f"✅ Название мероприятия сохранено!\n\n"
        f"📊 Теперь выберите статус мероприятия:",
        reply_markup=get_status_keyboard()
    )
    await state.set_state(PlanStates.waiting_for_event_status)

# Обработчики для плана
@dp.message(PlanStates.waiting_for_participants_count)
async def process_participants_count(message: types.Message, state: FSMContext):
    """Обработка количества участников для плана"""
    user_id = message.from_user.id
    
    draft = get_plan_draft(user_id)
    if not draft or not draft.get('plan_items'):
        await message.answer("❌ Нет активного черновика. Начните заново: /new_plan")
        await state.clear()
        return
    
    current_plan_index = draft.get('current_plan_index', 0)
    
    try:
        participants = int(message.text)
        if participants < 0:
            await message.answer("❌ Количество не может быть отрицательным.")
            return
        
        draft['plan_items'][current_plan_index]['participants'] = participants
        save_plan_draft(user_id, draft)
        
        await message.answer(
            "✅ Количество участников сохранено!\n\n"
            "🔗 Теперь выберите или введите ссылку на информационный ресурс:",
            reply_markup=get_standard_links_keyboard(for_plan=True)
        )
        await state.set_state(PlanStates.waiting_for_info_link)
        
    except ValueError:
        await message.answer("❌ Пожалуйста, введите корректное число:")

# Остальные обработчики для отчета
@dp.message(ReportStates.waiting_for_description)
async def process_description(message: types.Message, state: FSMContext):
    user_id = message.from_user.id
    
    draft = get_report_draft(user_id)
    if not draft or not draft.get('events'):
        await message.answer("❌ Нет активного черновика. Начните заново: /new_report")
        await state.clear()
        return
    
    current_event_index = draft.get('current_event_index', 0)
    draft['events'][current_event_index]['description'] = message.text
    save_report_draft(user_id, draft)
    
    await message.answer(
        "✅ Краткое содержание сохранено!\n\n"
        "🔗 Теперь введите ссылку на мероприятие:\n\n"
        "• Можно ввести несколько ссылок, каждую с новой строки\n"
        "• Или введите '-' для прочерка\n\n"
        "Пример:\nhttps://t.me/link1\nhttps://vk.com/link2",
        reply_markup=get_standard_links_keyboard(for_plan=False)  # Важно: for_plan=False
    )
    await state.set_state(ReportStates.waiting_for_link)

@dp.message(ReportStates.waiting_for_age_14_17)
async def process_age_14_17(message: types.Message, state: FSMContext):
    user_id = message.from_user.id
    
    draft = get_report_draft(user_id)
    if not draft or not draft.get('events'):
        await message.answer("❌ Нет активного черновика. Начните заново: /new_report")
        await state.clear()
        return
    
    current_event_index = draft.get('current_event_index', 0)
    
    try:
        age_14_17 = int(message.text)
        if age_14_17 < 0:
            await message.answer("❌ Количество не может быть отрицательным.")
            return
        
        draft['events'][current_event_index]['age_14_17'] = age_14_17
        save_report_draft(user_id, draft)
        
        await message.answer(
            "✅ Участники 14-17 лет сохранены!\n\n"
            "👥 Введите количество участников 18-29 лет:\n"
            "(если нет — введите 0)"
        )
        await state.set_state(ReportStates.waiting_for_age_18_29)
        
    except ValueError:
        await message.answer("❌ Пожалуйста, введите корректное число:")

@dp.message(ReportStates.waiting_for_age_18_29)
async def process_age_18_29(message: types.Message, state: FSMContext):
    user_id = message.from_user.id
    
    draft = get_report_draft(user_id)
    if not draft or not draft.get('events'):
        await message.answer("❌ Нет активного черновика. Начните заново: /new_report")
        await state.clear()
        return
    
    current_event_index = draft.get('current_event_index', 0)
    
    try:
        age_18_29 = int(message.text)
        if age_18_29 < 0:
            await message.answer("❌ Количество не может быть отрицательным.")
            return
        
        draft['events'][current_event_index]['age_18_29'] = age_18_29
        save_report_draft(user_id, draft)
        
        await message.answer(
            "✅ Участники 18-29 лет сохранены!\n\n"
            "👥 Введите количество участников 30-35 лет:\n"
            "(если нет — введите 0)"
        )
        await state.set_state(ReportStates.waiting_for_age_30_35)
        
    except ValueError:
        await message.answer("❌ Пожалуйста, введите корректное число:")

@dp.message(ReportStates.waiting_for_age_30_35)
async def process_age_30_35(message: types.Message, state: FSMContext):
    user_id = message.from_user.id
    
    draft = get_report_draft(user_id)
    if not draft or not draft.get('events'):
        await message.answer("❌ Нет активного черновика. Начните заново: /new_report")
        await state.clear()
        return
    
    current_event_index = draft.get('current_event_index', 0)
    
    try:
        age_30_35 = int(message.text)
        if age_30_35 < 0:
            await message.answer("❌ Количество не может быть отрицательным.")
            return
        
        draft['events'][current_event_index]['age_30_35'] = age_30_35
        save_report_draft(user_id, draft)
        
        event = draft['events'][current_event_index]
        total = event['age_14_17'] + event['age_18_29'] + event['age_30_35']
        dates_count = len(event['dates'])
        events_count = len(draft['events'])
        
        await message.answer(
            f"✅ Мероприятие #{events_count} сохранено!\n"
            f"📅 Количество выбранных дат: {dates_count}\n"
            f"📊 Общий охват: {total} человек\n\n"
            f"Выберите дальнейшее действие:",
            reply_markup=get_final_keyboard("report")
        )
        await state.set_state(ReportStates.waiting_for_more_events)
        
    except ValueError:
        await message.answer("❌ Пожалуйста, введите корректное число:")

# ═══════════════════════════════════════════════════════════════
# БЛОК: ТВОРЧЕСКИЕ ОТЧЁТЫ
# ═══════════════════════════════════════════════════════════════

# Папка с шаблонами творческих отчётов
CREATIVE_TEMPLATES_DIR = "Шаблоны творческих отчетов"

# Месяцы в родительном падеже (для подстановки в шаблон)
MONTH_GENITIVE = {
    "январь": "января", "февраль": "февраля", "март": "марта",
    "апрель": "апреля", "май": "мая", "июнь": "июня",
    "июль": "июля", "август": "августа", "сентябрь": "сентября",
    "октябрь": "октября", "ноябрь": "ноября", "декабрь": "декабря",
}


# ── Файловая система ───────────────────────────────────────────

def cr_get_clubs_with_templates() -> List[str]:
    """Возвращает клубы, у которых есть хотя бы один .docx-шаблон."""
    if not os.path.isdir(CREATIVE_TEMPLATES_DIR):
        return []
    result = []
    for name in sorted(os.listdir(CREATIVE_TEMPLATES_DIR)):
        club_dir = os.path.join(CREATIVE_TEMPLATES_DIR, name)
        if os.path.isdir(club_dir) and cr_list_templates(name):
            result.append(name)
    return result


def cr_list_templates(club_name: str) -> List[str]:
    """Список .docx файлов в папке клуба."""
    club_dir = os.path.join(CREATIVE_TEMPLATES_DIR, club_name)
    if not os.path.isdir(club_dir):
        return []
    return sorted(f for f in os.listdir(club_dir) if f.lower().endswith(".docx"))


def cr_template_label(filename: str) -> str:
    """'МТБ_даунхил.docx' → 'даунхил'."""
    name = os.path.splitext(filename)[0]
    parts = name.split("_", 1)
    return parts[1] if len(parts) > 1 else name


def cr_template_path(club_name: str, filename: str) -> str:
    return os.path.join(CREATIVE_TEMPLATES_DIR, club_name, filename)


# ── Работа с docx ─────────────────────────────────────────────

def cr_to_genitive(month: str) -> str:
    """'апрель' → 'апреля'."""
    return MONTH_GENITIVE.get(month.lower(), month)


def cr_calculate_total(a: str, b: str, c: str) -> int:
    """Суммирует возрастные группы; «-» = 0."""
    total = 0
    for v in (a, b, c):
        if v != "-":
            try:
                total += int(v)
            except ValueError:
                pass
    return total


def cr_replace_in_paragraph(paragraph, replacements: dict) -> None:
    """Заменяет плейсхолдеры в параграфе с сохранением форматирования."""
    full_text = "".join(r.text for r in paragraph.runs)
    changed = False
    for key, val in replacements.items():
        if key in full_text:
            full_text = full_text.replace(key, val)
            changed = True
    if changed and paragraph.runs:
        paragraph.runs[0].text = full_text
        for run in paragraph.runs[1:]:
            run.text = ""
    elif changed:
        paragraph.text = full_text


def cr_replace_in_doc(doc, replacements: dict) -> None:
    """Заменяет плейсхолдеры во всём документе (параграфы + таблицы)."""
    for para in doc.paragraphs:
        cr_replace_in_paragraph(para, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    cr_replace_in_paragraph(para, replacements)


def cr_replace_images(docx_bytes: bytes, photo_bytes_list: List[bytes]) -> bytes:
    """
    Заменяет первые N изображений в docx (передан как bytes) на фото пользователя.
    Работает через прямую ZIP-манипуляцию — безопасно для XML.
    """
    with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as zin:
        media_files = sorted(
            n for n in zin.namelist()
            if n.startswith("word/media/") and any(
                n.lower().endswith(ext) for ext in (".png", ".jpg", ".jpeg", ".gif", ".bmp")
            )
        )
        logger.info(f"[CR] Images in template: {media_files}")
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename in media_files:
                    idx = media_files.index(item.filename)
                    if idx < len(photo_bytes_list):
                        data = photo_bytes_list[idx]
                        logger.info(f"[CR] Replaced {item.filename} → photo #{idx+1}")
                zout.writestr(item, data)
    return buf.getvalue()


def cr_build_охват_phrase(total: int, age_14_17: str, age_18_29: str, age_30_35: str) -> str:
    """
    Строит финальную фразу охвата с учётом отсутствующих возрастных групп.

    Правила:
    - Первая присутствующая группа получает префикс «из них».
    - Если ни одна группа не указана — возвращает просто «N человек».
    - Пример (все три): «25 человек: из них 5 человек в возрасте от 14 до 17 лет,
                          10 человек в возрасте от 18 до 29 лет,
                          10 человек в возрасте от 30 до 35 лет»
    """
    age_parts = []

    if age_14_17 != "-":
        age_parts.append(f"из них {age_14_17} человек в возрасте от 14 до 17 лет")
    if age_18_29 != "-":
        prefix = "из них " if not age_parts else ""
        age_parts.append(f"{prefix}{age_18_29} человек в возрасте от 18 до 29 лет")
    if age_30_35 != "-":
        prefix = "из них " if not age_parts else ""
        age_parts.append(f"{prefix}{age_30_35} человек в возрасте от 30 до 35 лет")

    if not age_parts:
        return f"{total} человек"
    return f"{total} человек: {', '.join(age_parts)}"


def cr_fill_template(template_path: str, data: dict, photo_bytes_list: List[bytes]) -> str:
    """
    Заполняет шаблон творческого отчёта и возвращает путь к готовому .docx.

    Поля data: club_name, template_label, month, place, time,
               age_14_17, age_18_29, age_30_35, content, links,
               position, responsible

    Логика охвата:
    - В шаблоне ожидается строка вида:
        «КОЛИЧЕСТВО_ЛЮДЕЙ человек: из них 14_17 человек в возрасте от 14 до 17 лет,
         18_29 человек в возрасте от 18 до 29 лет, 30_35 человек в возрасте от 30 до 35 лет»
    - Вся эта фраза заменяется ОДНОЙ подстановкой (функция cr_build_охват_phrase).
      Это исключает конфликт между коротким плейсхолдером «14_17» и длинными фразами удаления.
    - Если точный шаблонный блок не нашёлся (другая редакция шаблона),
      применяются резервные замены: сначала длинные фразы, потом короткие числа.
    """
    with open(template_path, "rb") as fh:
        original_bytes = fh.read()

    # Заменяем изображения на уровне ZIP (до работы с текстом)
    modified_bytes = cr_replace_images(original_bytes, photo_bytes_list) if photo_bytes_list else original_bytes

    doc = Document(io.BytesIO(modified_bytes))

    month_gen  = cr_to_genitive(data["month"])
    age_14_17  = data["age_14_17"]
    age_18_29  = data["age_18_29"]
    age_30_35  = data["age_30_35"]
    total      = cr_calculate_total(age_14_17, age_18_29, age_30_35)

    # ── Строим финальную фразу охвата ─────────────────────────────────
    final_охват = cr_build_охват_phrase(total, age_14_17, age_18_29, age_30_35)

    # Точный шаблонный блок охвата (все три группы присутствуют в исходном шаблоне)
    TEMPLATE_ОХВАТ = (
        "КОЛИЧЕСТВО_ЛЮДЕЙ человек: из них 14_17 человек в возрасте от 14 до 17 лет, "
        "18_29 человек в возрасте от 18 до 29 лет, "
        "30_35 человек в возрасте от 30 до 35 лет"
    )

    # ── Словарь замен (ПОРЯДОК ВАЖЕН: длинные → короткие) ─────────────
    replacements = {
        # 1. Заменяем ВЕСЬ охватный блок одной подстановкой — главный путь
        TEMPLATE_ОХВАТ: final_охват,

        # 2. Остальные поля
        "МЕСЯЦ_ПРОВЕДЕНИЯ":       month_gen,
        "МЕСТО_ПРОВЕДЕНИЯ":       data["place"],
        "ВРЕМЯ_ПРОВЕДЕНИЯ":       data["time"],
        "СОДЕРЖАНИЕ_МЕРОПРИЯТИЯ": data["content"],
        "ССЫЛКА":                 data["links"],
        "ДОЛЖНОСТЬ":              data["position"],
        "ФИО":                    data["responsible"],

        # 3. Резервные замены на случай нестандартного шаблона:
        #    сначала длинные фразы удаления, потом короткий числовой плейсхолдер.
        #    Порядок внутри словаря гарантирует, что длинная фраза обрабатывается
        #    ДО того как будет заменён короткий плейсхолдер «14_17» и т.д.

        # Фразы удаления для 14-17
        "из них 14_17 человек в возрасте от 14 до 17 лет, ": "" if age_14_17 == "-" else None,
        ", из них 14_17 человек в возрасте от 14 до 17 лет": "" if age_14_17 == "-" else None,
        "из них 14_17 человек в возрасте от 14 до 17 лет":  "" if age_14_17 == "-" else None,

        # Фразы удаления для 18-29
        "18_29 человек в возрасте от 18 до 29 лет, ":       "" if age_18_29 == "-" else None,
        ", 18_29 человек в возрасте от 18 до 29 лет":       "" if age_18_29 == "-" else None,
        "18_29 человек в возрасте от 18 до 29 лет":         "" if age_18_29 == "-" else None,

        # Фразы удаления для 30-35
        "30_35 человек в возрасте от 30 до 35 лет, ":       "" if age_30_35 == "-" else None,
        ", 30_35 человек в возрасте от 30 до 35 лет":       "" if age_30_35 == "-" else None,
        "30_35 человек в возрасте от 30 до 35 лет":         "" if age_30_35 == "-" else None,

        # Короткие числовые плейсхолдеры — самые последние!
        "КОЛИЧЕСТВО_ЛЮДЕЙ": str(total),
        "14_17": age_14_17 if age_14_17 != "-" else "",
        "18_29": age_18_29 if age_18_29 != "-" else "",
        "30_35": age_30_35 if age_30_35 != "-" else "",
    }

    # Убираем записи с None (группа присутствует — фразу удалять не нужно)
    replacements = {k: v for k, v in replacements.items() if v is not None}

    cr_replace_in_doc(doc, replacements)

    # ── Имя выходного файла ────────────────────────────────────────────
    year_now   = datetime.now().year
    safe_label = "".join(
        c for c in data.get("template_label", "отчет")
        if c.isalnum() or c in (" ", "-", "_")
    ).strip()
    filename   = f"Творческий отчет за {data['month']} {year_now} {safe_label}.docx"
    output_path = os.path.join(TEMP_DIR, filename)
    os.makedirs(TEMP_DIR, exist_ok=True)
    doc.save(output_path)
    logger.info(f"[CR] Saved: {output_path}")
    return output_path


# ── Клавиатуры творческих отчётов ────────────────────────────

def cr_clubs_keyboard() -> InlineKeyboardMarkup:
    builder = InlineKeyboardBuilder()
    clubs = cr_get_clubs_with_templates()
    if not clubs:
        builder.add(InlineKeyboardButton(text="⚠️ Шаблоны не найдены", callback_data="cr_noop"))
    else:
        for i, club in enumerate(clubs):
            builder.add(InlineKeyboardButton(text=club, callback_data=f"cr_club_{i}"))
    builder.add(InlineKeyboardButton(text="❌ Отмена", callback_data="cr_cancel"))
    builder.adjust(2)
    return builder.as_markup()


def cr_templates_keyboard(club_name: str) -> InlineKeyboardMarkup:
    builder = InlineKeyboardBuilder()
    templates = cr_list_templates(club_name)
    if not templates:
        builder.add(InlineKeyboardButton(text="⚠️ Нет шаблонов", callback_data="cr_noop"))
    else:
        for i, tpl in enumerate(templates):
            label = cr_template_label(tpl)
            builder.add(InlineKeyboardButton(text=f"📄 {label}", callback_data=f"cr_tpl_{i}"))
    builder.add(InlineKeyboardButton(text="⬅️ К выбору клуба", callback_data="cr_back_clubs"))
    builder.adjust(1)
    return builder.as_markup()


def cr_months_keyboard() -> InlineKeyboardMarkup:
    builder = InlineKeyboardBuilder()
    months = ["январь","февраль","март","апрель","май","июнь",
              "июль","август","сентябрь","октябрь","ноябрь","декабрь"]
    for i, m in enumerate(months, start=1):
        builder.add(InlineKeyboardButton(text=m.capitalize(), callback_data=f"cr_month_{i}"))
    builder.add(InlineKeyboardButton(text="❌ Отмена", callback_data="cr_cancel"))
    builder.adjust(3)
    return builder.as_markup()


def cr_cancel_kb() -> InlineKeyboardMarkup:
    builder = InlineKeyboardBuilder()
    builder.add(InlineKeyboardButton(text="❌ Отмена", callback_data="cr_cancel"))
    return builder.as_markup()


def cr_progress(step: int) -> str:
    filled = "▓" * step
    empty  = "░" * (10 - step)
    return f"[{filled}{empty}] {step}/10"


def cr_validate_age(text: Optional[str]) -> Optional[str]:
    """Возвращает строку числа, «-» или None при ошибке."""
    if not text:
        return None
    text = text.strip()
    if text == "-":
        return "-"
    try:
        v = int(text)
        return str(v) if v >= 0 else None
    except ValueError:
        return None


# ── Обработчики ───────────────────────────────────────────────

@dp.callback_query(F.data == "start_creative_report")
async def cr_start(callback: types.CallbackQuery, state: FSMContext):
    """Вход в режим творческого отчёта. Если есть черновик — предлагаем продолжить."""
    clubs = cr_get_clubs_with_templates()
    if not clubs:
        await callback.message.edit_text(
            "❌ Папка «Шаблоны творческих отчетов» не найдена или пуста.\n\n"
            "Создайте структуру:\n"
            "📁 Шаблоны творческих отчетов/\n"
            "  📁 МТБ/\n"
            "    📄 МТБ_даунхил.docx"
        )
        await callback.answer()
        return

    # Проверяем незавершённый черновик
    if cr_draft_has(callback.from_user.id):
        draft = cr_draft_load(callback.from_user.id)
        if draft:
            draft_data, draft_step = draft
            club  = draft_data.get("club_name", "?")
            tmpl  = draft_data.get("template_label", "?")
            month = draft_data.get("month", "")
            step_labels = {
                "waiting_for_place":    "ввод места (шаг 4)",
                "waiting_for_time":     "ввод времени (шаг 5)",
                "waiting_for_age_14_17":"возраст 14–17 (шаг 6)",
                "waiting_for_age_18_29":"возраст 18–29 (шаг 7)",
                "waiting_for_age_30_35":"возраст 30–35 (шаг 8)",
                "waiting_for_content":  "содержание (шаг 9)",
                "waiting_for_links":    "ссылки (шаг 10)",
                "waiting_for_photos":   "фотографии (финал)",
            }
            step_display = step_labels.get(draft_step, draft_step)
            builder = InlineKeyboardBuilder()
            builder.add(
                InlineKeyboardButton(text="▶️ Продолжить с того же места", callback_data="cr_draft_resume"),
                InlineKeyboardButton(text="🗑 Начать заново",               callback_data="cr_draft_discard"),
            )
            builder.adjust(1)
            await state.set_state(CreativeReportStates.waiting_for_draft_choice)
            await callback.message.edit_text(
                f"📝 Найден незавершённый черновик:\n\n"
                f"🏢 Клуб: {club}\n"
                f"📄 Шаблон: {tmpl}\n"
                f"📅 Месяц: {month.capitalize() if month else '—'}\n"
                f"⏸ Остановился на: {step_display}\n\n"
                "Продолжить или начать заново?",
                reply_markup=builder.as_markup(),
            )
            await callback.answer()
            return

    await state.set_state(CreativeReportStates.waiting_for_club)
    await callback.message.edit_text(
        f"🎨 Творческий отчёт\n\n{cr_progress(1)}\nШаг 1 из 10 — выберите клуб:",
        reply_markup=cr_clubs_keyboard(),
    )
    await callback.answer()


@dp.callback_query(CreativeReportStates.waiting_for_draft_choice, F.data == "cr_draft_resume")
async def cr_draft_resume(callback: types.CallbackQuery, state: FSMContext):
    """Восстанавливает черновик и продолжает с нужного шага."""
    draft = cr_draft_load(callback.from_user.id)
    if not draft:
        await callback.answer("❌ Черновик не найден, начинаем заново", show_alert=True)
        await state.set_state(CreativeReportStates.waiting_for_club)
        await callback.message.edit_text(
            f"🎨 Творческий отчёт\n\n{cr_progress(1)}\nШаг 1 из 10 — выберите клуб:",
            reply_markup=cr_clubs_keyboard(),
        )
        await callback.answer()
        return

    draft_data, current_step = draft
    # Восстанавливаем данные в state (без фото — их запросим заново)
    await state.update_data(**draft_data, photos=[])

    # Переводим в нужное состояние и показываем нужный шаг
    step_to_state = {
        "waiting_for_place":    (CreativeReportStates.waiting_for_place,    4),
        "waiting_for_time":     (CreativeReportStates.waiting_for_time,     5),
        "waiting_for_age_14_17":(CreativeReportStates.waiting_for_age_14_17,6),
        "waiting_for_age_18_29":(CreativeReportStates.waiting_for_age_18_29,7),
        "waiting_for_age_30_35":(CreativeReportStates.waiting_for_age_30_35,8),
        "waiting_for_content":  (CreativeReportStates.waiting_for_content,  9),
        "waiting_for_links":    (CreativeReportStates.waiting_for_links,    10),
        "waiting_for_photos":   (CreativeReportStates.waiting_for_photos,   10),
    }
    step_prompts = {
        "waiting_for_place":     "введите место проведения:\n\nПример: г. Сочи, ул. Навагинская, 9",
        "waiting_for_time":      "введите время проведения:\n\nПример: с 9:00 до 15:00",
        "waiting_for_age_14_17": "введите количество участников 14–17 лет\n(или «-» если нет):",
        "waiting_for_age_18_29": "введите количество участников 18–29 лет\n(или «-» если нет):",
        "waiting_for_age_30_35": "введите количество участников 30–35 лет\n(или «-» если нет):",
        "waiting_for_content":   "опишите содержание мероприятия:",
        "waiting_for_links":     "введите ссылки на освещение (или «-»):",
        "waiting_for_photos":    "отправьте 4 фотографии с мероприятия.\n\nПрогресс: <b>0 / 4</b> фото",
    }

    fsm_state, step_num = step_to_state.get(
        current_step,
        (CreativeReportStates.waiting_for_place, 4),
    )
    prompt = step_prompts.get(current_step, "продолжите заполнение:")
    await state.set_state(fsm_state)

    month = draft_data.get("month", "")
    club  = draft_data.get("club_name", "")
    tmpl  = draft_data.get("template_label", "")

    await callback.message.edit_text(
        f"▶️ Продолжаем черновик\n\n"
        f"🏢 {club} / 📄 {tmpl} / 📅 {month.capitalize()}\n\n"
        f"{cr_progress(step_num)}\nШаг {step_num} из 10 — {prompt}",
        reply_markup=cr_cancel_kb(),
        parse_mode="HTML",
    )
    await callback.answer()


@dp.callback_query(CreativeReportStates.waiting_for_draft_choice, F.data == "cr_draft_discard")
async def cr_draft_discard(callback: types.CallbackQuery, state: FSMContext):
    """Удаляет черновик и начинает заново."""
    cr_draft_delete(callback.from_user.id)
    await state.set_state(CreativeReportStates.waiting_for_club)
    await callback.message.edit_text(
        f"🗑 Черновик удалён.\n\n🎨 Творческий отчёт\n\n{cr_progress(1)}\nШаг 1 из 10 — выберите клуб:",
        reply_markup=cr_clubs_keyboard(),
    )
    await callback.answer()


@dp.message(Command("creative_report"))
async def cr_cmd(message: types.Message, state: FSMContext):
    """Команда /creative_report."""
    clubs = cr_get_clubs_with_templates()
    if not clubs:
        await message.answer("❌ Папка с шаблонами не найдена или пуста.")
        return
    await state.set_state(CreativeReportStates.waiting_for_club)
    await message.answer(
        f"🎨 Творческий отчёт\n\n{cr_progress(1)}\nШаг 1 из 10 — выберите клуб:",
        reply_markup=cr_clubs_keyboard(),
    )


@dp.callback_query(CreativeReportStates.waiting_for_club, F.data.startswith("cr_club_"))
async def cr_select_club(callback: types.CallbackQuery, state: FSMContext):
    """Выбор клуба."""
    idx = int(callback.data.removeprefix("cr_club_"))
    clubs = cr_get_clubs_with_templates()
    if idx >= len(clubs):
        await callback.answer("❌ Клуб не найден", show_alert=True)
        return
    club_name = clubs[idx]
    templates = cr_list_templates(club_name)
    if not templates:
        await callback.message.edit_text(
            f"⚠️ Для клуба «{club_name}» шаблоны не найдены.\n"
            "Добавьте .docx файлы в папку клуба.",
            reply_markup=cr_clubs_keyboard(),
        )
        await callback.answer()
        return
    await state.update_data(club_name=club_name)
    await state.set_state(CreativeReportStates.waiting_for_template)
    await callback.message.edit_text(
        f"🏢 Клуб: {club_name}\n\n{cr_progress(2)}\nШаг 2 из 10 — выберите шаблон:",
        reply_markup=cr_templates_keyboard(club_name),
    )
    await callback.answer()


@dp.callback_query(CreativeReportStates.waiting_for_template, F.data.startswith("cr_tpl_"))
async def cr_select_template(callback: types.CallbackQuery, state: FSMContext):
    """Выбор шаблона."""
    idx = int(callback.data.removeprefix("cr_tpl_"))
    data = await state.get_data()
    club_name = data["club_name"]
    templates = cr_list_templates(club_name)
    if idx >= len(templates):
        await callback.answer("❌ Шаблон не найден", show_alert=True)
        return
    filename = templates[idx]
    tpl_path = cr_template_path(club_name, filename)
    if not os.path.exists(tpl_path):
        await callback.answer("❌ Файл не найден на диске", show_alert=True)
        return
    label = cr_template_label(filename)
    await state.update_data(template_path=tpl_path, template_label=label, photos=[])
    await state.set_state(CreativeReportStates.waiting_for_month)
    await callback.message.edit_text(
        f"🏢 Клуб: {club_name}\n📄 Шаблон: {label}\n\n"
        f"{cr_progress(3)}\nШаг 3 из 10 — выберите месяц проведения:",
        reply_markup=cr_months_keyboard(),
    )
    await callback.answer()


@dp.callback_query(CreativeReportStates.waiting_for_month, F.data.startswith("cr_month_"))
async def cr_select_month(callback: types.CallbackQuery, state: FSMContext):
    """Выбор месяца. Если за этот месяц по данному шаблону уже есть отчёт —
    отправляет его пользователю и предлагает перезаполнить или оставить."""
    _months = ["январь","февраль","март","апрель","май","июнь",
               "июль","август","сентябрь","октябрь","ноябрь","декабрь"]
    month_idx = int(callback.data.removeprefix("cr_month_")) - 1
    month = _months[month_idx]
    data = await state.get_data()
    club_name: str    = data["club_name"]
    template_label: str = data["template_label"]
    year_now = datetime.now().year
    month_year = f"{month} {year_now}"

    await state.update_data(month=month, month_year=month_year)

    # ── Проверяем, существует ли уже отчёт ────────────────────────
    existing = get_existing_creative_report(club_name, template_label, month_year)

    if existing:
        file_path, created_at = existing
        # Отправляем существующий файл (если файл на диске)
        file_path = resolve_file_path(file_path) or file_path
        if os.path.exists(file_path):
            try:
                await callback.message.answer_document(
                    FSInputFile(file_path),
                    caption=(
                        f"📄 Найден существующий творческий отчёт:\n\n"
                        f"🏢 Клуб: {club_name}\n"
                        f"📝 Шаблон: {template_label}\n"
                        f"📅 Месяц: {month.capitalize()} {year_now}\n"
                        f"🕐 Создан: {created_at}"
                    ),
                )
            except Exception as exc:
                logger.warning(f"[CR] Could not send existing file: {exc}")

        # Предлагаем выбор
        builder = InlineKeyboardBuilder()
        builder.add(
            InlineKeyboardButton(
                text="✏️ Перезаполнить (старый будет удалён)",
                callback_data="cr_overwrite_yes",
            ),
            InlineKeyboardButton(
                text="✅ Оставить как есть",
                callback_data="cr_overwrite_no",
            ),
        )
        builder.adjust(1)

        await state.set_state(CreativeReportStates.waiting_for_overwrite_confirm)
        await callback.message.answer(
            f"❓ Творческий отчёт «{template_label}» за {month.capitalize()} {year_now} "
            f"для клуба «{club_name}» уже сформирован.\n\n"
            "Хотите заполнить его заново?",
            reply_markup=builder.as_markup(),
        )
    else:
        # Отчёта нет — начинаем заполнение обычным путём
        await state.set_state(CreativeReportStates.waiting_for_place)
        await callback.message.edit_text(
            f"✅ Месяц: {month.capitalize()}\n\n"
            f"{cr_progress(4)}\nШаг 4 из 10 — введите место проведения:\n\n"
            "Пример: г. Сочи, ул. Навагинская, 9",
            reply_markup=cr_cancel_kb(),
        )

    await callback.answer()


@dp.callback_query(CreativeReportStates.waiting_for_overwrite_confirm, F.data == "cr_overwrite_yes")
async def cr_overwrite_yes(callback: types.CallbackQuery, state: FSMContext):
    """Пользователь подтвердил перезапись — удаляем старую запись и начинаем заново."""
    data = await state.get_data()
    club_name      = data["club_name"]
    template_label = data["template_label"]
    month_year     = data["month_year"]
    month          = data["month"]

    # Удаляем старый отчёт из БД и с диска
    delete_creative_report_from_db(club_name, template_label, month_year)

    await state.set_state(CreativeReportStates.waiting_for_place)
    await callback.message.edit_text(
        f"🗑 Старый отчёт удалён. Заполняем заново.\n\n"
        f"✅ Месяц: {month.capitalize()}\n\n"
        f"{cr_progress(4)}\nШаг 4 из 10 — введите место проведения:\n\n"
        "Пример: г. Сочи, ул. Навагинская, 9",
        reply_markup=cr_cancel_kb(),
    )
    await callback.answer()


@dp.callback_query(CreativeReportStates.waiting_for_overwrite_confirm, F.data == "cr_overwrite_no")
async def cr_overwrite_no(callback: types.CallbackQuery, state: FSMContext):
    """Пользователь решил оставить существующий отчёт."""
    await state.clear()
    await callback.message.edit_text(
        "✅ Существующий отчёт сохранён без изменений.\n\n"
        "Используйте /start для возврата в меню."
    )
    await callback.answer()


@dp.message(CreativeReportStates.waiting_for_place)
async def cr_enter_place(message: types.Message, state: FSMContext):
    """Ввод места проведения."""
    if not message.text or not message.text.strip():
        await message.answer("❌ Введите место проведения:")
        return
    await state.update_data(place=message.text.strip())
    cr_draft_save(message.from_user.id, await state.get_data(), "waiting_for_time")
    await state.set_state(CreativeReportStates.waiting_for_time)
    await message.answer(
        f"✅ Место: {message.text.strip()}\n\n"
        f"{cr_progress(5)}\nШаг 5 из 10 — введите время проведения:\n\n"
        "Пример: с 9:00 до 15:00",
        reply_markup=cr_cancel_kb(),
    )


@dp.message(CreativeReportStates.waiting_for_time)
async def cr_enter_time(message: types.Message, state: FSMContext):
    """Ввод времени проведения."""
    if not message.text or not message.text.strip():
        await message.answer("❌ Введите время:")
        return
    await state.update_data(time=message.text.strip())
    cr_draft_save(message.from_user.id, await state.get_data(), "waiting_for_age_14_17")
    await state.set_state(CreativeReportStates.waiting_for_age_14_17)
    await message.answer(
        f"✅ Время: {message.text.strip()}\n\n"
        f"{cr_progress(6)}\nШаг 6 из 10 — участники 14–17 лет.\n\n"
        "Введите число или <b>-</b> если таких не было.",
        parse_mode="HTML",
        reply_markup=cr_cancel_kb(),
    )


@dp.message(CreativeReportStates.waiting_for_age_14_17)
async def cr_enter_age_14_17(message: types.Message, state: FSMContext):
    """Ввод участников 14-17."""
    val = cr_validate_age(message.text)
    if val is None:
        await message.answer("❌ Введите целое число ≥ 0 или «-»:")
        return
    await state.update_data(age_14_17=val)
    cr_draft_save(message.from_user.id, await state.get_data(), "waiting_for_age_18_29")
    await state.set_state(CreativeReportStates.waiting_for_age_18_29)
    display = "нет" if val == "-" else val
    await message.answer(
        f"✅ Участники 14–17 лет: {display}\n\n"
        f"{cr_progress(7)}\nШаг 7 из 10 — участники 18–29 лет.\n\n"
        "Введите число или <b>-</b>.",
        parse_mode="HTML",
        reply_markup=cr_cancel_kb(),
    )


@dp.message(CreativeReportStates.waiting_for_age_18_29)
async def cr_enter_age_18_29(message: types.Message, state: FSMContext):
    """Ввод участников 18-29."""
    val = cr_validate_age(message.text)
    if val is None:
        await message.answer("❌ Введите целое число ≥ 0 или «-»:")
        return
    await state.update_data(age_18_29=val)
    cr_draft_save(message.from_user.id, await state.get_data(), "waiting_for_age_30_35")
    await state.set_state(CreativeReportStates.waiting_for_age_30_35)
    display = "нет" if val == "-" else val
    await message.answer(
        f"✅ Участники 18–29 лет: {display}\n\n"
        f"{cr_progress(8)}\nШаг 8 из 10 — участники 30–35 лет.\n\n"
        "Введите число или <b>-</b>.",
        parse_mode="HTML",
        reply_markup=cr_cancel_kb(),
    )


@dp.message(CreativeReportStates.waiting_for_age_30_35)
async def cr_enter_age_30_35(message: types.Message, state: FSMContext):
    """Ввод участников 30-35. Проверяет что хоть одна группа заполнена."""
    val = cr_validate_age(message.text)
    if val is None:
        await message.answer("❌ Введите целое число ≥ 0 или «-»:")
        return
    data = await state.get_data()
    if data.get("age_14_17") == "-" and data.get("age_18_29") == "-" and val == "-":
        await message.answer("❌ Хотя бы одна возрастная группа должна быть заполнена.")
        return
    await state.update_data(age_30_35=val)
    cr_draft_save(message.from_user.id, await state.get_data(), "waiting_for_content")
    total = cr_calculate_total(data.get("age_14_17", "-"), data.get("age_18_29", "-"), val)
    await state.set_state(CreativeReportStates.waiting_for_content)
    await message.answer(
        f"✅ Участники 30–35 лет: {'нет' if val == '-' else val}\n"
        f"📊 Итого участников: {total}\n\n"
        f"{cr_progress(9)}\nШаг 9 из 10 — содержание мероприятия.\n\n"
        "Подробно опишите, что происходило на мероприятии:",
        reply_markup=cr_cancel_kb(),
    )


@dp.message(CreativeReportStates.waiting_for_content)
async def cr_enter_content(message: types.Message, state: FSMContext):
    """Ввод содержания мероприятия."""
    if not message.text or len(message.text.strip()) < 5:
        await message.answer("❌ Описание слишком короткое. Напишите подробнее:")
        return
    await state.update_data(content=message.text.strip())
    cr_draft_save(message.from_user.id, await state.get_data(), "waiting_for_links")
    await state.set_state(CreativeReportStates.waiting_for_links)
    await message.answer(
        "✅ Содержание сохранено!\n\n"
        f"{cr_progress(10)}\nШаг 10 из 10 — ссылки на информационное освещение.\n\n"
        "• Несколько ссылок — каждую с новой строки\n"
        "• Если ссылок нет — отправьте: <b>-</b>\n\n"
        "Пример:\nhttps://t.me/molod_sochi\nhttps://vk.com/post123",
        parse_mode="HTML",
        reply_markup=cr_cancel_kb(),
    )


@dp.message(CreativeReportStates.waiting_for_links)
async def cr_enter_links(message: types.Message, state: FSMContext):
    """Ввод ссылок на освещение."""
    raw = (message.text or "").strip()
    if not raw:
        await message.answer("❌ Введите ссылки или «-»:")
        return
    links_text = "-" if raw == "-" else "\n".join(
        line.strip() for line in raw.splitlines() if line.strip()
    )
    await state.update_data(links=links_text)
    cr_draft_save(message.from_user.id, await state.get_data(), "waiting_for_photos")
    await state.set_state(CreativeReportStates.waiting_for_photos)
    await message.answer(
        "✅ Ссылки сохранены!\n\n"
        "📸 Отправьте 4 фотографии с мероприятия.\n\n"
        "Отправляйте по одному фото — бот посчитает сам.\n"
        "Прогресс: <b>0 / 4</b> фото",
        parse_mode="HTML",
        reply_markup=cr_cancel_kb(),
    )


@dp.message(CreativeReportStates.waiting_for_photos, F.photo)
async def cr_collect_photos(message: types.Message, state: FSMContext):
    """Сбор 4 фотографий с мероприятия."""
    data = await state.get_data()
    photos: list = data.get("photos", [])

    if len(photos) >= 4:
        await message.answer("✅ Уже получено 4 фото. Формирую документ...")
        await cr_generate_and_send(message, state)
        return

    # Скачиваем фото в максимальном разрешении
    photo = message.photo[-1]
    file_info = await bot.get_file(photo.file_id)
    downloaded = await bot.download_file(file_info.file_path)
    photos.append(downloaded.read())
    await state.update_data(photos=photos)

    received = len(photos)
    if received < 4:
        await message.answer(
            f"✅ Фото <b>{received}/4</b> получено. "
            f"Осталось: {4 - received} фото.",
            parse_mode="HTML",
        )
    else:
        await message.answer("✅ Все 4 фото получены! Формирую творческий отчёт...")
        await cr_generate_and_send(message, state)


@dp.message(CreativeReportStates.waiting_for_photos, ~F.photo)
async def cr_wrong_content(message: types.Message, state: FSMContext):
    """Пользователь прислал не фото во время сбора фото."""
    data = await state.get_data()
    received = len(data.get("photos", []))
    await message.answer(
        f"❌ Нужна <b>фотография</b> (не документ/файл).\n"
        f"Получено: {received}/4 фото.",
        parse_mode="HTML",
    )


async def cr_generate_and_send(message: types.Message, state: FSMContext):
    """Генерирует и отправляет готовый творческий отчёт."""
    data = await state.get_data()
    await state.clear()

    club_name = data.get("club_name", "")
    club_info = CLUBS_DATA.get(club_name, {})

    report_data = {
        "club_name":      club_name,
        "template_label": data.get("template_label", "отчет"),
        "month":          data.get("month", ""),
        "place":          data.get("place", ""),
        "time":           data.get("time", ""),
        "age_14_17":      data.get("age_14_17", "-"),
        "age_18_29":      data.get("age_18_29", "-"),
        "age_30_35":      data.get("age_30_35", "-"),
        "content":        data.get("content", ""),
        "links":          data.get("links", "-"),
        "position":       club_info.get("position", "Специалист"),
        "responsible":    club_info.get("responsible", ""),
    }

    try:
        output_path = cr_fill_template(
            template_path=data["template_path"],
            data=report_data,
            photo_bytes_list=data.get("photos", []),
        )
        total = cr_calculate_total(
            report_data["age_14_17"],
            report_data["age_18_29"],
            report_data["age_30_35"],
        )
        photos_count = len(data.get("photos", []))

        # Удаляем черновик — отчёт успешно создан
        cr_draft_delete(message.from_user.id)

        doc_file = FSInputFile(output_path)
        await message.answer_document(
            doc_file,
            caption=(
                f"✅ Творческий отчёт сформирован!\n\n"
                f"🏢 Клуб: {club_name}\n"
                f"📄 Тема: {report_data['template_label']}\n"
                f"📅 Месяц: {report_data['month'].capitalize()}\n"
                f"📍 Место: {report_data['place']}\n"
                f"⏰ Время: {report_data['time']}\n"
                f"👥 Участников: {total}\n"
                f"📸 Фото вставлено: {photos_count}/4\n\n"
                f"👤 {report_data['position']} {report_data['responsible']}"
            ),
        )

        # Меню редактирования полей
        builder = InlineKeyboardBuilder()
        builder.add(
            InlineKeyboardButton(text="🔗 Изменить ссылки",      callback_data="cr_edit_links"),
            InlineKeyboardButton(text="📍 Изменить место",        callback_data="cr_edit_place"),
            InlineKeyboardButton(text="⏰ Изменить время",        callback_data="cr_edit_time"),
            InlineKeyboardButton(text="📝 Изменить содержание",   callback_data="cr_edit_content"),
            InlineKeyboardButton(text="👥 Изменить участников",   callback_data="cr_edit_ages"),
            InlineKeyboardButton(text="✅ Всё верно, завершить",  callback_data="cr_edit_done"),
        )
        builder.adjust(1)
        await message.answer(
            "Хотите что-то изменить в отчёте?\n"
            "Выберите поле — бот перегенерирует документ с новыми данными.",
            reply_markup=builder.as_markup(),
        )

        # Уведомление администраторов
        await send_document_to_admin(
            message.from_user.id,
            output_path,
            "Творческий отчёт",
            club_name,
            report_data["month"],
            1,
            total,
        )

        # Сохраняем запись в БД (нужна для «Скачать документы за период»)
        # month_year берём из state — он уже задан при выборе месяца
        month_year_str = data.get("month_year") or f"{report_data['month']} {datetime.now().year}"
        save_creative_report_to_db(
            user_id=message.from_user.id,
            club_name=club_name,
            month_year=month_year_str,
            template_label=report_data["template_label"],
            file_path=output_path,
            total_participants=total,
        )

    except Exception as exc:
        logger.error(f"[CR] Error generating report: {exc}", exc_info=True)
        await message.answer(
            f"❌ Ошибка при формировании документа:\n<code>{exc}</code>\n\n"
            "Проверьте шаблон и попробуйте снова: /creative_report",
            parse_mode="HTML",
        )


@dp.callback_query(F.data == "cr_edit_done")
async def cr_edit_done(callback: types.CallbackQuery, state: FSMContext):
    """Пользователь доволен отчётом — завершаем."""
    await callback.message.edit_text("✅ Отчёт сохранён. Используйте /start для возврата в меню.")
    await callback.answer()


# Словарь: callback_data → (поле в state, метка для пользователя, подсказка)
CR_EDIT_FIELDS = {
    "cr_edit_links":   ("links",   "ссылки",        "Отправьте новые ссылки (каждую с новой строки) или «-» если ссылок нет:"),
    "cr_edit_place":   ("place",   "место",         "Введите новый адрес места проведения:"),
    "cr_edit_time":    ("time",    "время",          "Введите новое время проведения (например: с 9:00 до 15:00):"),
    "cr_edit_content": ("content", "содержание",    "Введите новое описание содержания мероприятия:"),
    "cr_edit_ages":    ("ages",    "участников",    "Введите через запятую: 14–17, 18–29, 30–35\nПример: 5, 10, 8\n(«-» если группы нет, например: -, 10, 8)"),
}


@dp.callback_query(F.data.startswith("cr_edit_") & ~F.data.in_({"cr_edit_done"}))
async def cr_edit_field_prompt(callback: types.CallbackQuery, state: FSMContext):
    """Запрашивает новое значение выбранного поля."""
    action = callback.data  # e.g. "cr_edit_links"
    if action not in CR_EDIT_FIELDS:
        await callback.answer()
        return

    field_key, field_label, prompt = CR_EDIT_FIELDS[action]
    await state.update_data(_editing_field=field_key, _editing_action=action)
    await state.set_state(CreativeReportStates.editing_field)

    builder = InlineKeyboardBuilder()
    builder.add(InlineKeyboardButton(text="❌ Отмена", callback_data="cr_edit_cancel"))
    await callback.message.answer(
        f"✏️ Редактирование: <b>{field_label}</b>\n\n{prompt}",
        parse_mode="HTML",
        reply_markup=builder.as_markup(),
    )
    await callback.answer()


@dp.message(CreativeReportStates.editing_field)
async def cr_edit_field_apply(message: types.Message, state: FSMContext):
    """Принимает новое значение, обновляет state, перегенерирует отчёт."""
    data = await state.get_data()
    field_key = data.get("_editing_field")

    if not field_key:
        await state.clear()
        return

    raw = (message.text or "").strip()
    if not raw:
        await message.answer("❌ Введите значение или нажмите «Отмена».")
        return

    # Обрабатываем специальные поля
    if field_key == "links":
        new_value = "-" if raw == "-" else "\n".join(
            line.strip() for line in raw.splitlines() if line.strip()
        )
        await state.update_data(links=new_value)

    elif field_key == "ages":
        # Формат: «5, 10, 8» или «-, 10, 8»
        parts = [p.strip() for p in raw.split(",")]
        if len(parts) != 3:
            await message.answer(
                "❌ Введите три значения через запятую.\n"
                "Пример: 5, 10, 8 (или «-» для пропуска группы, например: -, 10, 8)"
            )
            return
        validated = []
        for p in parts:
            v = cr_validate_age(p)
            if v is None:
                await message.answer(f"❌ Неверное значение «{p}». Введите число ≥ 0 или «-».")
                return
            validated.append(v)
        if all(v == "-" for v in validated):
            await message.answer("❌ Хотя бы одна группа должна быть заполнена.")
            return
        await state.update_data(age_14_17=validated[0], age_18_29=validated[1], age_30_35=validated[2])

    else:
        # place, time, content — просто обновляем
        await state.update_data(**{field_key: raw})

    # Перегенерируем документ с новыми данными
    await state.set_state(None)  # Снимаем состояние editing_field
    await message.answer("🔄 Перегенерирую отчёт с новыми данными...")

    updated_data = await state.get_data()
    club_name  = updated_data.get("club_name", "")
    club_info  = CLUBS_DATA.get(club_name, {})

    report_data = {
        "club_name":      club_name,
        "template_label": updated_data.get("template_label", "отчет"),
        "month":          updated_data.get("month", ""),
        "place":          updated_data.get("place", ""),
        "time":           updated_data.get("time", ""),
        "age_14_17":      updated_data.get("age_14_17", "-"),
        "age_18_29":      updated_data.get("age_18_29", "-"),
        "age_30_35":      updated_data.get("age_30_35", "-"),
        "content":        updated_data.get("content", ""),
        "links":          updated_data.get("links", "-"),
        "position":       club_info.get("position", "Специалист"),
        "responsible":    club_info.get("responsible", ""),
    }

    try:
        output_path = cr_fill_template(
            template_path=updated_data["template_path"],
            data=report_data,
            photo_bytes_list=updated_data.get("photos", []),
        )
        total = cr_calculate_total(report_data["age_14_17"], report_data["age_18_29"], report_data["age_30_35"])

        # Обновляем запись в БД
        month_year_str = updated_data.get("month_year") or f"{report_data['month']} {datetime.now().year}"
        delete_creative_report_from_db(club_name, report_data["template_label"], month_year_str)
        save_creative_report_to_db(
            user_id=message.from_user.id,
            club_name=club_name,
            month_year=month_year_str,
            template_label=report_data["template_label"],
            file_path=output_path,
            total_participants=total,
        )
        await state.update_data(template_path=updated_data["template_path"])

        doc_file = FSInputFile(output_path)
        await message.answer_document(
            doc_file,
            caption=(
                f"✅ Отчёт обновлён!\n\n"
                f"🏢 Клуб: {club_name}\n"
                f"📅 Месяц: {report_data['month'].capitalize()}\n"
                f"📍 Место: {report_data['place']}\n"
                f"⏰ Время: {report_data['time']}\n"
                f"👥 Участников: {total}\n"
            ),
        )

        # Снова показываем меню редактирования
        builder = InlineKeyboardBuilder()
        builder.add(
            InlineKeyboardButton(text="🔗 Изменить ссылки",      callback_data="cr_edit_links"),
            InlineKeyboardButton(text="📍 Изменить место",        callback_data="cr_edit_place"),
            InlineKeyboardButton(text="⏰ Изменить время",        callback_data="cr_edit_time"),
            InlineKeyboardButton(text="📝 Изменить содержание",   callback_data="cr_edit_content"),
            InlineKeyboardButton(text="👥 Изменить участников",   callback_data="cr_edit_ages"),
            InlineKeyboardButton(text="✅ Всё верно, завершить",  callback_data="cr_edit_done"),
        )
        builder.adjust(1)
        await message.answer(
            "Хотите изменить что-то ещё?",
            reply_markup=builder.as_markup(),
        )

    except Exception as exc:
        logger.error(f"[CR] Error re-generating report: {exc}", exc_info=True)
        await message.answer(f"❌ Ошибка при обновлении отчёта:\n<code>{exc}</code>", parse_mode="HTML")


@dp.callback_query(F.data == "cr_edit_cancel")
async def cr_edit_cancel(callback: types.CallbackQuery, state: FSMContext):
    """Отмена редактирования поля — возврат к меню редактирования."""
    await state.set_state(None)
    builder = InlineKeyboardBuilder()
    builder.add(
        InlineKeyboardButton(text="🔗 Изменить ссылки",      callback_data="cr_edit_links"),
        InlineKeyboardButton(text="📍 Изменить место",        callback_data="cr_edit_place"),
        InlineKeyboardButton(text="⏰ Изменить время",        callback_data="cr_edit_time"),
        InlineKeyboardButton(text="📝 Изменить содержание",   callback_data="cr_edit_content"),
        InlineKeyboardButton(text="👥 Изменить участников",   callback_data="cr_edit_ages"),
        InlineKeyboardButton(text="✅ Всё верно, завершить",  callback_data="cr_edit_done"),
    )
    builder.adjust(1)
    await callback.message.edit_text(
        "Редактирование отменено. Хотите изменить что-то ещё?",
        reply_markup=builder.as_markup(),
    )
    await callback.answer()


@dp.callback_query(F.data == "cr_cancel")
async def cr_cancel(callback: types.CallbackQuery, state: FSMContext):
    """Отмена творческого отчёта."""
    await state.clear()
    await callback.message.edit_text(
        "❌ Заполнение творческого отчёта отменено.\n"
        "Используйте /start для возврата в меню."
    )
    await callback.answer()


@dp.callback_query(F.data == "cr_back_clubs")
async def cr_back_clubs(callback: types.CallbackQuery, state: FSMContext):
    """Возврат к выбору клуба."""
    await state.set_state(CreativeReportStates.waiting_for_club)
    await callback.message.edit_text(
        f"🎨 Творческий отчёт\n\n{cr_progress(1)}\nШаг 1 из 10 — выберите клуб:",
        reply_markup=cr_clubs_keyboard(),
    )
    await callback.answer()


@dp.callback_query(F.data == "cr_noop")
async def cr_noop(callback: types.CallbackQuery):
    """Заглушка для неактивных кнопок."""
    await callback.answer("Нет доступных шаблонов", show_alert=True)


# ═══════════════════════════════════════════════════════════════
# КОНЕЦ БЛОКА: ТВОРЧЕСКИЕ ОТЧЁТЫ
# ═══════════════════════════════════════════════════════════════

# ═══════════════════════════════════════════════════════════════
# БЛОК: УВЕДОМЛЕНИЯ И НАПОМИНАНИЯ
# ═══════════════════════════════════════════════════════════════

# ── DB-утилиты ────────────────────────────────────────────────

def notif_get_settings() -> dict:
    """Возвращает текущие настройки расписания напоминаний."""
    conn = sqlite3.connect('reports_plans.db')
    cursor = conn.cursor()
    cursor.execute('SELECT settings_json FROM reminder_settings WHERE id = 1')
    row = cursor.fetchone()
    conn.close()
    if row:
        return json.loads(row[0])
    return {}


def notif_save_settings(settings: dict) -> None:
    conn = sqlite3.connect('reports_plans.db')
    cursor = conn.cursor()
    cursor.execute(
        '''INSERT INTO reminder_settings (id, settings_json, updated_at)
           VALUES (1, ?, CURRENT_TIMESTAMP)
           ON CONFLICT(id) DO UPDATE SET
               settings_json = excluded.settings_json,
               updated_at = CURRENT_TIMESTAMP''',
        (json.dumps(settings, ensure_ascii=False),),
    )
    conn.commit()
    conn.close()


def notif_register_specialist(user_id: int, club_name: str, username: str, full_name: str) -> None:
    conn = sqlite3.connect('reports_plans.db')
    cursor = conn.cursor()
    cursor.execute(
        '''INSERT INTO specialist_users (user_id, club_name, username, full_name)
           VALUES (?, ?, ?, ?)
           ON CONFLICT(user_id) DO UPDATE SET
               club_name = excluded.club_name,
               username  = excluded.username,
               full_name = excluded.full_name''',
        (user_id, club_name, username or "", full_name or ""),
    )
    conn.commit()
    conn.close()


def notif_get_all_specialists() -> List[Tuple]:
    """Возвращает [(user_id, club_name, full_name, notifications_enabled), ...]."""
    conn = sqlite3.connect('reports_plans.db')
    cursor = conn.cursor()
    cursor.execute(
        'SELECT user_id, club_name, full_name, notifications_enabled FROM specialist_users'
    )
    rows = cursor.fetchall()
    conn.close()
    return rows


def notif_toggle(user_id: int, enabled: bool) -> None:
    conn = sqlite3.connect('reports_plans.db')
    cursor = conn.cursor()
    cursor.execute(
        'UPDATE specialist_users SET notifications_enabled = ? WHERE user_id = ?',
        (1 if enabled else 0, user_id),
    )
    conn.commit()
    conn.close()


def notif_get_club_for_user(user_id: int) -> Optional[str]:
    conn = sqlite3.connect('reports_plans.db')
    cursor = conn.cursor()
    cursor.execute('SELECT club_name FROM specialist_users WHERE user_id = ?', (user_id,))
    row = cursor.fetchone()
    conn.close()
    return row[0] if row else None


def notif_missing_docs(month_year: str) -> Tuple[List[str], List[str], List[str]]:
    """
    Возвращает три списка клубов за month_year:
    (без_отчёта, без_плана, без_творческого_отчёта)
    """
    all_clubs = set(CLUBS_DATA.keys())
    conn = sqlite3.connect('reports_plans.db')
    cursor = conn.cursor()

    cursor.execute('SELECT DISTINCT club_name FROM reports WHERE month_year = ?', (month_year,))
    with_reports = {r[0] for r in cursor.fetchall()}

    cursor.execute('SELECT DISTINCT club_name FROM plans WHERE month_year = ?', (month_year,))
    with_plans = {r[0] for r in cursor.fetchall()}

    cursor.execute('SELECT DISTINCT club_name FROM creative_reports WHERE month_year = ?', (month_year,))
    with_creative = {r[0] for r in cursor.fetchall()}

    conn.close()
    return (
        sorted(all_clubs - with_reports),
        sorted(all_clubs - with_plans),
        sorted(all_clubs - with_creative),
    )


# ── Логика отправки уведомлений ───────────────────────────────

async def notif_send_specialist_reminders(send_bot: Bot) -> None:
    """
    Отправляет каждому зарегистрированному специалисту напоминание
    о том, какие документы его клуб ещё не сдал в текущем месяце.
    """
    now = datetime.now()
    month_names = [
        "январь","февраль","март","апрель","май","июнь",
        "июль","август","сентябрь","октябрь","ноябрь","декабрь",
    ]
    month_year = f"{month_names[now.month - 1]} {now.year}"
    missing_reports, missing_plans, missing_creative = notif_missing_docs(month_year)

    specialists = notif_get_all_specialists()
    sent = 0

    for user_id, club_name, full_name, enabled in specialists:
        if not enabled:
            continue

        needs_report   = club_name in missing_reports
        needs_plan     = club_name in missing_plans
        needs_creative = club_name in missing_creative

        if not (needs_report or needs_plan or needs_creative):
            # Клуб всё сдал — отправляем похвалу
            try:
                await send_bot.send_message(
                    user_id,
                    f"✅ {full_name or 'Привет'}!\n\n"
                    f"Клуб <b>{club_name}</b> сдал все документы за {month_year}. "
                    "Отличная работа! 🎉",
                    parse_mode="HTML",
                )
                sent += 1
            except Exception as exc:
                logger.warning(f"[NOTIF] Cannot send to {user_id}: {exc}")
            continue

        lines = [
            f"📋 <b>Напоминание о сдаче документов</b>\n",
            f"👋 {full_name or 'Привет'}!",
            f"🏢 Клуб: <b>{club_name}</b>",
            f"📅 Период: <b>{month_year}</b>\n",
            "Следующие документы ещё не сданы:\n",
        ]
        if needs_report:
            lines.append("📊 Ежемесячный отчёт")
        if needs_plan:
            lines.append("📋 План на месяц")
        if needs_creative:
            lines.append("🎨 Творческий отчёт")

        lines += [
            "",
            "Пожалуйста, заполните их до конца месяца.",
            "Нажмите /start чтобы начать.",
        ]

        try:
            await send_bot.send_message(user_id, "\n".join(lines), parse_mode="HTML")
            sent += 1
        except Exception as exc:
            logger.warning(f"[NOTIF] Cannot send to {user_id}: {exc}")

    logger.info(f"[NOTIF] Specialist reminders sent: {sent}/{len(specialists)}")


async def notif_send_admin_summary(send_bot: Bot) -> None:
    """
    Отправляет каждому администратору итоговую сводку
    по несданным документам за текущий месяц.
    """
    now = datetime.now()
    month_names = [
        "январь","февраль","март","апрель","май","июнь",
        "июль","август","сентябрь","октябрь","ноябрь","декабрь",
    ]
    month_year = f"{month_names[now.month - 1]} {now.year}"
    missing_reports, missing_plans, missing_creative = notif_missing_docs(month_year)

    total = len(CLUBS_DATA)
    ok_reports   = total - len(missing_reports)
    ok_plans     = total - len(missing_plans)
    ok_creative  = total - len(missing_creative)

    lines = [
        f"📊 <b>Итоговая сводка за {month_year}</b>\n",
        f"🏢 Всего клубов: {total}\n",
        f"📊 Отчёты сданы:           {ok_reports}/{total}",
        f"📋 Планы сданы:            {ok_plans}/{total}",
        f"🎨 Творч. отчёты сданы:    {ok_creative}/{total}\n",
    ]

    if missing_reports:
        lines.append("❌ <b>Не сдали отчёт:</b>")
        lines += [f"  • {c}" for c in missing_reports]
        lines.append("")

    if missing_plans:
        lines.append("❌ <b>Не сдали план:</b>")
        lines += [f"  • {c}" for c in missing_plans]
        lines.append("")

    if missing_creative:
        lines.append("❌ <b>Нет творческого отчёта:</b>")
        lines += [f"  • {c}" for c in missing_creative]
        lines.append("")

    if not missing_reports and not missing_plans and not missing_creative:
        lines.append("✅ <b>Все клубы сдали все документы!</b>")

    text = "\n".join(lines)
    for admin_id in ADMIN_IDS:
        try:
            await send_bot.send_message(admin_id, text, parse_mode="HTML")
        except Exception as exc:
            logger.warning(f"[NOTIF] Cannot send admin summary to {admin_id}: {exc}")

    logger.info(f"[NOTIF] Admin summary sent for {month_year}")


# ── Планировщик ───────────────────────────────────────────────

async def notification_scheduler(send_bot: Bot) -> None:
    """
    Фоновая задача — каждую минуту сверяет текущее время с настройками
    и запускает нужное уведомление. Не требует сторонних библиотек.
    """
    logger.info("[NOTIF] Scheduler started")
    last_specialist_day: Optional[int] = None
    last_admin_day:      Optional[int] = None

    while True:
        await asyncio.sleep(60)  # проверяем раз в минуту
        try:
            now      = datetime.now()
            settings = notif_get_settings()

            s_day  = settings.get("specialist_day",   20)
            s_hour = settings.get("specialist_hour",  10)
            s_min  = settings.get("specialist_minute", 0)

            a_day  = settings.get("admin_day",   25)
            a_hour = settings.get("admin_hour",  10)
            a_min  = settings.get("admin_minute", 0)

            # Специалистам
            if (now.day == s_day and now.hour == s_hour and now.minute == s_min
                    and last_specialist_day != (now.year, now.month, now.day)):
                logger.info("[NOTIF] Triggering specialist reminders")
                await notif_send_specialist_reminders(send_bot)
                last_specialist_day = (now.year, now.month, now.day)

            # Администратору
            if (now.day == a_day and now.hour == a_hour and now.minute == a_min
                    and last_admin_day != (now.year, now.month, now.day)):
                logger.info("[NOTIF] Triggering admin summary")
                await notif_send_admin_summary(send_bot)
                last_admin_day = (now.year, now.month, now.day)

        except Exception as exc:
            logger.error(f"[NOTIF] Scheduler error: {exc}", exc_info=True)


# ── Регистрация специалиста (/register) ───────────────────────

@dp.message(Command("register"))
async def cmd_register(message: types.Message, state: FSMContext):
    """Специалист регистрирует свой клуб для получения напоминаний."""
    clubs = sorted(CLUBS_DATA.keys())
    builder = InlineKeyboardBuilder()
    for i, club in enumerate(clubs):
        builder.add(InlineKeyboardButton(
            text=club,
            callback_data=f"notif_reg_{i}",
        ))
    builder.adjust(2)
    await message.answer(
        "👋 Регистрация для получения напоминаний.\n\n"
        "Выберите ваш клуб:",
        reply_markup=builder.as_markup(),
    )


@dp.callback_query(F.data.startswith("notif_reg_"))
async def notif_reg_club(callback: types.CallbackQuery):
    """Сохраняет привязку пользователя к клубу."""
    all_clubs_sorted = sorted(CLUBS_DATA.keys())
    try:
        club_idx = int(callback.data.removeprefix("notif_reg_"))
        club_name = all_clubs_sorted[club_idx]
    except (ValueError, IndexError):
        await callback.answer("❌ Клуб не найден", show_alert=True)
        return
    if club_name not in CLUBS_DATA:
        await callback.answer("❌ Клуб не найден", show_alert=True)
        return

    user = callback.from_user
    notif_register_specialist(
        user_id=user.id,
        club_name=club_name,
        username=user.username or "",
        full_name=user.full_name or "",
    )
    await callback.message.edit_text(
        f"✅ Вы зарегистрированы!\n\n"
        f"🏢 Клуб: <b>{club_name}</b>\n"
        f"👤 {user.full_name}\n\n"
        "Теперь бот будет присылать вам напоминания о сдаче документов.\n\n"
        "Управление уведомлениями:\n"
        "🔕 /notifications_off — отключить\n"
        "🔔 /notifications_on — включить",
        parse_mode="HTML",
    )
    await callback.answer()


@dp.message(Command("notifications_off"))
async def cmd_notif_off(message: types.Message):
    club = notif_get_club_for_user(message.from_user.id)
    if not club:
        await message.answer("❌ Вы не зарегистрированы. Используйте /register.")
        return
    notif_toggle(message.from_user.id, False)
    await message.answer("🔕 Уведомления отключены. Включить: /notifications_on")


@dp.message(Command("notifications_on"))
async def cmd_notif_on(message: types.Message):
    club = notif_get_club_for_user(message.from_user.id)
    if not club:
        await message.answer("❌ Вы не зарегистрированы. Используйте /register.")
        return
    notif_toggle(message.from_user.id, True)
    await message.answer("🔔 Уведомления включены.")


# ── Панель управления напоминаниями для администратора ────────

@dp.callback_query(F.data == "admin_notifications")
async def admin_notifications_menu(callback: types.CallbackQuery):
    """Меню управления напоминаниями в панели администратора."""
    if callback.from_user.id not in ADMIN_IDS:
        await callback.answer("❌ Доступ запрещён", show_alert=True)
        return

    settings = notif_get_settings()
    specialists = notif_get_all_specialists()
    enabled_count = sum(1 for _, _, _, en in specialists if en)

    builder = InlineKeyboardBuilder()
    builder.add(
        InlineKeyboardButton(text="📅 Изменить расписание",     callback_data="admin_notif_schedule"),
        InlineKeyboardButton(text="🔔 Отправить сейчас (спец)", callback_data="admin_notif_send_spec"),
        InlineKeyboardButton(text="📊 Отправить сводку (адм.)",  callback_data="admin_notif_send_admin"),
        InlineKeyboardButton(text="👥 Список специалистов",      callback_data="admin_notif_list"),
        InlineKeyboardButton(text="⬅️ Назад",                    callback_data="start_admin"),
    )
    builder.adjust(1)

    await callback.message.edit_text(
        f"🔔 <b>Управление уведомлениями</b>\n\n"
        f"📅 Напоминание специалистам: <b>{settings.get('specialist_day', 20)}-е число</b> "
        f"в {settings.get('specialist_hour', 10):02d}:{settings.get('specialist_minute', 0):02d}\n"
        f"📊 Сводка администратору:    <b>{settings.get('admin_day', 25)}-е число</b> "
        f"в {settings.get('admin_hour', 10):02d}:{settings.get('admin_minute', 0):02d}\n\n"
        f"👥 Зарегистрированных специалистов: {len(specialists)} "
        f"(активных: {enabled_count})",
        reply_markup=builder.as_markup(),
        parse_mode="HTML",
    )
    await callback.answer()


@dp.callback_query(F.data == "admin_notif_list")
async def admin_notif_list(callback: types.CallbackQuery):
    """Показывает список зарегистрированных специалистов."""
    if callback.from_user.id not in ADMIN_IDS:
        await callback.answer("❌ Доступ запрещён", show_alert=True)
        return

    specialists = notif_get_all_specialists()
    if not specialists:
        await callback.answer("Специалистов ещё нет", show_alert=True)
        return

    lines = ["👥 <b>Зарегистрированные специалисты:</b>\n"]
    for user_id, club_name, full_name, enabled in sorted(specialists, key=lambda x: x[1]):
        status = "🔔" if enabled else "🔕"
        lines.append(f"{status} <b>{club_name}</b> — {full_name or '—'}")

    builder = InlineKeyboardBuilder()
    builder.add(InlineKeyboardButton(text="⬅️ Назад", callback_data="admin_notifications"))
    await callback.message.edit_text(
        "\n".join(lines),
        reply_markup=builder.as_markup(),
        parse_mode="HTML",
    )
    await callback.answer()


@dp.callback_query(F.data == "admin_notif_send_spec")
async def admin_notif_send_spec(callback: types.CallbackQuery):
    """Немедленная отправка напоминаний специалистам."""
    if callback.from_user.id not in ADMIN_IDS:
        await callback.answer("❌ Доступ запрещён", show_alert=True)
        return
    await callback.answer("⏳ Отправляю напоминания...")
    await notif_send_specialist_reminders(bot)
    await callback.message.answer("✅ Напоминания специалистам отправлены.")


@dp.callback_query(F.data == "admin_notif_send_admin")
async def admin_notif_send_admin(callback: types.CallbackQuery):
    """Немедленная отправка сводки администраторам."""
    if callback.from_user.id not in ADMIN_IDS:
        await callback.answer("❌ Доступ запрещён", show_alert=True)
        return
    await callback.answer("⏳ Отправляю сводку...")
    await notif_send_admin_summary(bot)


@dp.callback_query(F.data == "admin_notif_schedule")
async def admin_notif_schedule(callback: types.CallbackQuery, state: FSMContext):
    """Показывает форму изменения расписания."""
    if callback.from_user.id not in ADMIN_IDS:
        await callback.answer("❌ Доступ запрещён", show_alert=True)
        return

    settings = notif_get_settings()
    await state.update_data(_notif_editing=True)
    await state.set_state(AdminStates.waiting_for_admin_action)

    await callback.message.answer(
        "📅 <b>Изменение расписания</b>\n\n"
        "Отправьте новое расписание в формате:\n"
        "<code>специалисты: ДД ЧЧ:ММ\nадмин: ДД ЧЧ:ММ</code>\n\n"
        "Пример:\n"
        "<code>специалисты: 20 10:00\nадмин: 25 10:00</code>\n\n"
        f"Текущее:\n"
        f"специалисты: {settings.get('specialist_day',20)} "
        f"{settings.get('specialist_hour',10):02d}:{settings.get('specialist_minute',0):02d}\n"
        f"админ: {settings.get('admin_day',25)} "
        f"{settings.get('admin_hour',10):02d}:{settings.get('admin_minute',0):02d}",
        parse_mode="HTML",
    )
    await callback.answer()


@dp.message(AdminStates.waiting_for_admin_action, F.text.regexp(r'специалисты:\s*\d+\s+\d+:\d+'))
async def admin_notif_schedule_save(message: types.Message, state: FSMContext):
    """Сохраняет новое расписание напоминаний."""
    if message.from_user.id not in ADMIN_IDS:
        return

    data = await state.get_data()
    if not data.get("_notif_editing"):
        return

    await state.update_data(_notif_editing=False)

    import re as _re
    spec_m = _re.search(r'специалисты:\s*(\d+)\s+(\d+):(\d+)', message.text)
    admin_m = _re.search(r'админ:\s*(\d+)\s+(\d+):(\d+)', message.text)

    settings = notif_get_settings()

    if spec_m:
        settings["specialist_day"]    = int(spec_m.group(1))
        settings["specialist_hour"]   = int(spec_m.group(2))
        settings["specialist_minute"] = int(spec_m.group(3))

    if admin_m:
        settings["admin_day"]    = int(admin_m.group(1))
        settings["admin_hour"]   = int(admin_m.group(2))
        settings["admin_minute"] = int(admin_m.group(3))

    notif_save_settings(settings)

    await message.answer(
        f"✅ Расписание обновлено!\n\n"
        f"📅 Специалистам: {settings['specialist_day']}-е число "
        f"в {settings['specialist_hour']:02d}:{settings['specialist_minute']:02d}\n"
        f"📊 Сводка админу: {settings['admin_day']}-е число "
        f"в {settings['admin_hour']:02d}:{settings['admin_minute']:02d}",
    )

# ═══════════════════════════════════════════════════════════════
# КОНЕЦ БЛОКА: УВЕДОМЛЕНИЯ
# ═══════════════════════════════════════════════════════════════


# ВАЖНО: этот обработчик ДОЛЖЕН быть последним — он перехватывает
# все сообщения, не обработанные выше (catch-all).
@dp.message()
async def handle_other_messages(message: types.Message, state: FSMContext):
    is_admin = message.from_user.id in ADMIN_IDS
    builder = InlineKeyboardBuilder()
    builder.add(
        InlineKeyboardButton(text="📊 Заполнить отчет", callback_data="start_new_report"),
        InlineKeyboardButton(text="📋 Заполнить план", callback_data="start_new_plan"),
        InlineKeyboardButton(text="📈 Квартальный отчет", callback_data="start_new_quarterly"),
        InlineKeyboardButton(text="🎨 Творческий отчет", callback_data="start_creative_report"),
    )
    if is_admin:
        builder.add(InlineKeyboardButton(text="⚙️ Панель администратора", callback_data="start_admin"))
    builder.adjust(1)
    await message.answer(
        "Используйте кнопки для начала работы или /cancel для отмены текущей операции.",
        reply_markup=builder.as_markup()
    )


async def main():
    # Инициализация базы данных
    init_database()

    # Запуск фоновых задач
    asyncio.create_task(notification_scheduler(bot))

    logger.info("Бот запущен")
    await dp.start_polling(bot)

if __name__ == "__main__":
    import asyncio
    asyncio.run(main())